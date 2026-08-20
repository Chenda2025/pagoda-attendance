import time
import secrets
from collections import defaultdict
from flask import request, render_template, Blueprint, jsonify, send_file, session, redirect, url_for, abort
from conn import connect_db
from create_table import (create_monks_table, insert_monk, get_all_monks,
                          update_monk, delete_monk, insert_pending_submission,
                          update_monk_living_status)
from datetime import date as _date
from sorting import sort_attendance_monks, ROLE_RANK
from khmer_lunar import khmer_lunar_date
from auth_service import (
    get_user_by_username, get_user_by_id, list_users, create_user, update_user,
    delete_user, save_face_enrollment, touch_last_login, find_user_by_face,
    user_allowed, user_home, log_activity, list_activity, purge_old_activity,
    verify_password, unlock_user, register_failed_password, bind_device,
    device_is_trusted, trusted_devices,
    MODULE_LABELS, ALL_PERMISSIONS, ACTIVITY_RETENTION_DAYS,
    MAX_PASSWORD_ATTEMPTS, LOCK_REASON_PASSWORD,
)

# ============ MINISTRY DOCUMENT HEADER (shared by docx/html/excel exports) ============

_HDR_LEFT_LINES  = ['មន្ទីរធម្មការ និងសាសនា រាជធានី', 'ភ្នំពេញ', 'សាលា ពុ.អ.វិ.ស.ព្រ.ន.វ.និ']
_HDR_RIGHT_LINES = ['ព្រះរាជាណាចក្រកម្ពុជា', 'ជាតិ សាសនា ព្រះមហាក្សត្រ']
_HDR_MAIN_TITLE  = 'ស្ថិតព្រះសង្ឃក្នុងវត្តនិរោធរង្សី'

main_bp = Blueprint('main_bp', __name__)

from form_options_service import (
    FIELD_LABELS,
    list_grouped,
    valid_set,
    add_option as add_form_option,
    delete_option as delete_form_option,
    update_option as update_form_option,
)

# ============ PUBLIC SUBMISSION — ALLOWED VALUES (from DB) ============

def _valid_monk_types():
    return valid_set('monk_type')

def _valid_residences():
    return valid_set('residence')

def _valid_positions():
    return valid_set('position')

def _valid_education_levels():
    return valid_set('education_level')

def _valid_academic_years():
    return valid_set('academic_year')

def _validate_entry_dropdowns(monk_type, residence, position, education_level, academic_year):
    errors = []
    if monk_type not in _valid_monk_types():
        errors.append('ប្រភេទព្រះសង្ឃ មិនត្រឹមត្រូវ')
    if residence not in _valid_residences():
        errors.append('ស្នាក់នៅកុដិ មិនត្រឹមត្រូវ')
    if position not in _valid_positions():
        errors.append('តួនាទី មិនត្រឹមត្រូវ')
    if education_level not in _valid_education_levels():
        errors.append('កម្រិតសិក្សា មិនត្រឹមត្រូវ')
    if academic_year not in _valid_academic_years():
        errors.append('ឆ្នាំសិក្សា មិនត្រឹមត្រូវ')
    return errors

LIVING_STATUSES = (
    'កំពុងស្នាក់នៅ',
    'ឈប់ស្នាក់នៅ',
    'នៅស្រុក',
    'ឈឺនៅពេទ្យ',
    'ឈឺនៅស្រុក',
)
_VALID_LIVING_STATUSES  = set(LIVING_STATUSES)
_ACTIVE_LIVING_STATUS   = 'កំពុងស្នាក់នៅ'

# Discipline / contract thresholds (bi-weekly report rules)
DISC_ABSENT_MIN = 2   # absences  >= 2
DISC_PERM_MIN   = 3   # permissions >= 3


def _contract_eligible(absent_count, perm_count):
    """True when monk exceeds absent or permission criteria (for contract / Telegram)."""
    return absent_count >= DISC_ABSENT_MIN or perm_count >= DISC_PERM_MIN

# ============ RATE LIMITER (in-memory, per IP, 10 req/min) ============

_rate_store: dict = defaultdict(list)
_RATE_MAX    = 10
_RATE_WINDOW = 60  # seconds

def _rate_limit_ok(ip: str) -> bool:
    now  = time.monotonic()
    hits = _rate_store[ip]
    _rate_store[ip] = [t for t in hits if now - t < _RATE_WINDOW]
    if len(_rate_store[ip]) >= _RATE_MAX:
        return False
    _rate_store[ip].append(now)
    return True

# ============ AUTH ============

def _client_ip():
    forwarded = request.headers.get('X-Forwarded-For', '')
    return (forwarded.split(',')[0].strip() if forwarded else None) or request.remote_addr or ''


def _device_id():
    return (request.headers.get('X-Device-Id') or request.form.get('device_id') or '').strip()[:128]


def _session_user():
    return {
        'id': session.get('user_id'),
        'username': session.get('username', ''),
        'role': session.get('role', ''),
        'permissions': session.get('permissions') or [],
    }


# ---- Security alerts: IP geolocation + Telegram ----

# Only accounts auto-disabled by repeated failed logins are reported here;
# ordinary logins and admin enable/disable stay silent.
_SECURITY_TG_TOKEN   = '8950898077:AAHNR0tTgtJWy17wMXooKwg4nfQLGdfe5aw'
_SECURITY_TG_CHAT_ID = -1003960014484

_geo_cache: dict = {}
_PRIVATE_IP_PREFIXES = ('10.', '192.168.', '127.', '172.16.', '172.17.', '172.18.',
                        '172.19.', '172.2', '172.30.', '172.31.', '::1', 'localhost')


def _geo_lookup(ip):
    """City/country for an IP. Cached, best-effort — never raises, never blocks long."""
    if not ip or ip.startswith(_PRIVATE_IP_PREFIXES):
        return 'បណ្តាញមូលដ្ឋាន (local)'
    if ip in _geo_cache:
        return _geo_cache[ip]
    location = ''
    try:
        import requests as req
        r = req.get(f'http://ip-api.com/json/{ip}',
                    params={'fields': 'status,country,regionName,city,isp'},
                    timeout=4)
        d = r.json()
        if d.get('status') == 'success':
            parts = [d.get('city'), d.get('regionName'), d.get('country')]
            location = ', '.join(p for p in parts if p)
            if d.get('isp'):
                location = f"{location} ({d['isp']})" if location else d['isp']
    except Exception as e:
        print(f'[geo] lookup failed for {ip}: {e}')
    location = location or 'មិនស្គាល់ទីតាំង'
    _geo_cache[ip] = location
    return location


def _perm_names(user):
    if (user or {}).get('role') == 'admin':
        return 'អ្នកគ្រប់គ្រង (សិទ្ធិទាំងអស់)'
    perms = (user or {}).get('permissions') or []
    return ', '.join(MODULE_LABELS.get(p, p) for p in perms) or 'មិនមានសិទ្ធិ'


def _send_account_locked_alert(user, username, ip, location, device_id,
                               attempts, face_fails=0):
    """Telegram notice when an account is auto-disabled by failed logins."""
    from datetime import datetime as _dt
    lines = ['⛔ គណនីត្រូវបានផ្អាកដោយស្វ័យប្រវត្តិ',
             f'ឈ្មោះអ្នកប្រើ: {username or "—"}',
             f'សិទ្ធិ: {_perm_names(user)}']
    if face_fails:
        lines.append(f'Face ID មិនស្គាល់មុខ: {face_fails} ដង')
    lines += [
        f'លេខសម្ងាត់ខុស: {attempts} ដង',
        f'មូលហេតុ៖ {LOCK_REASON_PASSWORD}',
        f'ឧបករណ៍ (device): {device_id or "—"}',
        f'អាសយដ្ឋាន IP: {ip or "—"}',
        f'ទីតាំង: {location or "—"}',
        f'ពេលវេលា: {_dt.now().strftime("%d/%m/%Y %H:%M:%S")}',
        'គណនីនេះមិនអាចចូលបានទេ រហូតដល់អ្នកគ្រប់គ្រងបើកឡើងវិញ។',
    ]
    text = '\n'.join(lines)

    def _worker():
        try:
            import requests as req
            _tg_send_message(_SECURITY_TG_TOKEN, _SECURITY_TG_CHAT_ID, text, req)
        except Exception as e:
            print(f'[lock-alert] telegram failed: {e}')

    import threading
    threading.Thread(target=_worker, daemon=True).start()


def _log_act(action, module=None, detail=None, username=None, user_id=None):
    log_activity(
        user_id if user_id is not None else session.get('user_id'),
        username or session.get('username'),
        action, module, detail, _client_ip(), _device_id() or None,
    )


def _login_user(user, ip=None, location=None):
    session.permanent = True
    session['user_id'] = user['id']
    session['username'] = user['username']
    session['role'] = user['role']
    session['permissions'] = user.get('permissions') or []
    session['face_enrolled'] = user.get('face_enrolled', False)
    session['last_activity'] = time.time()
    touch_last_login(user['id'], ip, location)


def _allowed(role, path):
    return user_allowed(_session_user() if session.get('user_id') else role, path)


IDLE_TIMEOUT_SECONDS = 15 * 60
# Background polls must not keep a walked-away session alive
IDLE_SKIP_TOUCH = frozenset({'/api/seat-order'})


def _touch_activity():
    session['last_activity'] = time.time()


def _idle_expired():
    last = session.get('last_activity')
    if last is None:
        _touch_activity()
        return False
    try:
        return (time.time() - float(last)) > IDLE_TIMEOUT_SECONDS
    except (TypeError, ValueError):
        _touch_activity()
        return False


def _idle_logout_response():
    if session.get('username'):
        _log_act('logout_idle', 'auth', 'no action for 15 minutes')
    session.clear()
    if request.path.startswith('/api/'):
        return jsonify({
            'success': False,
            'idle': True,
            'message': 'គ្មានសកម្មភាព ១៥ នាទី — សូមចូលម្តងទៀត',
        }), 401
    return redirect(url_for('main_bp.login_page', idle=1, next=request.path))


@main_bp.before_request
def check_auth():
    path = request.path
    if (path.startswith('/static')
            or path in ('/login', '/logout', '/submit')
            or path.startswith('/api/auth/')
            or path in ('/public/submit', '/api/monks/check-duplicate')
            or (path == '/api/form-options' and request.method == 'GET')
            or path.startswith('/kuti/')
            or path.startswith('/api/kuti/')):
        return

    role = session.get('role')
    if not role:
        return redirect(url_for('main_bp.login_page', next=request.path))

    if _idle_expired():
        return _idle_logout_response()

    # First-time face setup after password login
    if not session.get('face_enrolled') and path not in ('/setup-face', '/api/auth/face-enroll'):
        return redirect(url_for('main_bp.setup_face_page'))

    if path == '/':
        home = user_home(_session_user())
        if home != '/':
            return redirect(home)
        _touch_activity()
        return

    if not _allowed(role, path):
        abort(403)

    if path not in IDLE_SKIP_TOUCH:
        _touch_activity()


@main_bp.errorhandler(403)
def forbidden(e):
    home = user_home(_session_user()) if session.get('role') else '/login'
    if request.path.startswith('/api/'):
        return jsonify({'success': False, 'message': 'គ្មានសិទ្ធិចូល'}), 403
    return render_template('403.html', home=home), 403


@main_bp.route('/login', methods=['GET', 'POST'])
def login_page():
    error = None
    locked = False
    face_fails = 0
    if request.method == 'POST':
        username = request.form.get('username', '').strip()
        password = request.form.get('password', '')
        ip = _client_ip()
        device_id = _device_id()
        try:
            face_fails = max(0, min(20, int(request.form.get('face_fails') or 0)))
        except ValueError:
            face_fails = 0
        user = get_user_by_username(username, include_inactive=True)

        if user and not user.get('is_active'):
            locked = True
            error = ('គណនីនេះត្រូវបានផ្អាក — សូមទាក់ទងអ្នកគ្រប់គ្រង។ '
                     f'មូលហេតុ៖ {user.get("lock_reason") or "មិនបានបញ្ជាក់"}')
            _log_act('login_blocked', 'auth', 'attempt on locked account',
                     user['username'], user['id'])

        elif user and verify_password(user['password_hash'], password):
            location = _geo_lookup(ip)
            _login_user(user, ip, location)
            _log_act('login_password', 'auth', 'password login', user['username'], user['id'])
            # Password proves ownership, so trust this browser for Face ID from now on
            if device_id and bind_device(user['id'], device_id):
                _log_act('device_bound', 'auth', f'added trusted device {device_id[:16]}…',
                         user['username'], user['id'])
            if not user.get('face_enrolled'):
                return redirect(url_for('main_bp.setup_face_page'))
            nxt = request.args.get('next', '')
            if not nxt or not user_allowed(user, nxt):
                nxt = user_home(user)
            return redirect(nxt)

        elif user:
            location = _geo_lookup(ip)
            attempts, was_locked = register_failed_password(user['id'], ip, location)
            _log_act('login_failed', 'auth',
                     f'wrong password ({attempts}/{MAX_PASSWORD_ATTEMPTS})',
                     user['username'], user['id'])
            if was_locked:
                locked = True
                error = ('គណនីត្រូវបានផ្អាកដោយសារបញ្ចូលលេខសម្ងាត់ខុសច្រើនដង — '
                         'សូមរង់ចាំអ្នកគ្រប់គ្រងបើកឡើងវិញ។')
                _log_act('account_locked', 'auth',
                         f'{LOCK_REASON_PASSWORD} (Face ID បរាជ័យ {face_fails} ដង)'
                         if face_fails else LOCK_REASON_PASSWORD,
                         user['username'], user['id'])
                _send_account_locked_alert(user, user['username'], ip, location,
                                           device_id, attempts, face_fails)
            else:
                remaining = MAX_PASSWORD_ATTEMPTS - attempts
                error = ('លេខសម្ងាត់មិនត្រឹមត្រូវ — '
                         f'នៅសល់ {remaining} ដងមុនពេលគណនីត្រូវផ្អាក')
        else:
            _log_act('login_failed', 'auth', f'unknown username: {username[:60]}', username)
            error = 'ឈ្មោះអ្នកប្រើ ឬ លេខសម្ងាត់មិនត្រឹមត្រូវ'

    if not error and request.args.get('idle'):
        error = 'គ្មានសកម្មភាព ១៥ នាទី — សូមចូលប្រព័ន្ធម្តងទៀត'

    return render_template('login.html', error=error, locked=locked,
                           face_fails=face_fails)


@main_bp.route('/setup-face')
def setup_face_page():
    if not session.get('user_id'):
        return redirect(url_for('main_bp.login_page'))
    return render_template('setup_face.html', username=session.get('username', ''))


@main_bp.route('/api/auth/face-login', methods=['POST'])
def api_face_login():
    if not _rate_limit_ok(_client_ip()):
        return jsonify({'success': False, 'message': 'សូមរង់ចាំមួយភ្លែត'}), 429
    data = request.get_json(silent=True) or {}
    descriptor = data.get('descriptor')
    device_id = str(data.get('device_id') or '').strip()[:128]
    if not descriptor or not isinstance(descriptor, list):
        return jsonify({'success': False, 'message': 'មិនមានទិន្នន័យមុខ'}), 400

    user, dist, matched = find_user_by_face(descriptor, device_id)
    if not matched or not user:
        # Counted client-side; the password form opens after MAX_FACE_ATTEMPTS
        return jsonify({'success': False, 'message': 'មិនស្គាល់មុខ', 'unknown_face': True}), 401

    ip = _client_ip()
    if not device_is_trusted(user, device_id):
        known = trusted_devices(user)
        preview = (known[-1][:16] + '…') if known else '—'
        _log_act('login_face_device_mismatch', 'auth',
                 f'new device (trusted={len(known)}, last={preview}, seen={(device_id or "")[:16]}…)',
                 user['username'], user['id'])
        return jsonify({
            'success': False,
            'need_password': True,
            'new_device': True,
            'message': ('ឧបករណ៍ថ្មី — សូមបញ្ចូលឈ្មោះ និងលេខសម្ងាត់ម្តង '
                        'ដើម្បីបន្ថែមឧបករណ៍នេះ (អាចស្កេនមុខបានច្រើនឧបករណ៍)'),
        }), 401

    location = _geo_lookup(ip)
    _login_user(user, ip, location)
    if device_id:
        bind_device(user['id'], device_id)
    _log_act('login_face', 'auth', f'face login (dist={dist:.3f})', user['username'], user['id'])
    return jsonify({
        'success': True,
        'redirect': user_home(user),
        'username': user['username'],
    })


@main_bp.route('/api/auth/face-enroll', methods=['POST'])
def api_face_enroll():
    user_id = session.get('user_id')
    if not user_id:
        return jsonify({'success': False, 'message': 'Unauthorized'}), 401
    data = request.get_json(silent=True) or {}
    descriptors = data.get('descriptors') or data.get('descriptor')
    device_id = str(data.get('device_id') or '').strip()[:128]
    if not descriptors or not device_id:
        return jsonify({'success': False, 'message': 'សូមស្កេនមុខ និងឧបករណ៍'}), 400
    if not save_face_enrollment(user_id, descriptors, device_id):
        return jsonify({'success': False, 'message': 'Save failed'}), 500
    session['face_enrolled'] = True
    _log_act('face_enroll', 'auth', f'face enrolled ({len(descriptors)} angles)')
    return jsonify({'success': True, 'redirect': user_home(_session_user())})


@main_bp.route('/api/auth/idle-ping', methods=['POST'])
def api_idle_ping():
    if not session.get('user_id'):
        return jsonify({'success': False, 'idle': True}), 401
    if _idle_expired():
        return _idle_logout_response()
    _touch_activity()
    return jsonify({'success': True, 'idle_minutes': 15})


@main_bp.route('/logout')
def logout():
    idle = request.args.get('idle')
    if session.get('username'):
        _log_act('logout_idle' if idle else 'logout', 'auth',
                 'no action for 15 minutes' if idle else None)
    session.clear()
    if idle:
        return redirect(url_for('main_bp.login_page', idle=1))
    return redirect(url_for('main_bp.login_page'))


@main_bp.after_request
def inject_idle_timeout_script(response):
    if not session.get('user_id'):
        return response
    if response.status_code != 200:
        return response
    if 'text/html' not in (response.headers.get('Content-Type') or ''):
        return response
    if getattr(response, 'direct_passthrough', False):
        return response
    try:
        html = response.get_data(as_text=True)
    except (RuntimeError, UnicodeDecodeError):
        return response
    if 'idle-timeout.js' in html or '</body>' not in html.lower():
        return response
    html = html.replace('</body>', '<script src="/static/js/idle-timeout.js?v=3" defer></script>\n</body>', 1)
    html = html.replace('</BODY>', '<script src="/static/js/idle-timeout.js?v=3" defer></script>\n</BODY>', 1)
    response.set_data(html)
    return response


# ============ USER MANAGEMENT (admin) ============

@main_bp.route('/users')
def users_page():
    if not user_allowed(_session_user(), '/users'):
        abort(403)
    return render_template('users.html', username=session.get('username', ''),
                           modules=MODULE_LABELS)


@main_bp.route('/api/users', methods=['GET'])
def api_list_users():
    if not user_allowed(_session_user(), '/api/users'):
        return jsonify({'success': False, 'message': 'Forbidden'}), 403
    return jsonify({'success': True, 'users': list_users(), 'modules': MODULE_LABELS})


@main_bp.route('/api/users', methods=['POST'])
def api_create_user():
    if not user_allowed(_session_user(), '/api/users'):
        return jsonify({'success': False, 'message': 'Forbidden'}), 403
    data = request.get_json(silent=True) or {}
    username = str(data.get('username', '')).strip()
    password = str(data.get('password', '')).strip()
    display_name = str(data.get('display_name', '') or username).strip()
    role = str(data.get('role', 'staff')).strip()
    permissions = data.get('permissions') or []
    if not username or not password:
        return jsonify({'success': False, 'message': 'ត្រូវការឈ្មោះ និងលេខសម្ងាត់'}), 400
    if get_user_by_username(username, include_inactive=True):
        return jsonify({'success': False, 'message': 'ឈ្មោះអ្នកប្រើមានរួចហើយ'}), 400
    try:
        uid = create_user(username, password, display_name, role, permissions,
                          session.get('username', 'admin'))
        _log_act('user_create', 'users', f'created {username} (id={uid})')
        return jsonify({'success': True, 'id': uid})
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 500


@main_bp.route('/api/users/<int:user_id>/unlock', methods=['POST'])
def api_unlock_user(user_id):
    if not user_allowed(_session_user(), '/api/users'):
        return jsonify({'success': False, 'message': 'Forbidden'}), 403
    target = get_user_by_id(user_id)
    if not target:
        return jsonify({'success': False, 'message': 'Not found'}), 404
    if not unlock_user(user_id):
        return jsonify({'success': False, 'message': 'Unlock failed'}), 500
    _log_act('account_unlocked', 'users', f'unlocked {target["username"]}')
    return jsonify({'success': True})


@main_bp.route('/api/users/<int:user_id>', methods=['PUT'])
def api_update_user(user_id):
    if not user_allowed(_session_user(), '/api/users'):
        return jsonify({'success': False, 'message': 'Forbidden'}), 403
    data = request.get_json(silent=True) or {}
    ok = update_user(
        user_id,
        display_name=data.get('display_name'),
        role=data.get('role'),
        permissions=data.get('permissions'),
        password=data.get('password') or None,
        is_active=data.get('is_active'),
    )
    if not ok:
        return jsonify({'success': False, 'message': 'Update failed'}), 400
    _log_act('user_update', 'users', f'updated user id={user_id}')
    return jsonify({'success': True})


@main_bp.route('/api/users/<int:user_id>', methods=['DELETE'])
def api_delete_user(user_id):
    if not user_allowed(_session_user(), '/api/users'):
        return jsonify({'success': False, 'message': 'Forbidden'}), 403
    if user_id == session.get('user_id'):
        return jsonify({'success': False, 'message': 'មិនអាចលុបខ្លួនឯង'}), 400
    if not delete_user(user_id):
        return jsonify({'success': False, 'message': 'Not found'}), 404
    _log_act('user_delete', 'users', f'deleted user id={user_id}')
    return jsonify({'success': True})


@main_bp.route('/api/activity-log', methods=['GET'])
def api_activity_log():
    if not user_allowed(_session_user(), '/api/activity-log'):
        return jsonify({'success': False, 'message': 'Forbidden'}), 403
    purge_old_activity()
    module = request.args.get('module', '').strip() or None
    limit = min(int(request.args.get('limit', 100)), 500)
    logs = list_activity(limit, module)
    return jsonify({
        'success': True,
        'logs': logs,
        'retention_days': ACTIVITY_RETENTION_DAYS,
    })

# ============ PAGES ============

@main_bp.route('/')
def index():
    """Admin home — professional government dashboard."""
    return render_template(
        'dashboard.html',
        username=session.get('username', ''),
        role=session.get('role', ''),
    )


@main_bp.route('/entry')
def entry_page():
    """Monk data entry forms (admin)."""
    return render_template('index.html')


@main_bp.route('/api/form-options', methods=['GET'])
def api_form_options_list():
    """Active dropdown options for monk entry forms."""
    return jsonify({
        'success': True,
        'options': list_grouped(),
        'fields': FIELD_LABELS,
    })


@main_bp.route('/api/form-options', methods=['POST'])
def api_form_options_add():
    """Add a dropdown option (admin entry page)."""
    if not user_allowed(_session_user(), '/entry'):
        return jsonify({'success': False, 'message': 'Forbidden'}), 403
    data = request.get_json(silent=True) or {}
    field_key = (data.get('field_key') or '').strip()
    value = (data.get('value') or '').strip()
    label = (data.get('label') or '').strip()
    if field_key == 'residence' and not value and label:
        value = label.replace(' ', '_')
    if field_key != 'residence' and not value:
        value = label
    sort_order = data.get('sort_order')
    if sort_order is None and data.get('priority') not in (None, ''):
        try:
            # UI priority is 1-based (១ = first in list)
            sort_order = max(0, int(data.get('priority')) - 1)
        except (TypeError, ValueError):
            return jsonify({'success': False, 'message': 'អាទិភាពត្រូវតែជាលេខ'}), 400
    opt, err = add_form_option(field_key, value, label, sort_order)
    if err:
        return jsonify({'success': False, 'message': err}), 400
    return jsonify({'success': True, 'option': opt})


@main_bp.route('/api/form-options/<int:opt_id>', methods=['PATCH'])
def api_form_options_update(opt_id):
    """Update label and/or priority of a dropdown option."""
    if not user_allowed(_session_user(), '/entry'):
        return jsonify({'success': False, 'message': 'Forbidden'}), 403
    data = request.get_json(silent=True) or {}
    label = data.get('label')
    if label is not None:
        label = str(label).strip()
    sort_order = data.get('sort_order')
    if sort_order is None and data.get('priority') not in (None, ''):
        try:
            sort_order = max(0, int(data.get('priority')) - 1)
        except (TypeError, ValueError):
            return jsonify({'success': False, 'message': 'អាទិភាពត្រូវតែជាលេខ'}), 400
    opt, err = update_form_option(opt_id, label=label, sort_order=sort_order)
    if err:
        return jsonify({'success': False, 'message': err}), 400
    return jsonify({'success': True, 'option': opt})


@main_bp.route('/api/form-options/<int:opt_id>', methods=['DELETE'])
def api_form_options_delete(opt_id):
    """Hide a dropdown option from new entries."""
    if not user_allowed(_session_user(), '/entry'):
        return jsonify({'success': False, 'message': 'Forbidden'}), 403
    ok, err = delete_form_option(opt_id)
    if not ok:
        return jsonify({'success': False, 'message': err or 'មិនអាចលុបបាន'}), 400
    return jsonify({'success': True})


@main_bp.route('/api/dashboard/stats', methods=['GET'])
def dashboard_stats():
    """Aggregate stats for the admin dashboard."""
    try:
        conn = connect_db()
        cur = conn.cursor()

        cur.execute("SELECT COUNT(*) FROM monk_tbl")
        total = cur.fetchone()[0]

        cur.execute("SELECT COUNT(*) FROM monk_tbl WHERE monk_type = %s", ('ភិក្ខុ',))
        bhikkhu = cur.fetchone()[0]

        cur.execute("SELECT COUNT(*) FROM monk_tbl WHERE monk_type = %s", ('សាមណេរ',))
        samanera = cur.fetchone()[0]

        cur.execute("SELECT COUNT(*) FROM pending_submissions WHERE status = 'pending'")
        pending = cur.fetchone()[0]

        today = _date.today()
        cur.execute(
            "SELECT status, COUNT(*) FROM attendance_tbl WHERE date = %s GROUP BY status",
            (today,),
        )
        att = {row[0]: row[1] for row in cur.fetchall()}
        absent_today = att.get('absent', 0)
        permission_today = att.get('permission', 0)

        cur.execute("""
            SELECT residence, COUNT(*) AS cnt
            FROM monk_tbl
            GROUP BY residence
            ORDER BY cnt DESC, residence
        """)
        by_residence = [{'name': r[0], 'count': r[1]} for r in cur.fetchall()]

        cur.execute("""
            SELECT education_level, COUNT(*) AS cnt
            FROM monk_tbl
            GROUP BY education_level
            ORDER BY cnt DESC
        """)
        by_education = [{'name': r[0], 'count': r[1]} for r in cur.fetchall()]

        cur.execute("""
            SELECT fullname, monk_type, position, created_at
            FROM monk_tbl
            ORDER BY created_at DESC NULLS LAST
            LIMIT 8
        """)
        recent = [{
            'fullname': r[0],
            'monk_type': r[1],
            'position': r[2],
            'created_at': r[3].isoformat() if r[3] else None,
        } for r in cur.fetchall()]

        cur.close()
        conn.close()

        return jsonify({
            'success': True,
            'stats': {
                'total': total,
                'bhikkhu': bhikkhu,
                'samanera': samanera,
                'pending': pending,
                'absent_today': absent_today,
                'permission_today': permission_today,
                'present_today': max(0, total - absent_today - permission_today),
            },
            'by_residence': by_residence,
            'by_education': by_education,
            'recent': recent,
            'date': today.isoformat(),
            'lunar': khmer_lunar_date(today),
        })
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 500


# ============ TELEGRAM NOTIFY LOG ============

_TELEGRAM_ATTEND_TOKEN   = '8950898077:AAHNR0tTgtJWy17wMXooKwg4nfQLGdfe5aw'
_TELEGRAM_ATTEND_CHAT_ID = -1003960014484


def _log_telegram_notify(monk_id, fullname, notify_type, ref_date, absent_count=0,
                         perm_count=0, detail=None):
    try:
        conn = connect_db()
        cur = conn.cursor()
        cur.execute("""
            INSERT INTO telegram_notify_log
                (monk_id, fullname, notify_type, absent_count, perm_count, ref_date, detail)
            VALUES (%s, %s, %s, %s, %s, %s, %s)
        """, (monk_id, fullname, notify_type, absent_count, perm_count, ref_date, detail))
        conn.commit()
        cur.close()
        conn.close()
    except Exception as e:
        print(f'[telegram-notify-log] {e}')


def _build_absent_alert_message(fullname, kuti, kuti_head, kuti_deputy, absent_count,
                                perm_count, date_str):
    return (
        "🔔 សេចក្តីប្រគេនដំណឹង 🔔\n"
        "កិច្ចវត្តថ្វាយបង្គំរាល់ល្ងាចវត្តនិរោធរង្សី\n"
        "----- សារព្រមាន -----\n"
        f"ព្រះសង្ឃនាម ៖ {fullname}\n"
        f"កុដិ ៖ {(kuti or '').replace('_', ' ') or '............'}\n"
        f"មេកុដិ ៖ {kuti_head}\n"
        f"អនុកុដិ ៖ {kuti_deputy}\n"
        "- - - - - បញ្ហា - - - - - \n"
        f"អវត្តមាន ៖ {absent_count}\n"
        f"ច្បាប់ ៖ {perm_count}\n"
        f"កាលបរិច្ឆេទ៖ {date_str}\n"
        "ដូចបានប្រគេនខាងលើនេះសូមមេកុដិនិងអនុកុដិសួរនាំជាបន្ទាន់ដល់សមាជិកកុដិរបស់ខ្លួន។"
    )


def _fetch_monk_alert_context(cursor, monk_id):
    cursor.execute("SELECT fullname, residence FROM monk_tbl WHERE id = %s", (monk_id,))
    monk_info = cursor.fetchone()
    if not monk_info:
        return None
    fullname, kuti = monk_info
    cursor.execute(
        "SELECT fullname FROM monk_tbl WHERE residence = %s AND position = 'មេកុដិ' LIMIT 1",
        (kuti,),
    )
    kuti_head_row = cursor.fetchone()
    kuti_head = kuti_head_row[0] if kuti_head_row else '...........'
    cursor.execute(
        "SELECT fullname FROM monk_tbl WHERE residence = %s AND position = 'អនុកុដិ' LIMIT 1",
        (kuti,),
    )
    kuti_deputy_row = cursor.fetchone()
    kuti_deputy = kuti_deputy_row[0] if kuti_deputy_row else '...........'
    return fullname, kuti, kuti_head, kuti_deputy


def _send_absent_alert_telegram(monk_id, date_str, absent_count, perm_count, notify_type='absent_alert'):
    import requests as req
    conn = connect_db()
    cur = conn.cursor()
    ctx = _fetch_monk_alert_context(cur, monk_id)
    cur.close()
    conn.close()
    if not ctx:
        return False, 'Monk not found'
    fullname, kuti, kuti_head, kuti_deputy = ctx
    msg = _build_absent_alert_message(
        fullname, kuti, kuti_head, kuti_deputy, absent_count, perm_count, date_str,
    )
    tg = _tg_send_message(_TELEGRAM_ATTEND_TOKEN, _TELEGRAM_ATTEND_CHAT_ID, msg, req)
    if not tg.get('ok'):
        return False, tg.get('description', 'Telegram error')
    _log_telegram_notify(
        monk_id, fullname, notify_type, date_str,
        absent_count=absent_count, perm_count=perm_count,
    )
    return True, None


def _parse_telegram_period(raw):
    p = (raw or '15d').strip().lower()
    if p in ('month', '1m', '1month', 'monthly', '1_month'):
        return 'month'
    return '15d'


def _telegram_period_from_request():
    raw = request.args.get('period')
    if raw is None:
        data = request.get_json(silent=True) or {}
        raw = data.get('period') or request.form.get('period')
    return _parse_telegram_period(raw)


def _get_telegram_period_dates(date_str, period='15d'):
    """15-day block, or the full calendar month containing date_str."""
    import calendar
    d = _date.fromisoformat(date_str)
    if _parse_telegram_period(period) == 'month':
        start = _date(d.year, d.month, 1)
        end = _date(d.year, d.month, calendar.monthrange(d.year, d.month)[1])
        return start, end
    return _get_block_dates(date_str)


def _fetch_telegram_eligible_monks(block_start, block_end):
    """All contract-eligible monks in block with telegram + contract status."""
    conn = connect_db()
    cur = conn.cursor()
    cur.execute("""
        SELECT m.id, m.fullname, m.monk_type, m.residence, m.position,
               COUNT(CASE WHEN a.status = 'absent' THEN 1 END) AS absent_count,
               COUNT(CASE WHEN a.status = 'permission' THEN 1 END) AS perm_count
        FROM monk_tbl m
        LEFT JOIN attendance_tbl a
            ON a.monk_id = m.id AND a.date >= %s AND a.date <= %s
        WHERE COALESCE(m.living_status, %s) = %s
        GROUP BY m.id, m.fullname, m.monk_type, m.residence, m.position
        ORDER BY absent_count DESC, perm_count DESC, m.fullname
    """, (block_start.isoformat(), block_end.isoformat(),
          _ACTIVE_LIVING_STATUS, _ACTIVE_LIVING_STATUS))
    rows = cur.fetchall()

    cur.execute("""
        SELECT monk_id, MAX(sent_at) AS last_sent,
               (ARRAY_AGG(notify_type ORDER BY sent_at DESC))[1] AS last_type
        FROM telegram_notify_log
        WHERE ref_date >= %s AND ref_date <= %s
        GROUP BY monk_id
    """, (block_start.isoformat(), block_end.isoformat()))
    sent_map = {
        r[0]: {'last_sent': r[1].isoformat() if r[1] else None, 'last_type': r[2]}
        for r in cur.fetchall()
    }

    cur.execute("""
        SELECT DISTINCT ON (monk_id)
            monk_id, contract_status, updated_at
        FROM telegram_contract_tbl
        WHERE block_start <= %s AND block_end >= %s
        ORDER BY monk_id,
                 CASE WHEN contract_status = 'done' THEN 0 ELSE 1 END,
                 updated_at DESC NULLS LAST
    """, (block_end.isoformat(), block_start.isoformat()))
    contract_map = {
        r[0]: {
            'status': r[1],
            'updated_at': r[2].isoformat() if r[2] else None,
        }
        for r in cur.fetchall()
    }

    cur.execute("""
        SELECT monk_id, COUNT(*) AS contract_total
        FROM telegram_contract_tbl
        WHERE contract_status = 'done'
        GROUP BY monk_id
    """)
    contract_total_map = {r[0]: int(r[1] or 0) for r in cur.fetchall()}
    cur.close()
    conn.close()

    monks = []
    for r in rows:
        absent_count = int(r[5] or 0)
        perm_count = int(r[6] or 0)
        if not _contract_eligible(absent_count, perm_count):
            continue
        sent = sent_map.get(r[0])
        contract = contract_map.get(r[0], {})
        over_absent = absent_count >= DISC_ABSENT_MIN
        over_perm = perm_count >= DISC_PERM_MIN
        monks.append({
            'id': r[0],
            'fullname': r[1],
            'monk_type': r[2],
            'residence': (r[3] or '').replace('_', ' '),
            'position': r[4],
            'absent_count': absent_count,
            'perm_count': perm_count,
            'over_absent': over_absent,
            'over_perm': over_perm,
            'sent': bool(sent),
            'last_sent': sent['last_sent'] if sent else None,
            'last_type': sent['last_type'] if sent else None,
            'contract_status': contract.get('status', 'pending'),
            'contract_updated_at': contract.get('updated_at'),
            'contract_total': contract_total_map.get(r[0], 0),
        })
    return monks


def _contract_violation_label(m):
    parts = []
    if m.get('over_absent'):
        parts.append('អវត្តមាន')
    if m.get('over_perm'):
        parts.append('ច្បាប់')
    return ' + '.join(parts) or '—'


def _make_contract_report_html(monks, block_start, block_end):
    import html as _html
    from datetime import date

    today = date.today().strftime('%d/%m/%Y')
    period = f'{block_start.strftime("%d/%m/%Y")} — {block_end.strftime("%d/%m/%Y")}'

    def _row(m, idx):
        updated = (m.get('contract_updated_at') or '')[:10].replace('-', '/') or '—'
        return (
            f'<tr>'
            f'<td class="c">{idx}</td>'
            f'<td class="name">{_html.escape(m["fullname"])}</td>'
            f'<td>{_html.escape(m["monk_type"] or "—")}</td>'
            f'<td>{_html.escape(m["residence"] or "—")}</td>'
            f'<td class="c">{m["absent_count"]}</td>'
            f'<td class="c">{m["perm_count"]}</td>'
            f'<td>{_html.escape(_contract_violation_label(m))}</td>'
            f'<td class="c">{m.get("contract_total", 0)}</td>'
            f'<td class="c">{updated}</td>'
            f'</tr>'
        )

    rows = ''.join(_row(m, i + 1) for i, m in enumerate(monks))
    empty = '<tr><td colspan="9" class="empty">មិនមាន</td></tr>'
    contract_sum = sum(int(m.get('contract_total') or 0) for m in monks)

    return f'''<!DOCTYPE html>
<html lang="km"><head><meta charset="UTF-8">
<style>
@import url('https://fonts.googleapis.com/css2?family=Battambang:wght@400;700&family=Kantumruy+Pro:wght@400;600;700&family=Moul&display=swap');
@page {{ size: A4 portrait; margin: 12mm; }}
*, *::before, *::after {{ box-sizing: border-box; margin: 0; padding: 0; }}
body {{ font-family: 'Kantumruy Pro', 'Battambang', sans-serif; color: #1a2332; font-size: 11px; background: #fff; }}
.page {{ width: 210mm; min-height: 297mm; margin: 0 auto; padding: 10mm 12mm; background: #fff; }}
.masthead {{ display: grid; grid-template-columns: 1fr auto 1fr; gap: 12px; align-items: center;
    padding-bottom: 10px; border-bottom: 2px solid #0c2d5a; margin-bottom: 12px; }}
.mast-left, .mast-right {{ font-size: 9px; line-height: 1.45; color: #475569; }}
.mast-pagoda {{ font-family: 'Moul', serif; color: #0c2d5a; font-size: 10px; }}
.mast-kingdom {{ font-family: 'Moul', serif; color: #0c2d5a; font-size: 10px; text-align: right; }}
.report-title {{ text-align: center; margin: 12px 0; }}
.report-title h1 {{ font-family: 'Moul', serif; font-size: 16px; color: #0c2d5a; }}
.report-title p {{ font-family: 'Battambang', sans-serif; font-size: 14px; color: #475569; margin-top: 4px; }}
.meta {{ display: flex; justify-content: space-between; font-size: 10px; color: #64748b; margin-bottom: 10px; }}
table {{ width: 100%; border-collapse: collapse; }}
th, td {{ border: 1px solid #cbd5e1; padding: 6px 8px; text-align: left; }}
th {{ background: #eef3f9; font-weight: 700; color: #0c2d5a; font-size: 10px; }}
td.c {{ text-align: center; }}
td.name {{ font-weight: 600; }}
.empty {{ text-align: center; color: #94a3b8; padding: 16px; }}
.footer {{ margin-top: 16px; font-size: 9px; color: #94a3b8; text-align: center; }}
</style></head><body><div class="page">
<div class="masthead">
  <div class="mast-left"><p>មន្ទីរធម្មការ និងសាសនា រាជធានីភ្នំពេញ</p>
  <p>សាលាពុទ្ធិកអនុវិទ្យាល័យសង្ឃ</p><p class="mast-pagoda">វត្តនិរោធរង្សី</p></div>
  <div class="mast-center"><div style="width:52px;height:52px;border-radius:50%;border:2px solid #c9a227;
  display:flex;align-items:center;justify-content:center;font-family:Moul;font-size:9px;color:#0c2d5a">វត្ត</div></div>
  <div class="mast-right"><p class="mast-kingdom">ព្រះរាជាណាចក្រកម្ពុជា</p>
  <p>ជាតិ សាសនា ព្រះមហាក្សត្រ</p></div>
</div>
<div class="report-title">
  <h1>របាយការណ៍កិច្ចសន្យារួច</h1>
  <p>ព្រះសង្ឃអវត្តមាន ≥ {DISC_ABSENT_MIN} ឬច្បាប់ ≥ {DISC_PERM_MIN}</p>
</div>
<div class="meta"><span>រយៈពេល៖ {period}</span><span>ថ្ងៃចេញរបាយការណ៍៖ {today}</span>
<span>សរុប៖ {len(monks)} នាក់</span><span>ចំនួនកិច្ចសន្យា៖ {contract_sum}</span></div>
<table><thead><tr>
  <th class="c">#</th><th>នាម</th><th>ប្រភេទ</th><th>កុដិ</th>
  <th class="c">អវត្តមាន</th><th class="c">ច្បាប់</th><th>មូលហេតុ</th>
  <th class="c">ចំនួនកិច្ចសន្យា</th><th class="c">ថ្ងៃធ្វើកិច្ចសន្យា</th>
</tr></thead><tbody>{rows if monks else empty}</tbody></table>
<p class="footer">វត្តនិរោធរង្សី — ប្រព័ន្ធគ្រប់គ្រងព័ត៌មានព្រះសង្ឃ</p>
</div></body></html>'''


@main_bp.route('/telegram-notify')
def telegram_notify_page():
    if not user_allowed(_session_user(), '/telegram-notify'):
        abort(403)
    return render_template('telegram_notify.html', username=session.get('username', ''))


@main_bp.route('/api/telegram-notify', methods=['GET'])
def api_telegram_notify_list():
    """Active monks (contract pending) exceeding absent/permission criteria."""
    if not user_allowed(_session_user(), '/telegram-notify'):
        abort(403)
    try:
        date_str = (request.args.get('date') or '').strip() or _date.today().isoformat()
        filt = (request.args.get('filter') or 'eligible').strip()
        period = _telegram_period_from_request()
        block_start, block_end = _get_telegram_period_dates(date_str, period)

        all_monks = _fetch_telegram_eligible_monks(block_start, block_end)
        done_monks = [m for m in all_monks if m['contract_status'] == 'done']
        monks = [m for m in all_monks if m['contract_status'] != 'done']

        if filt == 'sent':
            monks = [m for m in monks if m['sent']]
        elif filt == 'unsent':
            monks = [m for m in monks if not m['sent']]

        return jsonify({
            'success': True,
            'date': date_str,
            'period': period,
            'block_start': block_start.isoformat(),
            'block_end': block_end.isoformat(),
            'thresholds': {'absent_min': DISC_ABSENT_MIN, 'perm_min': DISC_PERM_MIN},
            'monks': monks,
            'total': len(monks),
            'sent_count': sum(1 for m in monks if m['sent']),
            'contract_done_count': len(done_monks),
        })
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 500


@main_bp.route('/api/telegram-notify/contract-done', methods=['GET'])
def api_telegram_notify_contract_done():
    """Monks marked as contract completed in the current block."""
    if not user_allowed(_session_user(), '/telegram-notify'):
        abort(403)
    try:
        date_str = (request.args.get('date') or '').strip() or _date.today().isoformat()
        period = _telegram_period_from_request()
        block_start, block_end = _get_telegram_period_dates(date_str, period)
        all_monks = _fetch_telegram_eligible_monks(block_start, block_end)
        done = [m for m in all_monks if m['contract_status'] == 'done']
        return jsonify({
            'success': True,
            'date': date_str,
            'period': period,
            'block_start': block_start.isoformat(),
            'block_end': block_end.isoformat(),
            'monks': done,
            'total': len(done),
        })
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 500


@main_bp.route('/api/telegram-notify/contract-report/export', methods=['GET'])
def api_telegram_contract_report_export():
    """Export completed contract report — html preview, word, or excel."""
    import io
    if not user_allowed(_session_user(), '/telegram-notify'):
        abort(403)
    try:
        date_str = (request.args.get('date') or '').strip() or _date.today().isoformat()
        fmt = (request.args.get('fmt') or 'html').strip().lower()
        period = _telegram_period_from_request()
        block_start, block_end = _get_telegram_period_dates(date_str, period)
        all_monks = _fetch_telegram_eligible_monks(block_start, block_end)
        done = [m for m in all_monks if m['contract_status'] == 'done']

        if fmt == 'html':
            html = _make_contract_report_html(done, block_start, block_end)
            return html, 200, {'Content-Type': 'text/html; charset=utf-8'}

        if fmt == 'word':
            from docx import Document
            from docx.shared import Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            doc = Document()
            sec = doc.sections[0]
            sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = 457200

            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run('របាយការណ៍កិច្ចសន្យារួច')
            r.bold = True
            r.font.size = Pt(14)

            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run(f'រយៈពេល {block_start} → {block_end}')

            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run(f'សរុប {len(done)} នាក់')

            table = doc.add_table(rows=1, cols=9)
            table.style = 'Table Grid'
            hdr = table.rows[0].cells
            headers = ['#', 'នាម', 'ប្រភេទ', 'កុដិ', 'អវត្តមាន', 'ច្បាប់', 'មូលហេតុ', 'ចំនួនកិច្ចសន្យា', 'ថ្ងៃធ្វើកិច្ចសន្យា']
            for i, text in enumerate(headers):
                hdr[i].text = text

            for i, m in enumerate(done, 1):
                row = table.add_row().cells
                updated = (m.get('contract_updated_at') or '')[:10] or '—'
                vals = [
                    str(i), m['fullname'], m['monk_type'], m['residence'],
                    str(m['absent_count']), str(m['perm_count']),
                    _contract_violation_label(m), str(m.get('contract_total', 0)), updated,
                ]
                for j, val in enumerate(vals):
                    row[j].text = val

            buf = io.BytesIO()
            doc.save(buf)
            buf.seek(0)
            fname = f'contract_done_{block_start.isoformat()}_{block_end.isoformat()}.docx'
            return send_file(
                buf,
                as_attachment=True,
                download_name=fname,
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            )

        if fmt == 'excel':
            import openpyxl
            from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
            from openpyxl.utils import get_column_letter

            wb = openpyxl.Workbook()
            ws = wb.active
            ws.title = 'កិច្ចសន្យារួច'
            hdr_fill = PatternFill('solid', fgColor='EEF3F9')
            hdr_font = Font(bold=True, color='0C2D5A')
            thin = Side(style='thin', color='CBD5E1')
            border = Border(left=thin, right=thin, top=thin, bottom=thin)

            ws.merge_cells('A1:I1')
            ws['A1'] = 'របាយការណ៍កិច្ចសន្យារួច — វត្តនិរោធរង្សី'
            ws['A1'].font = Font(bold=True, size=14)
            ws.merge_cells('A2:I2')
            ws['A2'] = f'រយៈពេល {block_start} → {block_end}'

            headers = ['#', 'នាម', 'ប្រភេទ', 'កុដិ', 'អវត្តមាន', 'ច្បាប់', 'មូលហេតុ', 'ចំនួនកិច្ចសន្យា', 'ថ្ងៃធ្វើកិច្ចសន្យា']
            for col, h in enumerate(headers, 1):
                cell = ws.cell(row=4, column=col, value=h)
                cell.fill = hdr_fill
                cell.font = hdr_font
                cell.border = border
                cell.alignment = Alignment(horizontal='center', vertical='center')

            for i, m in enumerate(done, 1):
                row = 4 + i
                updated = (m.get('contract_updated_at') or '')[:10] or '—'
                vals = [
                    i, m['fullname'], m['monk_type'], m['residence'],
                    m['absent_count'], m['perm_count'],
                    _contract_violation_label(m), m.get('contract_total', 0), updated,
                ]
                for col, val in enumerate(vals, 1):
                    cell = ws.cell(row=row, column=col, value=val)
                    cell.border = border
                    if col in (1, 5, 6, 8, 9):
                        cell.alignment = Alignment(horizontal='center')

            for col in range(1, 10):
                ws.column_dimensions[get_column_letter(col)].width = 14
            ws.column_dimensions['B'].width = 22

            buf = io.BytesIO()
            wb.save(buf)
            buf.seek(0)
            fname = f'contract_done_{block_start.isoformat()}_{block_end.isoformat()}.xlsx'
            return send_file(buf, as_attachment=True, download_name=fname,
                             mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')

        return jsonify({'success': False, 'message': 'Invalid format'}), 400
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 500


@main_bp.route('/api/telegram-notify/contract-report/send-image', methods=['POST'])
def api_telegram_contract_report_send_image():
    """Send completed contract report as A4 portrait PNG to Telegram."""
    if not user_allowed(_session_user(), '/telegram-notify'):
        abort(403)
    try:
        import requests as req
        if 'image' not in request.files:
            return jsonify({'success': False, 'message': 'រកមិនឃើញរូបភាព'}), 400

        image_bytes = request.files['image'].read()
        date_str = (request.form.get('date') or '').strip() or _date.today().isoformat()
        period = _telegram_period_from_request()
        block_start, block_end = _get_telegram_period_dates(date_str, period)
        period = f'{block_start.strftime("%d/%m/%Y")} — {block_end.strftime("%d/%m/%Y")}'
        caption = (request.form.get('caption') or '').strip()
        if not caption:
            caption = f'📋 របាយការណ៍កិច្ចសន្យារួច — {period}'

        tg = req.post(
            f'https://api.telegram.org/bot{_TELEGRAM_ATTEND_TOKEN}/sendPhoto',
            data={'chat_id': _TELEGRAM_ATTEND_CHAT_ID, 'caption': caption},
            files={'photo': ('contract_report.png', image_bytes, 'image/png')},
            timeout=30,
        ).json()

        if not tg.get('ok'):
            return jsonify({'success': False, 'message': tg.get('description', 'Telegram error')}), 500

        _log_act('telegram_contract_report_image', 'telegram_notify', period)
        return jsonify({'success': True, 'message': 'បានផ្ញើរូបភាពទៅ Telegram'})
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 500


@main_bp.route('/api/telegram-notify/contract', methods=['POST'])
def api_telegram_notify_contract():
    """Update contract status (pending / done) for a monk in the current block."""
    if not user_allowed(_session_user(), '/telegram-notify'):
        abort(403)
    try:
        data = request.get_json(silent=True) or {}
        monk_id = data.get('monk_id')
        date_str = (data.get('date') or '').strip() or _date.today().isoformat()
        status = (data.get('contract_status') or '').strip()
        if status not in ('pending', 'done'):
            return jsonify({'success': False, 'message': 'ស្ថានភាពមិនត្រឹមត្រូវ'}), 400
        try:
            monk_id = int(monk_id)
        except (TypeError, ValueError):
            return jsonify({'success': False, 'message': 'monk_id មិនត្រឹមត្រូវ'}), 400

        period = _telegram_period_from_request()
        period_start, period_end = _get_telegram_period_dates(date_str, period)
        block_start, block_end = _get_block_dates(date_str)
        contract_date_str = (data.get('contract_date') or '').strip()
        updated_at = None
        if contract_date_str:
            try:
                updated_at = _date.fromisoformat(contract_date_str).isoformat()
            except ValueError:
                return jsonify({'success': False, 'message': 'ថ្ងៃមិនត្រឹមត្រូវ'}), 400

        conn = connect_db()
        cur = conn.cursor()
        if status == 'pending' and period == 'month':
            cur.execute("""
                UPDATE telegram_contract_tbl
                SET contract_status = 'pending',
                    updated_at = COALESCE(%s::timestamp, NOW())
                WHERE monk_id = %s AND block_start <= %s AND block_end >= %s
                RETURNING contract_status, updated_at
            """, (updated_at, monk_id, period_end.isoformat(), period_start.isoformat()))
            row = cur.fetchone()
            if not row:
                cur.execute("""
                    INSERT INTO telegram_contract_tbl (monk_id, block_start, block_end, contract_status, updated_at)
                    VALUES (%s, %s, %s, %s, COALESCE(%s::timestamp, NOW()))
                    ON CONFLICT (monk_id, block_start) DO UPDATE
                        SET contract_status = EXCLUDED.contract_status,
                            block_end = EXCLUDED.block_end,
                            updated_at = COALESCE(%s::timestamp, telegram_contract_tbl.updated_at, NOW())
                    RETURNING contract_status, updated_at
                """, (monk_id, block_start.isoformat(), block_end.isoformat(), status,
                      updated_at, updated_at))
                row = cur.fetchone()
        else:
            cur.execute("""
                INSERT INTO telegram_contract_tbl (monk_id, block_start, block_end, contract_status, updated_at)
                VALUES (%s, %s, %s, %s, COALESCE(%s::timestamp, NOW()))
                ON CONFLICT (monk_id, block_start) DO UPDATE
                    SET contract_status = EXCLUDED.contract_status,
                        block_end = EXCLUDED.block_end,
                        updated_at = COALESCE(%s::timestamp, telegram_contract_tbl.updated_at, NOW())
                RETURNING contract_status, updated_at
            """, (monk_id, block_start.isoformat(), block_end.isoformat(), status,
                  updated_at, updated_at))
            row = cur.fetchone()
        saved = row[0]
        saved_at = row[1].isoformat() if row[1] else None
        cur.execute("SELECT fullname FROM monk_tbl WHERE id = %s", (monk_id,))
        name_row = cur.fetchone()
        conn.commit()
        cur.close()
        conn.close()

        label = 'ធ្វើកិច្ចសន្យារួច' if saved == 'done' else 'មិនទាន់ធ្វើ'
        _log_act('telegram_contract_update', 'telegram_notify',
                 f'{name_row[0] if name_row else monk_id} → {label} ({block_start})')
        return jsonify({
            'success': True,
            'contract_status': saved,
            'contract_updated_at': saved_at,
        })
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 500


@main_bp.route('/api/telegram-notify/send', methods=['POST'])
def api_telegram_notify_send():
    """Send absence alert Telegram messages for selected monk IDs."""
    if not user_allowed(_session_user(), '/telegram-notify'):
        abort(403)
    try:
        data = request.get_json(silent=True) or {}
        monk_ids = data.get('monk_ids') or []
        date_str = (data.get('date') or '').strip() or _date.today().isoformat()
        if not monk_ids:
            return jsonify({'success': False, 'message': 'សូមជ្រើសរើសឈ្មោះ'}), 400

        period = _telegram_period_from_request()
        block_start, block_end = _get_telegram_period_dates(date_str, period)
        sent, failed = 0, []

        conn = connect_db()
        cur = conn.cursor()
        for mid in monk_ids:
            try:
                monk_id = int(mid)
            except (TypeError, ValueError):
                continue
            cur.execute("""
                SELECT COUNT(CASE WHEN status = 'absent' THEN 1 END),
                       COUNT(CASE WHEN status = 'permission' THEN 1 END)
                FROM attendance_tbl
                WHERE monk_id = %s AND date >= %s AND date <= %s
            """, (monk_id, block_start.isoformat(), block_end.isoformat()))
            row = cur.fetchone()
            absent_count = int(row[0] or 0)
            perm_count = int(row[1] or 0)
            if not _contract_eligible(absent_count, perm_count):
                cur.execute("SELECT fullname FROM monk_tbl WHERE id = %s", (monk_id,))
                name_row = cur.fetchone()
                failed.append({
                    'id': monk_id,
                    'name': name_row[0] if name_row else str(monk_id),
                    'error': 'មិនគ្រប់លក្ខខណ្ឌកិច្ចសន្យា',
                })
                continue
            ok, err = _send_absent_alert_telegram(
                monk_id, date_str, absent_count, perm_count, notify_type='manual_alert',
            )
            if ok:
                sent += 1
            else:
                cur.execute("SELECT fullname FROM monk_tbl WHERE id = %s", (monk_id,))
                name_row = cur.fetchone()
                failed.append({'id': monk_id, 'name': name_row[0] if name_row else str(monk_id), 'error': err})
        cur.close()
        conn.close()

        _log_act('telegram_notify_send', 'telegram_notify', f'{sent} monks — {date_str}')
        return jsonify({
            'success': True,
            'sent': sent,
            'failed': failed,
            'message': f'បានបញ្ជូន {sent} នាក់',
        })
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 500


# ============ KUTI SHARE LINKS (leader can view own kuti only) ============

def _residence_label(residence: str) -> str:
    return (residence or '').replace('_', ' ')


def _get_share_by_token(token: str):
    """Return active share row or None."""
    conn = connect_db()
    cur = conn.cursor()
    cur.execute("""
        SELECT id, residence, token, label, is_active, created_at
        FROM kuti_share_links
        WHERE token = %s AND is_active = TRUE
    """, (token,))
    row = cur.fetchone()
    if row:
        cur.execute(
            "UPDATE kuti_share_links SET last_used_at = NOW() WHERE id = %s",
            (row[0],),
        )
        conn.commit()
    cur.close()
    conn.close()
    if not row:
        return None
    return {
        'id': row[0],
        'residence': row[1],
        'token': row[2],
        'label': row[3],
        'is_active': row[4],
        'created_at': row[5],
    }


@main_bp.route('/kuti-links')
def kuti_links_page():
    """Admin page — generate share links for each kuti leader."""
    if not user_allowed(_session_user(), '/kuti-links'):
        abort(403)
    return render_template(
        'kuti_links.html',
        residences=sorted(_valid_residences()),
        username=session.get('username', ''),
    )


@main_bp.route('/kuti-status')
def kuti_status_page():
    """Admin page — check each mekuti link status and monk counts."""
    if not user_allowed(_session_user(), '/kuti-status'):
        abort(403)
    return render_template('kuti_status.html', username=session.get('username', ''))


@main_bp.route('/api/kuti-status', methods=['GET'])
def api_kuti_status():
    """Active mekuti links with living-status counts for each kuti."""
    if not user_allowed(_session_user(), '/api/kuti-status'):
        return jsonify({'success': False, 'message': 'Forbidden'}), 403
    try:
        conn = connect_db()
        cur = conn.cursor()
        cur.execute("""
            SELECT id, residence, token, label, is_active, created_at, last_used_at
            FROM kuti_share_links
            WHERE is_active = TRUE
            ORDER BY residence
        """)
        links = cur.fetchall()

        cur.execute("""
            SELECT residence,
                   COUNT(*) AS total,
                   COUNT(*) FILTER (WHERE COALESCE(living_status, %s) = %s) AS active,
                   COUNT(*) FILTER (WHERE living_status = 'ឈប់ស្នាក់នៅ') AS left_status,
                   COUNT(*) FILTER (WHERE living_status = 'នៅស្រុក') AS hometown
            FROM monk_tbl
            GROUP BY residence
        """, (_ACTIVE_LIVING_STATUS, _ACTIVE_LIVING_STATUS))
        counts = {
            r[0]: {
                'total': r[1],
                'active': r[2],
                'left': r[3],
                'hometown': r[4],
            }
            for r in cur.fetchall()
        }
        cur.close()
        conn.close()

        base = request.host_url.rstrip('/')
        empty = {'total': 0, 'active': 0, 'left': 0, 'hometown': 0}
        items = []
        for row in links:
            residence = row[1]
            c = counts.get(residence, empty)
            items.append({
                'id': row[0],
                'residence': residence,
                'residence_label': _residence_label(residence),
                'token': row[2],
                'label': row[3] or '',
                'is_active': row[4],
                'created_at': row[5].isoformat() if row[5] else None,
                'last_used_at': row[6].isoformat() if row[6] else None,
                'url': f'{base}/kuti/{row[2]}',
                'total': c['total'],
                'active': c['active'],
                'left': c['left'],
                'hometown': c['hometown'],
            })

        linked = {i['residence'] for i in items}
        missing = []
        for res in sorted(_valid_residences()):
            if res in linked:
                continue
            c = counts.get(res, empty)
            missing.append({
                'residence': res,
                'residence_label': _residence_label(res),
                'total': c['total'],
                'active': c['active'],
                'left': c['left'],
                'hometown': c['hometown'],
            })

        return jsonify({
            'success': True,
            'links': items,
            'missing': missing,
            'summary': {
                'links': len(items),
                'missing': len(missing),
                'monks_active': sum(i['active'] for i in items),
                'monks_total': sum(i['total'] for i in items),
            },
        })
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 500


def _make_kuti_status_export_html(label, monks, stay, home, left, sick_hosp=None, sick_home=None):
    """Pagoda-formatted HTML report — columns by living status."""
    import html as _html
    from datetime import date

    sick_hosp = sick_hosp or []
    sick_home = sick_home or []
    today = date.today().strftime('%d/%m/%Y')
    bhikkhu = sum(1 for m in monks if m['monk_type'] == 'ភិក្ខុ')
    samanera = sum(1 for m in monks if m['monk_type'] == 'សាមណេរ')

    def _row(m, idx):
        edu = _html.escape(f"{m['education_level']} {m['academic_year']}".strip()) or '—'
        return (
            f'<tr>'
            f'<td class="c">{idx}</td>'
            f'<td class="name">{_html.escape(m["fullname"])}</td>'
            f'<td class="c">{m["vassa_years"]}</td>'
            f'<td>{_html.escape(m["monk_type"] or "—")}</td>'
            f'<td>{_html.escape(m["position"] or "—")}</td>'
            f'<td>{edu}</td>'
            f'</tr>'
        )

    def _col(title, group, accent, head_bg, row_bg):
        rows = ''.join(_row(m, i + 1) for i, m in enumerate(group))
        empty = '<tr><td colspan="6" class="empty">មិនមាន</td></tr>'
        return f'''
        <section class="status-col" style="--accent:{accent};--head-bg:{head_bg};--row-bg:{row_bg}">
            <header class="status-col-head">
                <span class="status-dot"></span>
                <h3>{_html.escape(title)}</h3>
                <span class="status-count">{len(group)}</span>
            </header>
            <table>
                <thead><tr>
                    <th class="c">#</th><th>ឈ្មោះ</th><th class="c">វស្សា</th>
                    <th>ប្រភេទ</th><th>តួនាទី</th><th>កម្រិតសិក្សា</th>
                </tr></thead>
                <tbody>{rows if group else empty}</tbody>
            </table>
        </section>'''

    return f'''<!DOCTYPE html>
<html lang="km"><head><meta charset="UTF-8">
<style>
@import url('https://fonts.googleapis.com/css2?family=Battambang:wght@400;700&family=Kantumruy+Pro:wght@400;600;700&family=Moul&display=swap');
@page {{ size: A4 portrait; margin: 12mm; }}
*, *::before, *::after {{ box-sizing: border-box; margin: 0; padding: 0; }}
body {{
    font-family: 'Kantumruy Pro', 'Battambang', sans-serif;
    color: #1a2332; background: #fff; font-size: 11px;
}}
.page {{ width: 210mm; min-height: 297mm; margin: 0 auto; padding: 10mm 12mm; }}

.masthead {{
    display: grid; grid-template-columns: 1fr auto 1fr; gap: 12px;
    align-items: center; padding-bottom: 10px;
    border-bottom: 2px solid #0c2d5a; margin-bottom: 12px;
}}
.mast-left, .mast-right {{ font-size: 9px; line-height: 1.45; color: #475569; }}
.mast-left p, .mast-right p {{ margin: 0; }}
.mast-pagoda {{ font-family: 'Moul', 'Battambang', serif; color: #0c2d5a; font-size: 10px; margin-top: 2px; }}
.mast-center {{ text-align: center; }}
.mast-seal {{
    width: 52px; height: 52px; border-radius: 50%;
    border: 2px solid #c9a227; background: #fff8e8;
    display: inline-flex; align-items: center; justify-content: center;
    font-family: 'Moul', serif; font-size: 9px; color: #0c2d5a;
}}
.mast-kingdom {{ font-family: 'Moul', serif; color: #0c2d5a; font-size: 10px; text-align: right; }}
.mast-motto {{ font-size: 9px; color: #64748b; text-align: right; }}

.report-title {{
    text-align: center; margin-bottom: 12px;
}}
.report-title h1 {{
    font-family: 'Moul', 'Battambang', serif;
    font-size: 16px; font-weight: 400; color: #0c2d5a; margin-bottom: 4px;
}}
.report-title p {{
    font-family: 'Battambang', serif;
    font-size: 14px; color: #64748b;
}}

.chips {{
    display: flex; gap: 8px; justify-content: center; margin-bottom: 14px; flex-wrap: wrap;
}}
.chip {{
    min-width: 88px; padding: 8px 12px; border-radius: 8px; text-align: center;
    border: 1px solid #e2e8f0; background: #f8fafc;
}}
.chip span {{ display: block; font-size: 9px; color: #64748b; margin-bottom: 2px; }}
.chip strong {{ font-size: 16px; color: #0c2d5a; }}
.chip-stay {{ background: #e6f5ee; border-color: #a7dcc4; }}
.chip-stay strong {{ color: #1f6b4a; }}
.chip-home {{ background: #fef3e2; border-color: #f5c98a; }}
.chip-home strong {{ color: #b45309; }}
.chip-left {{ background: #fee2e2; border-color: #f5a5a5; }}
.chip-left strong {{ color: #9b1c1c; }}
.chip-sick-hosp {{ background: #ede9fe; border-color: #c4b5fd; }}
.chip-sick-hosp strong {{ color: #6d28d9; }}
.chip-sick-home {{ background: #e0f2fe; border-color: #7dd3fc; }}
.chip-sick-home strong {{ color: #0369a1; }}

.cols {{
    display: grid; grid-template-columns: 1fr; gap: 12px; align-items: start;
}}
.status-col {{
    border: 1px solid #e2e8f0; border-radius: 8px; overflow: hidden;
    border-top: 3px solid var(--accent);
}}
.status-col-head {{
    display: flex; align-items: center; gap: 8px;
    padding: 8px 10px; background: var(--head-bg);
    border-bottom: 1px solid #e2e8f0;
}}
.status-dot {{
    width: 8px; height: 8px; border-radius: 50%; background: var(--accent); flex-shrink: 0;
}}
.status-col-head h3 {{
    flex: 1; font-size: 11px; font-weight: 700; color: var(--accent);
}}
.status-count {{
    min-width: 24px; padding: 2px 8px; border-radius: 999px;
    background: #fff; font-size: 10px; font-weight: 700; color: var(--accent);
    box-shadow: 0 0 0 1px rgba(0,0,0,.06); text-align: center;
}}
table {{ width: 100%; border-collapse: collapse; }}
th {{
    background: #0c2d5a; color: #fff; padding: 6px 5px;
    font-size: 9px; font-weight: 700; text-align: center;
}}
td {{ padding: 5px 5px; border-bottom: 1px solid #edf2f7; vertical-align: middle; font-size: 9.5px; }}
tbody tr:nth-child(even) {{ background: var(--row-bg); }}
td.c, th.c {{ text-align: center; }}
td.name {{ font-weight: 700; text-align: left; }}
td.empty {{ text-align: center; color: #94a3b8; padding: 16px; }}

.footer {{
    display: flex; justify-content: space-between; align-items: center;
    margin-top: 12px; padding-top: 8px; border-top: 1px solid #e2e8f0;
    font-size: 9px; color: #94a3b8;
}}
</style></head><body>
<div class="page">
    <div class="masthead">
        <div class="mast-left">
            <p>មន្ទីរធម្មការ និងសាសនា រាជធានីភ្នំពេញ</p>
            <p>សាលាពុទ្ធិកអនុវិទ្យាល័យសង្ឃ</p>
            <p class="mast-pagoda">វត្តនិរោធរង្សី</p>
        </div>
        <div class="mast-center"><div class="mast-seal">វត្ត</div></div>
        <div class="mast-right">
            <p class="mast-kingdom">ព្រះរាជាណាចក្រកម្ពុជា</p>
            <p class="mast-motto">ជាតិ · សាសនា · ព្រះមហាក្សត្រ</p>
        </div>
    </div>

    <div class="report-title">
        <h1>បញ្ជីព្រះសង្ឃ — {_html.escape(label)}</h1>
        <p>ថ្ងៃទី {today} · តាមស្ថានភាពស្នាក់នៅ · ភិក្ខុ {bhikkhu} · សាមណេរ {samanera}</p>
    </div>

    <div class="chips">
        <div class="chip chip-stay"><span>កំពុងស្នាក់នៅ</span><strong>{len(stay)}</strong></div>
        <div class="chip chip-home"><span>ទៅស្រុក</span><strong>{len(home)}</strong></div>
        <div class="chip chip-left"><span>ឈប់ស្នាក់នៅ</span><strong>{len(left)}</strong></div>
        <div class="chip chip-sick-hosp"><span>ឈឺនៅពេទ្យ</span><strong>{len(sick_hosp)}</strong></div>
        <div class="chip chip-sick-home"><span>ឈឺនៅស្រុក</span><strong>{len(sick_home)}</strong></div>
        <div class="chip"><span>សរុប</span><strong>{len(monks)}</strong></div>
    </div>

    <div class="cols">
        {_col('កំពុងស្នាក់នៅ', stay, '#1f6b4a', '#f0faf5', '#f7fdf9')}
        {_col('ទៅស្រុក', home, '#b45309', '#fffaf3', '#fffbf5')}
        {_col('ឈប់ស្នាក់នៅ', left, '#9b1c1c', '#fff5f5', '#fffafa')}
        {_col('ឈឺនៅពេទ្យ', sick_hosp, '#6d28d9', '#f5f3ff', '#faf9ff')}
        {_col('ឈឺនៅស្រុក', sick_home, '#0369a1', '#f0f9ff', '#f8fcff')}
    </div>

    <div class="footer">
        <span>វត្តនិរោធរង្សី — ពិនិត្យស្ថានភាពមេកុដិ</span>
        <span>ថ្ងៃទី {today}</span>
    </div>
</div>
</body></html>'''


@main_bp.route('/api/kuti-status/export', methods=['GET'])
def api_kuti_status_export():
    """Admin export of one kuti's monk list — excel (xlsx) or html preview."""
    import io
    import html as _html
    from datetime import date

    if not user_allowed(_session_user(), '/api/kuti-status'):
        return jsonify({'success': False, 'message': 'Forbidden'}), 403

    residence = (request.args.get('residence') or '').strip()
    fmt = (request.args.get('fmt') or 'excel').strip().lower()
    if residence not in _valid_residences():
        return jsonify({'success': False, 'message': 'កុដិមិនត្រឹមត្រូវ'}), 400
    if fmt not in ('excel', 'html'):
        return jsonify({'success': False, 'message': 'Invalid format'}), 400

    try:
        conn = connect_db()
        cur = conn.cursor()
        cur.execute("""
            SELECT fullname, vassa_years, monk_type, position, living_status,
                   education_level, academic_year
            FROM monk_tbl
            WHERE residence = %s
            ORDER BY
                CASE living_status
                     WHEN 'កំពុងស្នាក់នៅ' THEN 0
                     WHEN 'នៅស្រុក' THEN 1
                     WHEN 'ឈឺនៅពេទ្យ' THEN 2
                     WHEN 'ឈឺនៅស្រុក' THEN 3
                     ELSE 4 END,
                monk_type, position, vassa_years DESC, fullname
        """, (residence,))
        rows = cur.fetchall()
        cur.close()
        conn.close()
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 500

    monks = [{
        'fullname': r[0],
        'vassa_years': r[1],
        'monk_type': r[2] or '',
        'position': r[3] or '',
        'living_status': r[4] or _ACTIVE_LIVING_STATUS,
        'education_level': r[5] or '',
        'academic_year': r[6] or '',
    } for r in rows]

    stay = [m for m in monks if m['living_status'] == _ACTIVE_LIVING_STATUS]
    home = [m for m in monks if m['living_status'] == 'នៅស្រុក']
    left = [m for m in monks if m['living_status'] == 'ឈប់ស្នាក់នៅ']
    sick_hosp = [m for m in monks if m['living_status'] == 'ឈឺនៅពេទ្យ']
    sick_home = [m for m in monks if m['living_status'] == 'ឈឺនៅស្រុក']
    label = _residence_label(residence)

    if fmt == 'html':
        return _make_kuti_status_export_html(label, monks, stay, home, left, sick_hosp, sick_home), 200, {
            'Content-Type': 'text/html; charset=utf-8',
        }

    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter

    wb = openpyxl.Workbook()
    thin = Side(border_style='thin', color='D1D5DB')
    bdr = Border(left=thin, right=thin, top=thin, bottom=thin)
    hfont = Font(bold=True, color='FFFFFF', size=11, name='Calibri')
    dfont = Font(size=10, name='Calibri')
    hdrs = ['#', 'ឈ្មោះ', 'វស្សា', 'ប្រភេទ', 'តួនាទី', 'កម្រិតសិក្សា']

    sheets = [
        ('កំពុងស្នាក់នៅ', stay, '1F6B4A', 'E6F5EE'),
        ('ទៅស្រុក', home, 'B45309', 'FEF3E2'),
        ('ឈប់ស្នាក់នៅ', left, '9B1C1C', 'FEE2E2'),
        ('ឈឺនៅពេទ្យ', sick_hosp, '6D28D9', 'EDE9FE'),
        ('ឈឺនៅស្រុក', sick_home, '0369A1', 'E0F2FE'),
    ]

    first = True
    for sheet_name, group, hdr_hex, row_hex in sheets:
        ws = wb.active if first else wb.create_sheet(title=sheet_name[:31])
        if first:
            ws.title = sheet_name[:31]
            first = False
        hfill = PatternFill(start_color=hdr_hex, end_color=hdr_hex, fill_type='solid')
        rfill = PatternFill(start_color=row_hex, end_color=row_hex, fill_type='solid')

        label = _residence_label(residence)
        ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=len(hdrs))
        title = ws.cell(row=1, column=1,
                        value=f'{label} — {sheet_name} ({len(group)} នាក់)')
        title.font = Font(bold=True, size=13, color='0C2D5A', name='Calibri')
        title.alignment = Alignment(horizontal='center', vertical='center')
        ws.row_dimensions[1].height = 22

        ws.merge_cells(start_row=2, start_column=1, end_row=2, end_column=len(hdrs))
        meta = ws.cell(row=2, column=1,
                       value=f'ថ្ងៃទី {date.today().strftime("%d/%m/%Y")}  |  សរុបកុដិ {len(monks)} នាក់')
        meta.font = Font(size=10, color='718096', name='Calibri')
        meta.alignment = Alignment(horizontal='center')

        for col, h in enumerate(hdrs, 1):
            c = ws.cell(row=3, column=col, value=h)
            c.font = hfont
            c.fill = hfill
            c.border = bdr
            c.alignment = Alignment(horizontal='center', vertical='center')

        for i, m in enumerate(group, 1):
            vals = [i, m['fullname'], m['vassa_years'], m['monk_type'],
                    m['position'], m['education_level']]
            for col, val in enumerate(vals, 1):
                c = ws.cell(row=3 + i, column=col, value=val)
                c.fill = rfill
                c.border = bdr
                c.font = dfont
                c.alignment = Alignment(
                    horizontal='center' if col in (1, 3) else 'left',
                    vertical='center',
                )

        for col, w in enumerate([5, 28, 8, 12, 20, 16], 1):
            ws.column_dimensions[get_column_letter(col)].width = w

    # Overview sheet first
    overview = wb.create_sheet(title='សរុប', index=0)
    overview['A1'] = f'បញ្ជីព្រះសង្ឃ — {_residence_label(residence)}'
    overview['A1'].font = Font(bold=True, size=14, color='0C2D5A')
    overview['A3'] = 'ស្ថានភាព'
    overview['B3'] = 'ចំនួន'
    overview['A3'].font = overview['B3'].font = Font(bold=True)
    overview['A4'] = 'កំពុងស្នាក់នៅ'
    overview['B4'] = len(stay)
    overview['A5'] = 'ទៅស្រុក'
    overview['B5'] = len(home)
    overview['A6'] = 'ឈប់ស្នាក់នៅ'
    overview['B6'] = len(left)
    overview['A7'] = 'សរុប'
    overview['B7'] = len(monks)
    overview['A7'].font = overview['B7'].font = Font(bold=True)
    overview.column_dimensions['A'].width = 22
    overview.column_dimensions['B'].width = 12

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    safe = residence.replace(' ', '_').replace('/', '-')
    return send_file(
        buf,
        mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        as_attachment=True,
        download_name=f"kuti_status_{safe}_{date.today().isoformat()}.xlsx",
    )


@main_bp.route('/api/kuti-links/monks', methods=['GET'])
def kuti_links_monks():
    """Admin: list all monks in one residence (including departed statuses)."""
    if not user_allowed(_session_user(), '/api/kuti-links'):
        return jsonify({'success': False, 'message': 'Forbidden'}), 403
    residence = (request.args.get('residence') or '').strip()
    if residence not in _valid_residences():
        return jsonify({'success': False, 'message': 'កុដិមិនត្រឹមត្រូវ'}), 400
    try:
        conn = connect_db()
        cur = conn.cursor()
        cur.execute("""
            SELECT id, fullname, vassa_years, monk_type, residence, position,
                   education_level, academic_year, living_status, created_at, updated_at
            FROM monk_tbl
            WHERE residence = %s
            ORDER BY
                CASE living_status WHEN 'កំពុងស្នាក់នៅ' THEN 0 ELSE 1 END,
                monk_type, position, vassa_years DESC, fullname
        """, (residence,))
        rows = cur.fetchall()
        cur.close()
        conn.close()
        monks = [{
            'id': r[0],
            'fullname': r[1],
            'vassa_years': r[2],
            'monk_type': r[3],
            'residence': r[4],
            'residence_label': _residence_label(r[4]),
            'position': r[5],
            'education_level': r[6],
            'academic_year': r[7],
            'living_status': r[8] or _ACTIVE_LIVING_STATUS,
            'created_at': r[9].isoformat() if r[9] else None,
            'updated_at': r[10].isoformat() if r[10] else None,
        } for r in rows]
        return jsonify({
            'success': True,
            'residence': residence,
            'residence_label': _residence_label(residence),
            'count': len(monks),
            'monks': monks,
        })
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 500


@main_bp.route('/api/kuti-links/leaders', methods=['GET'])
def kuti_leaders_for_residence():
    """Return monks with position មេកុដិ (and អនុកុដិ) in one residence."""
    if not user_allowed(_session_user(), '/api/kuti-links'):
        return jsonify({'success': False, 'message': 'Forbidden'}), 403
    residence = (request.args.get('residence') or '').strip()
    if residence not in _valid_residences():
        return jsonify({'success': False, 'message': 'កុដិមិនត្រឹមត្រូវ'}), 400
    try:
        conn = connect_db()
        cur = conn.cursor()
        cur.execute("""
            SELECT id, fullname, position
            FROM monk_tbl
            WHERE residence = %s
              AND position IN ('មេកុដិ', 'អនុកុដិ')
            ORDER BY
                CASE position WHEN 'មេកុដិ' THEN 0 ELSE 1 END,
                fullname
        """, (residence,))
        rows = cur.fetchall()
        cur.close()
        conn.close()
        leaders = [{
            'id': r[0],
            'fullname': r[1],
            'position': r[2],
        } for r in rows]
        return jsonify({'success': True, 'leaders': leaders})
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 500


@main_bp.route('/api/kuti-links', methods=['GET'])
def list_kuti_links():
    if not user_allowed(_session_user(), '/api/kuti-links'):
        return jsonify({'success': False, 'message': 'Forbidden'}), 403
    try:
        conn = connect_db()
        cur = conn.cursor()
        cur.execute("""
            SELECT id, residence, token, label, is_active, created_at, last_used_at
            FROM kuti_share_links
            ORDER BY residence, created_at DESC
        """)
        rows = cur.fetchall()
        cur.close()
        conn.close()
        base = request.host_url.rstrip('/')
        links = [{
            'id': r[0],
            'residence': r[1],
            'residence_label': _residence_label(r[1]),
            'token': r[2],
            'label': r[3] or '',
            'is_active': r[4],
            'created_at': r[5].isoformat() if r[5] else None,
            'last_used_at': r[6].isoformat() if r[6] else None,
            'url': f'{base}/kuti/{r[2]}',
        } for r in rows]
        return jsonify({'success': True, 'links': links})
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 500


@main_bp.route('/api/kuti-links', methods=['POST'])
def create_kuti_link():
    """Create (or recreate) an active share link for one residence."""
    if not user_allowed(_session_user(), '/api/kuti-links'):
        return jsonify({'success': False, 'message': 'Forbidden'}), 403
    data = request.get_json(silent=True) or {}
    residence = str(data.get('residence', '') or '').strip()
    label = str(data.get('label', '') or '').strip() or None
    if residence not in _valid_residences():
        return jsonify({'success': False, 'message': 'កុដិមិនត្រឹមត្រូវ'}), 400
    try:
        token = secrets.token_urlsafe(24)
        conn = connect_db()
        cur = conn.cursor()
        # Deactivate previous active links for this kuti
        cur.execute("""
            UPDATE kuti_share_links SET is_active = FALSE
            WHERE residence = %s AND is_active = TRUE
        """, (residence,))
        cur.execute("""
            INSERT INTO kuti_share_links (residence, token, label, is_active)
            VALUES (%s, %s, %s, TRUE)
            RETURNING id, residence, token, label, created_at
        """, (residence, token, label))
        row = cur.fetchone()
        conn.commit()
        cur.close()
        conn.close()
        base = request.host_url.rstrip('/')
        _log_act('kuti_link_create', 'kuti_links', f'{residence} — {token[:8]}…')
        return jsonify({
            'success': True,
            'link': {
                'id': row[0],
                'residence': row[1],
                'residence_label': _residence_label(row[1]),
                'token': row[2],
                'label': row[3] or '',
                'is_active': True,
                'created_at': row[4].isoformat() if row[4] else None,
                'url': f'{base}/kuti/{row[2]}',
            },
        })
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 500


@main_bp.route('/api/kuti-links/<int:link_id>', methods=['DELETE'])
def delete_kuti_link(link_id):
    """Permanently remove a share-link record."""
    if not user_allowed(_session_user(), '/api/kuti-links'):
        return jsonify({'success': False, 'message': 'Forbidden'}), 403
    try:
        conn = connect_db()
        cur = conn.cursor()
        cur.execute("DELETE FROM kuti_share_links WHERE id = %s", (link_id,))
        conn.commit()
        deleted = cur.rowcount
        cur.close()
        conn.close()
        if not deleted:
            return jsonify({'success': False, 'message': 'រកមិនឃើញតំណ'}), 404
        _log_act('kuti_link_delete', 'kuti_links', f'id={link_id}')
        return jsonify({'success': True})
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 500


@main_bp.route('/kuti/<token>')
def kuti_public_view(token):
    """Public kuti leader view (via share link) — manage monks in this kuti only."""
    share = _get_share_by_token(token)
    if not share:
        return render_template('kuti_view.html', invalid=True, share=None), 404
    return render_template(
        'kuti_view.html',
        invalid=False,
        share=share,
        residence=share['residence'],
        residence_label=_residence_label(share['residence']),
        living_statuses=list(LIVING_STATUSES),
    )


@main_bp.route('/api/kuti/<token>/monks', methods=['GET'])
def kuti_public_monks(token):
    """Return monks for the token's residence only — never other kutis."""
    share = _get_share_by_token(token)
    if not share:
        return jsonify({'success': False, 'message': 'តំណភ្ជាប់មិនត្រឹមត្រូវ ឬត្រូវបានបិទ'}), 404
    residence = share['residence']
    try:
        conn = connect_db()
        cur = conn.cursor()
        cur.execute("""
            SELECT id, fullname, vassa_years, monk_type, residence, position,
                   education_level, academic_year, created_at, living_status
            FROM monk_tbl
            WHERE residence = %s
            ORDER BY
                CASE living_status WHEN 'កំពុងស្នាក់នៅ' THEN 0 ELSE 1 END,
                monk_type, position, vassa_years DESC, fullname
        """, (residence,))
        rows = cur.fetchall()
        cur.close()
        conn.close()
        monks = [{
            'id': r[0],
            'fullname': r[1],
            'vassa_years': r[2],
            'monk_type': r[3],
            'residence': r[4],
            'residence_label': _residence_label(r[4]),
            'position': r[5],
            'education_level': r[6],
            'academic_year': r[7],
            'created_at': r[8].isoformat() if r[8] else None,
            'living_status': r[9] or _ACTIVE_LIVING_STATUS,
        } for r in rows]
        monks = sort_attendance_monks(monks)
        monks = [m for m in monks if m['residence'] == residence]
        return jsonify({
            'success': True,
            'residence': residence,
            'residence_label': _residence_label(residence),
            'label': share.get('label') or '',
            'count': len(monks),
            'monks': monks,
        })
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 500


def _kuti_monk_payload(data, forced_residence):
    """Parse monk JSON for token-scoped writes; residence is always forced."""
    fullname = (data.get('fullname') or '').strip()
    vassa_years = data.get('total-monk')
    monk_type = data.get('type')
    position = data.get('position')
    education_level = data.get('education_level')
    academic_year = data.get('academic_level')
    living_status = data.get('living_status') or _ACTIVE_LIVING_STATUS
    try:
        vassa_years = int(vassa_years)
    except (TypeError, ValueError):
        return None, 'ចំនួនវស្សាមិនត្រឹមត្រូវ'
    if not all([fullname, monk_type, position, education_level, academic_year]):
        return None, 'Missing required fields'
    dropdown_errors = _validate_entry_dropdowns(
        monk_type, forced_residence, position, education_level, academic_year,
    )
    if dropdown_errors:
        return None, dropdown_errors[0]
    if living_status not in _VALID_LIVING_STATUSES:
        return None, 'ស្ថានភាពមិនត្រឹមត្រូវ'
    return {
        'fullname': fullname,
        'vassa_years': vassa_years,
        'monk_type': monk_type,
        'residence': forced_residence,
        'position': position,
        'education_level': education_level,
        'academic_year': academic_year,
        'living_status': living_status,
    }, None


def _monk_in_residence(monk_id, residence):
    conn = connect_db()
    cur = conn.cursor()
    cur.execute("SELECT id FROM monk_tbl WHERE id = %s AND residence = %s", (monk_id, residence))
    row = cur.fetchone()
    cur.close()
    conn.close()
    return bool(row)


def _queue_kuti_pending(payload, target_monk_id=None):
    """Mekuti add/edit goes to the admin approval queue (INSERT or UPDATE)."""
    if target_monk_id is None:
        conn = connect_db()
        cur = conn.cursor()
        cur.execute("""
            SELECT id FROM monk_tbl
            WHERE residence = %s AND LOWER(fullname) = LOWER(%s)
            LIMIT 1
        """, (payload['residence'], payload['fullname']))
        row = cur.fetchone()
        cur.close()
        conn.close()
        if row:
            target_monk_id = row[0]
    return insert_pending_submission(
        payload['fullname'], payload['vassa_years'], payload['monk_type'],
        payload['residence'], payload['position'], payload['education_level'],
        payload['academic_year'], source='kuti', target_monk_id=target_monk_id,
    )


@main_bp.route('/api/kuti/<token>/monks', methods=['POST'])
def kuti_public_add_monk(token):
    """មេកុដិ: queue a new monk for admin approval (INSERT)."""
    share = _get_share_by_token(token)
    if not share:
        return jsonify({'success': False, 'message': 'តំណភ្ជាប់មិនត្រឹមត្រូវ ឬត្រូវបានបិទ'}), 404
    data = request.get_json(silent=True) or {}
    payload, err = _kuti_monk_payload(data, share['residence'])
    if err:
        return jsonify({'success': False, 'message': err}), 400
    sub_id = _queue_kuti_pending(payload)
    if not sub_id:
        return jsonify({'success': False, 'message': 'មិនអាចដាក់ស្នើបាន'}), 500
    return jsonify({
        'success': True,
        'pending': True,
        'pending_id': sub_id,
        'message': 'បានដាក់ស្នើ — រង់ចាំអនុម័តពីអ្នកគ្រប់គ្រង',
    })


@main_bp.route('/api/kuti/<token>/monks/<int:monk_id>', methods=['PUT'])
def kuti_public_update_monk(token, monk_id):
    """មេកុដិ: queue an edit for admin approval (UPDATE)."""
    share = _get_share_by_token(token)
    if not share:
        return jsonify({'success': False, 'message': 'តំណភ្ជាប់មិនត្រឹមត្រូវ ឬត្រូវបានបិទ'}), 404
    residence = share['residence']
    if not _monk_in_residence(monk_id, residence):
        return jsonify({'success': False, 'message': 'មិនមានសិទ្ធិកែព្រះសង្ឃនេះ'}), 403
    data = request.get_json(silent=True) or {}
    payload, err = _kuti_monk_payload(data, residence)
    if err:
        return jsonify({'success': False, 'message': err}), 400
    sub_id = _queue_kuti_pending(payload, target_monk_id=monk_id)
    if not sub_id:
        return jsonify({'success': False, 'message': 'មិនអាចដាក់ស្នើបាន'}), 500
    return jsonify({
        'success': True,
        'pending': True,
        'pending_id': sub_id,
        'message': 'បានដាក់ស្នើកែប្រែ — រង់ចាំអនុម័តពីអ្នកគ្រប់គ្រង',
    })


@main_bp.route('/api/kuti/<token>/monks/<int:monk_id>/living-status', methods=['PATCH'])
def kuti_public_patch_living_status(token, monk_id):
    """មេកុដិ: change living status for a monk in this kuti only."""
    share = _get_share_by_token(token)
    if not share:
        return jsonify({'success': False, 'message': 'តំណភ្ជាប់មិនត្រឹមត្រូវ ឬត្រូវបានបិទ'}), 404
    residence = share['residence']
    if not _monk_in_residence(monk_id, residence):
        return jsonify({'success': False, 'message': 'មិនមានសិទ្ធិកែព្រះសង្ឃនេះ'}), 403
    data = request.get_json(silent=True) or {}
    living_status = (data.get('living_status') or '').strip()
    if living_status not in _VALID_LIVING_STATUSES:
        return jsonify({'success': False, 'message': 'ស្ថានភាពមិនត្រឹមត្រូវ'}), 400
    ok = update_monk_living_status(monk_id, living_status)
    if not ok:
        return jsonify({'success': False, 'message': 'រកមិនឃើញព្រះសង្ឃ'}), 404
    return jsonify({'success': True, 'living_status': living_status})


@main_bp.route('/api/kuti/<token>/export', methods=['GET'])
def kuti_public_export(token):
    """Export monks for this share-link residence only (docx / excel / html→PDF)."""
    import io
    import html as _html
    from datetime import date

    share = _get_share_by_token(token)
    if not share:
        return jsonify({'success': False, 'message': 'តំណភ្ជាប់មិនត្រឹមត្រូវ ឬត្រូវបានបិទ'}), 404

    fmt = (request.args.get('fmt') or 'docx').strip().lower()
    if fmt not in ('docx', 'excel', 'html'):
        return jsonify({'success': False, 'message': 'Invalid format'}), 400

    residence = share['residence']
    residence_label = _residence_label(residence)
    try:
        conn = connect_db()
        cur = conn.cursor()
        cur.execute("""
            SELECT fullname, vassa_years, monk_type, residence, position,
                   education_level, academic_year, living_status, created_at
            FROM monk_tbl
            WHERE residence = %s
            ORDER BY
                CASE living_status WHEN 'កំពុងស្នាក់នៅ' THEN 0 ELSE 1 END,
                monk_type, position, vassa_years DESC, fullname
        """, (residence,))
        rows = cur.fetchall()
        cur.close()
        conn.close()
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 500

    monks = [{
        'fullname': r[0],
        'vassa_years': r[1],
        'monk_type': r[2],
        'residence': (r[3] or '').replace('_', ' '),
        'position': r[4] or '',
        'education_level': r[5] or '',
        'academic_year': r[6] or '',
        'living_status': r[7] or _ACTIVE_LIVING_STATUS,
        'created_at': r[8],
    } for r in rows]
    monks = sort_attendance_monks(monks)

    today = date.today().strftime('%d/%m/%Y')
    safe_name = residence.replace(' ', '_').replace('/', '-')
    fname_base = f"kuti_{safe_name}_{date.today().isoformat()}"
    leader = share.get('label') or ''
    title = f'បញ្ជីព្រះសង្ឃ — {residence_label}'

    if fmt == 'docx':
        from docx import Document
        from docx.shared import Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        def _shade(cell, hex_color):
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), hex_color)
            tcPr.append(shd)

        doc = Document()
        sec = doc.sections[0]
        sec.page_width = Cm(21.0)   # A4
        sec.page_height = Cm(29.7)
        sec.left_margin = sec.right_margin = Cm(1.5)
        sec.top_margin = sec.bottom_margin = Cm(1.5)

        h = doc.add_heading(title, 0)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub_txt = f"ថ្ងៃទី {today}  |  ចំនួនសរុប: {len(monks)} នាក់"
        if leader:
            sub_txt += f"  |  មេកុដិ: {leader}"
        sub = doc.add_paragraph(sub_txt)
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub.runs[0].font.size = Pt(10)
        sub.runs[0].font.color.rgb = RGBColor(0x71, 0x80, 0x96)
        doc.add_paragraph()

        hdrs = ['#', 'ឈ្មោះ', 'វស្សា', 'ប្រភេទ', 'តួនាទី', 'ស្ថានភាព', 'កម្រិតសិក្សា']
        widths = [0.8, 3.6, 1.3, 1.6, 3.0, 2.4, 2.6]
        tbl = doc.add_table(rows=1, cols=len(hdrs))
        tbl.style = 'Table Grid'
        for i, (hdr, w) in enumerate(zip(hdrs, widths)):
            cell = tbl.rows[0].cells[i]
            cell.width = Cm(w)
            _shade(cell, '0C2D5A')
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(hdr)
            run.bold = True
            run.font.size = Pt(8.5)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

        for idx, m in enumerate(monks, 1):
            bg = 'FFF8E1' if m['monk_type'] == 'ភិក្ខុ' else 'F1F8E9'
            if m['living_status'] != _ACTIVE_LIVING_STATUS:
                bg = 'F8F9FA'
            vals = [
                str(idx), m['fullname'], f"{m['vassa_years']} ឆ្នាំ", m['monk_type'],
                m['position'], m['living_status'], m['education_level'],
            ]
            row = tbl.add_row()
            for j, (val, w) in enumerate(zip(vals, widths)):
                cell = row.cells[j]
                cell.width = Cm(w)
                _shade(cell, bg)
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if j in (0, 2) else WD_ALIGN_PARAGRAPH.LEFT
                run = p.add_run(val)
                run.font.size = Pt(8.5)
                if j == 1:
                    run.bold = True

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return send_file(
            buf,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=f"{fname_base}.docx",
        )

    if fmt == 'excel':
        import openpyxl
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        from openpyxl.utils import get_column_letter

        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = 'កុដិ'
        hdrs = ['#', 'ឈ្មោះ', 'វស្សា', 'ប្រភេទ', 'តួនាទី', 'ស្ថានភាព',
                'កម្រិតសិក្សា', 'ថ្នាក់', 'ថ្ងៃបញ្ចូល']
        hfill = PatternFill(start_color='0C2D5A', end_color='0C2D5A', fill_type='solid')
        hfont = Font(bold=True, color='FFFFFF', size=11, name='Calibri')
        thin = Side(border_style='thin', color='D1D5DB')
        bdr = Border(left=thin, right=thin, top=thin, bottom=thin)

        ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=len(hdrs))
        title_cell = ws.cell(row=1, column=1, value=title)
        title_cell.font = Font(bold=True, size=14, color='0C2D5A', name='Calibri')
        title_cell.alignment = Alignment(horizontal='center', vertical='center')
        ws.row_dimensions[1].height = 24

        meta = f"ថ្ងៃទី {today}  |  សរុប {len(monks)} នាក់"
        if leader:
            meta += f"  |  មេកុដិ: {leader}"
        ws.merge_cells(start_row=2, start_column=1, end_row=2, end_column=len(hdrs))
        meta_cell = ws.cell(row=2, column=1, value=meta)
        meta_cell.alignment = Alignment(horizontal='center')
        meta_cell.font = Font(size=10, color='718096', name='Calibri')

        for col, h in enumerate(hdrs, 1):
            c = ws.cell(row=3, column=col, value=h)
            c.font = hfont
            c.fill = hfill
            c.border = bdr
            c.alignment = Alignment(horizontal='center', vertical='center')
        ws.row_dimensions[3].height = 22

        bfill = PatternFill(start_color='FFF8E1', end_color='FFF8E1', fill_type='solid')
        sfill = PatternFill(start_color='F1F8E9', end_color='F1F8E9', fill_type='solid')
        gfill = PatternFill(start_color='F8F9FA', end_color='F8F9FA', fill_type='solid')
        dfont = Font(size=10, name='Calibri')

        for row_i, m in enumerate(monks, 4):
            if m['living_status'] != _ACTIVE_LIVING_STATUS:
                fill = gfill
            else:
                fill = bfill if m['monk_type'] == 'ភិក្ខុ' else sfill
            cre = m['created_at'].strftime('%d/%m/%Y') if m['created_at'] else '—'
            vals = [
                row_i - 3, m['fullname'], m['vassa_years'], m['monk_type'],
                m['position'], m['living_status'], m['education_level'],
                m['academic_year'], cre,
            ]
            for col, val in enumerate(vals, 1):
                c = ws.cell(row=row_i, column=col, value=val)
                c.fill = fill
                c.border = bdr
                c.font = dfont
                c.alignment = Alignment(
                    horizontal='center' if col in (1, 3, 9) else 'left',
                    vertical='center',
                )

        for col, w in enumerate([5, 24, 8, 12, 18, 16, 15, 10, 13], 1):
            ws.column_dimensions[get_column_letter(col)].width = w

        buf = io.BytesIO()
        wb.save(buf)
        buf.seek(0)
        return send_file(
            buf,
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            as_attachment=True,
            download_name=f"{fname_base}.xlsx",
        )

    # html → used by client PDF export
    bhikkhus = [m for m in monks if m['monk_type'] == 'ភិក្ខុ']
    samaneras = [m for m in monks if m['monk_type'] == 'សាមណេរ']

    def _trow(m, idx, stripe):
        bg = '#f8fafc' if stripe else '#ffffff'
        if m['living_status'] != _ACTIVE_LIVING_STATUS:
            bg = '#f1f3f5'
        return (
            f'<tr style="background:{bg}">'
            f'<td class="c">{idx}</td>'
            f'<td class="name">{_html.escape(m["fullname"])}</td>'
            f'<td class="c">{m["vassa_years"]}</td>'
            f'<td>{_html.escape(m["monk_type"])}</td>'
            f'<td>{_html.escape(m["position"])}</td>'
            f'<td>{_html.escape(m["living_status"])}</td>'
            f'<td>{_html.escape(m["education_level"])}</td>'
            f'</tr>'
        )

    def _section(section_monks, sec_title, hdr_bg):
        if not section_monks:
            return ''
        rows = ''.join(_trow(m, i + 1, i % 2 == 1) for i, m in enumerate(section_monks))
        return f'''
        <div class="sec-title" style="background:{hdr_bg}">{_html.escape(sec_title)}
            <span class="sec-count">({len(section_monks)} នាក់)</span></div>
        <table>
            <thead><tr>
                <th class="c">#</th><th>ឈ្មោះ</th><th class="c">វស្សា</th>
                <th>ប្រភេទ</th><th>តួនាទី</th><th>ស្ថានភាព</th><th>កម្រិតសិក្សា</th>
            </tr></thead>
            <tbody>{rows}</tbody>
        </table>'''

    leader_html = f' · មេកុដិ៖ {_html.escape(leader)}' if leader else ''
    html_out = f'''<!DOCTYPE html><html lang="km"><head><meta charset="UTF-8">
    <style>
    @import url('https://fonts.googleapis.com/css2?family=Battambang:wght@400;700&family=Moul&display=swap');
    @page {{ size: A4 portrait; margin: 12mm; }}
    * {{ box-sizing: border-box; }}
    html, body {{
      margin: 0; padding: 0;
      background: #fff;
      font-family: 'Battambang', sans-serif;
      color: #1a2332;
      font-size: 11px;
    }}
    .page {{
      width: 210mm;
      min-height: 297mm;
      margin: 0 auto;
      padding: 14mm 12mm;
      background: #fff;
      display: flex;
      flex-direction: column;
      align-items: stretch;
      justify-content: flex-start;
    }}
    .page-inner {{
      width: 100%;
      margin: 0 auto;
    }}
    .header {{
      background: linear-gradient(135deg, #0c2d5a, #163f73);
      color: #fff;
      padding: 14px 18px;
      border-radius: 10px;
      margin-bottom: 14px;
      text-align: center;
    }}
    .header h1 {{
      font-family: 'Moul', 'Battambang', serif;
      font-size: 16px; margin: 0 0 6px; font-weight: 400;
    }}
    .header p {{
      font-family: 'Battambang', serif;
      margin: 0; opacity: .9; font-size: 14px;
    }}
    .chips {{ display: flex; gap: 10px; margin-bottom: 14px; justify-content: center; }}
    .chip {{
      flex: 1; max-width: 160px; padding: 10px; border-radius: 8px;
      text-align: center; background: #f0f4fa;
    }}
    .chip strong {{ display: block; font-size: 18px; color: #0c2d5a; }}
    .sec-title {{
      padding: 8px 12px; border-radius: 6px; color: #fff;
      font-weight: 700; margin: 14px 0 8px; text-align: center;
    }}
    .sec-count {{ font-weight: 400; opacity: .9; }}
    table {{ width: 100%; border-collapse: collapse; margin: 0 auto 8px; }}
    th {{
      background: #0c2d5a; color: #fff; padding: 8px 10px;
      text-align: center; font-size: 11px;
    }}
    td {{ padding: 7px 10px; border-bottom: 1px solid #e8edf3; vertical-align: middle; }}
    td.c, th.c {{ text-align: center; }}
    td.name {{ font-weight: 700; text-align: left; }}
    </style></head><body>
    <div class="page"><div class="page-inner">
    <div class="header">
      <h1>{_html.escape(title)}</h1>
      <p>ថ្ងៃទី {today}{leader_html} · វត្តនិរោធរង្សី</p>
    </div>
    <div class="chips">
      <div class="chip"><span>សរុប</span><strong>{len(monks)}</strong></div>
      <div class="chip"><span>ភិក្ខុ</span><strong>{len(bhikkhus)}</strong></div>
      <div class="chip"><span>សាមណេរ</span><strong>{len(samaneras)}</strong></div>
    </div>
    {_section(bhikkhus, 'ភិក្ខុ', '#8a6d12')}
    {_section(samaneras, 'សាមណេរ', '#1f6b4a')}
    </div></div>
    </body></html>'''
    return html_out, 200, {'Content-Type': 'text/html; charset=utf-8'}


# ============ PUBLIC SUBMISSION ROUTES ============

@main_bp.route('/submit', methods=['GET'])
def public_submit_page():
    success = request.args.get('success') == '1'
    return render_template('submit.html', success=success, error=None, form_data={})


@main_bp.route('/public/submit', methods=['POST'])
def public_submit():
    ip = request.headers.get('X-Forwarded-For', request.remote_addr or '').split(',')[0].strip()

    if not _rate_limit_ok(ip):
        return jsonify({'success': False, 'message': 'ការដាក់ស្នើច្រើនពេក។ សូមព្យាយាមមកវិញក្រោយ ១ នាទី។'}), 429

    data = request.get_json(silent=True) or {}

    fullname        = str(data.get('fullname',        '') or '').strip()
    vassa_years     = str(data.get('total-monk',      '') or '').strip()
    monk_type       = str(data.get('type',            '') or '').strip()
    residence       = str(data.get('home',            '') or '').strip()
    position        = str(data.get('position',        '') or '').strip()
    education_level = str(data.get('education_level', '') or '').strip()
    academic_year   = str(data.get('academic_level',  '') or '').strip()

    errors = []
    if not fullname or len(fullname) > 200:
        errors.append('នាមត្រកូល និង នាម មិនត្រឹមត្រូវ')

    vassa_int = None
    try:
        vassa_int = int(vassa_years)
        if vassa_int < 1 or vassa_int > 100:
            raise ValueError
    except (ValueError, TypeError):
        errors.append('ចំនួនវស្សា ត្រូវតែជាលេខចន្លោះ ១ ដល់ ១០០')

    errors.extend(_validate_entry_dropdowns(
        monk_type, residence, position, education_level, academic_year,
    ))

    if errors:
        return jsonify({'success': False, 'message': ' | '.join(errors)}), 422

    # Guard against duplicate pending submissions for the same name
    try:
        _conn = connect_db()
        _cur  = _conn.cursor()
        _cur.execute(
            "SELECT id FROM pending_submissions WHERE LOWER(fullname)=LOWER(%s) AND status='pending' LIMIT 1",
            (fullname,)
        )
        if _cur.fetchone():
            _cur.close(); _conn.close()
            return jsonify({'success': False,
                            'message': 'ឈ្មោះនេះ​មាន​ក្នុង​បញ្ជីរ​ង់ចាំ​ហើយ។ សូម​រង់ចាំ​ការ​អនុម័ត។'}), 409
        _cur.close(); _conn.close()
    except Exception:
        pass

    try:
        sub_id = insert_pending_submission(
            fullname, vassa_int, monk_type, residence,
            position, education_level, academic_year, source='public',
        )
        if sub_id:
            return jsonify({'success': True, 'pending_id': sub_id})
        return jsonify({'success': False, 'message': 'មានបញ្ហាក្នុងការរក្សាទុក។ សូមព្យាយាមមកវិញ។'}), 500
    except Exception:
        return jsonify({'success': False, 'message': 'កំហុសសេវាកម្ម។ សូមព្យាយាមមកវិញ។'}), 500


# ============ APPROVAL QUEUE ============

@main_bp.route('/approve')
def approve_page():
    conn = None
    try:
        conn = connect_db()
        cursor = conn.cursor()
        cursor.execute("""
            SELECT
                ps.id, ps.fullname, ps.vassa_years, ps.monk_type,
                ps.residence, ps.position, ps.education_level, ps.academic_year,
                ps.submitted_at,
                COALESCE(ps.target_monk_id,
                    (SELECT m.id FROM monk_tbl m
                     WHERE LOWER(m.fullname) = LOWER(ps.fullname) LIMIT 1)
                ) AS matched_id,
                COALESCE(ps.source, 'public') AS source
            FROM pending_submissions ps
            WHERE ps.status = 'pending'
            ORDER BY ps.submitted_at DESC;
        """)
        rows = cursor.fetchall()
        cursor.close()
        submissions = []
        for r in rows:
            submissions.append({
                'id':              r[0],
                'fullname':        r[1],
                'vassa_years':     r[2],
                'monk_type':       r[3],
                'residence':       r[4],
                'position':        r[5],
                'education_level': r[6],
                'academic_year':   r[7],
                'submitted_at':    r[8].strftime('%d/%m/%Y %H:%M') if r[8] else '',
                'matched':         r[9] is not None,
                'source':          r[10] or 'public',
                'residence_label': _residence_label(r[4]),
            })
        return render_template('approve.html', submissions=submissions, error=None)
    except Exception as e:
        return render_template('approve.html', submissions=[], error=str(e))
    finally:
        if conn:
            conn.close()


@main_bp.route('/api/submissions/bulk-action', methods=['POST'])
def submissions_bulk_action():
    data           = request.get_json(silent=True) or {}
    action         = str(data.get('action', '')).strip()
    ids            = data.get('ids', [])
    rejection_note = str(data.get('rejection_note', '') or '').strip()[:500] or None

    valid_actions = {'approve-selected', 'approve-all', 'reject-selected', 'reject-all'}
    if action not in valid_actions:
        return jsonify({'success': False, 'message': 'សកម្មភាព​មិន​ត្រឹម​ត្រូវ'}), 400

    if action.endswith('-selected') and not ids:
        return jsonify({'success': False, 'message': 'សូម​រើស​យ៉ាង​ហោច​ណាស់​មួយ​ជួរ'}), 400

    conn = None
    try:
        conn = connect_db()
        conn.autocommit = False
        cursor = conn.cursor()

        if action in ('reject-all', 'reject-selected'):
            if action == 'reject-all':
                cursor.execute("""
                    UPDATE pending_submissions
                    SET status = 'rejected', rejection_note = %s, reviewed_at = NOW()
                    WHERE status = 'pending'
                """, (rejection_note,))
            else:
                cursor.execute("""
                    UPDATE pending_submissions
                    SET status = 'rejected', rejection_note = %s, reviewed_at = NOW()
                    WHERE status = 'pending' AND id = ANY(%s)
                """, (rejection_note, ids))
            count = cursor.rowcount
            conn.commit()
            return jsonify({'success': True, 'count': count})

        # Approve path — fetch rows to process
        if action == 'approve-all':
            cursor.execute("""
                SELECT id, fullname, vassa_years, monk_type, residence,
                       position, education_level, academic_year, target_monk_id
                FROM pending_submissions
                WHERE status = 'pending'
                FOR UPDATE
            """)
        else:
            cursor.execute("""
                SELECT id, fullname, vassa_years, monk_type, residence,
                       position, education_level, academic_year, target_monk_id
                FROM pending_submissions
                WHERE status = 'pending' AND id = ANY(%s)
                FOR UPDATE
            """, (ids,))
        rows = cursor.fetchall()

        approved = 0
        for row in rows:
            sub_id, fullname, vassa_years, monk_type, residence, \
                position, education_level, academic_year, target_monk_id = row

            match_id = target_monk_id
            if not match_id:
                cursor.execute(
                    "SELECT id FROM monk_tbl WHERE LOWER(fullname) = LOWER(%s) LIMIT 1",
                    (fullname,)
                )
                found = cursor.fetchone()
                match_id = found[0] if found else None
            if match_id:
                cursor.execute("""
                    UPDATE monk_tbl
                    SET fullname=%s, vassa_years=%s, monk_type=%s, residence=%s, position=%s,
                        education_level=%s, academic_year=%s, updated_at=NOW()
                    WHERE id=%s
                """, (fullname, vassa_years, monk_type, residence, position,
                      education_level, academic_year, match_id))
            else:
                cursor.execute("""
                    INSERT INTO monk_tbl
                        (fullname, vassa_years, monk_type, residence, position, education_level, academic_year)
                    VALUES (%s, %s, %s, %s, %s, %s, %s)
                """, (fullname, vassa_years, monk_type, residence,
                      position, education_level, academic_year))

            cursor.execute("""
                UPDATE pending_submissions
                SET status = 'approved', reviewed_at = NOW()
                WHERE id = %s
            """, (sub_id,))
            approved += 1

        conn.commit()
        return jsonify({'success': True, 'count': approved})

    except Exception as e:
        if conn:
            conn.rollback()
        return jsonify({'success': False, 'message': str(e)}), 500
    finally:
        if conn:
            conn.close()


@main_bp.route('/api/monks', methods=['POST'])
def add_monk():
    """API endpoint to add a new monk"""
    try:
        data = request.get_json()

        # Extract form data
        fullname = data.get('fullname')
        vassa_years = data.get('total-monk')  # Match form field name
        monk_type = data.get('type')
        residence = data.get('home')
        position = data.get('position')
        education_level = data.get('education_level')
        academic_year = data.get('academic_level')
        living_status = data.get('living_status') or _ACTIVE_LIVING_STATUS

        # Validate required fields
        if not all([fullname, vassa_years, monk_type, residence, position, education_level, academic_year]):
            return jsonify({'success': False, 'message': 'Missing required fields'}), 400
        dropdown_errors = _validate_entry_dropdowns(
            monk_type, residence, position, education_level, academic_year,
        )
        if dropdown_errors:
            return jsonify({'success': False, 'message': ' | '.join(dropdown_errors)}), 400
        if living_status not in _VALID_LIVING_STATUSES:
            return jsonify({'success': False, 'message': 'ស្ថានភាពមិនត្រឹមត្រូវ'}), 400

        # Insert into database
        monk_id = insert_monk(
            fullname, vassa_years, monk_type, residence, position,
            education_level, academic_year, living_status,
        )

        if monk_id:
            return jsonify({'success': True, 'monk_id': monk_id})
        else:
            return jsonify({'success': False, 'message': 'Failed to insert monk'}), 500

    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 500


@main_bp.route('/api/monks', methods=['GET'])
def list_monks():
    """API endpoint to get all monks.
    Pass ?residing=1 to exclude monks who left (ឈប់ស្នាក់នៅ) — still includes នៅស្រុក.
    """
    try:
        residing_only = request.args.get('residing', '').strip() in ('1', 'true', 'yes')
        monks = get_all_monks()
        # Convert to list of dictionaries
        monks_list = []
        for monk in monks:
            living = monk[10] if len(monk) > 10 else _ACTIVE_LIVING_STATUS
            living = living or _ACTIVE_LIVING_STATUS
            if residing_only and living == 'ឈប់ស្នាក់នៅ':
                continue
            monks_list.append({
                'id': monk[0],
                'fullname': monk[1],
                'vassa_years': monk[2],
                'monk_type': monk[3],
                'residence': monk[4],
                'position': monk[5],
                'education_level': monk[6],
                'academic_year': monk[7],
                'created_at': monk[8].isoformat() if monk[8] else None,
                'updated_at': monk[9].isoformat() if monk[9] else None,
                'living_status': living,
            })
        monks_list = sort_attendance_monks(monks_list)
        return jsonify({'success': True, 'monks': monks_list})
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 500


@main_bp.route('/api/monks/<int:monk_id>', methods=['PUT'])
def update_monk_route(monk_id):
    try:
        data = request.get_json()
        fullname = data.get('fullname')
        vassa_years = data.get('total-monk')
        monk_type = data.get('type')
        residence = data.get('home')
        position = data.get('position')
        education_level = data.get('education_level')
        academic_year = data.get('academic_level')
        living_status = data.get('living_status')

        if not all([fullname, vassa_years, monk_type, residence, position, education_level, academic_year]):
            return jsonify({'success': False, 'message': 'Missing required fields'}), 400
        dropdown_errors = _validate_entry_dropdowns(
            monk_type, residence, position, education_level, academic_year,
        )
        if dropdown_errors:
            return jsonify({'success': False, 'message': ' | '.join(dropdown_errors)}), 400
        if living_status is not None and living_status not in _VALID_LIVING_STATUSES:
            return jsonify({'success': False, 'message': 'ស្ថានភាពមិនត្រឹមត្រូវ'}), 400

        success = update_monk(
            monk_id, fullname, vassa_years, monk_type, residence, position,
            education_level, academic_year, living_status,
        )
        if success:
            return jsonify({'success': True})
        return jsonify({'success': False, 'message': 'Failed to update'}), 500
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 500


@main_bp.route('/api/monks/<int:monk_id>/living-status', methods=['PATCH'])
def patch_monk_living_status(monk_id):
    if not user_allowed(_session_user(), '/view'):
        return jsonify({'success': False, 'message': 'Forbidden'}), 403
    data = request.get_json(silent=True) or {}
    living_status = (data.get('living_status') or '').strip()
    if living_status not in _VALID_LIVING_STATUSES:
        return jsonify({'success': False, 'message': 'ស្ថានភាពមិនត្រឹមត្រូវ'}), 400
    try:
        ok = update_monk_living_status(monk_id, living_status)
        if not ok:
            return jsonify({'success': False, 'message': 'រកមិនឃើញព្រះសង្ឃ'}), 404
        return jsonify({'success': True, 'living_status': living_status})
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 500


@main_bp.route('/api/monks/<int:monk_id>', methods=['DELETE'])
def delete_monk_route(monk_id):
    try:
        success = delete_monk(monk_id)
        if success:
            return jsonify({'success': True})
        return jsonify({'success': False, 'message': 'Failed to delete'}), 500
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 500


@main_bp.route('/api/monks/bulk-delete', methods=['POST'])
def bulk_delete_monks():
    data = request.get_json(silent=True) or {}
    ids  = [int(i) for i in data.get('ids', []) if str(i).isdigit()]
    if not ids:
        return jsonify({'success': False, 'message': 'No ids provided'}), 400

    conn = None
    try:
        conn = connect_db()
        conn.autocommit = False
        cursor = conn.cursor()
        cursor.execute("DELETE FROM monk_tbl WHERE id = ANY(%s)", (ids,))
        count = cursor.rowcount
        conn.commit()
        return jsonify({'success': True, 'count': count})
    except Exception as e:
        if conn:
            conn.rollback()
        return jsonify({'success': False, 'message': str(e)}), 500
    finally:
        if conn:
            conn.close()


@main_bp.route('/api/monks/check-duplicate', methods=['GET'])
def check_duplicate():
    fullname    = request.args.get('fullname', '').strip()
    monk_type   = request.args.get('monk_type', '').strip()
    vassa_years = request.args.get('vassa_years', '').strip()
    residence   = request.args.get('residence', '').strip()
    position    = request.args.get('position', '').strip()

    if not all([fullname, monk_type, vassa_years, residence, position]):
        return jsonify({'exists': False})

    try:
        vassa_int = int(vassa_years)
    except (ValueError, TypeError):
        return jsonify({'exists': False})

    try:
        conn   = connect_db()
        cursor = conn.cursor()
        cursor.execute("""
            SELECT id, fullname, monk_type, vassa_years, residence, position
            FROM monk_tbl
            WHERE LOWER(fullname) = LOWER(%s)
              AND monk_type   = %s
              AND vassa_years = %s
              AND residence   = %s
              AND position    = %s
            LIMIT 1
        """, (fullname, monk_type, vassa_int, residence, position))
        row = cursor.fetchone()
        cursor.close()
        conn.close()
        if row:
            return jsonify({
                'exists': True,
                'match': {
                    'id':          row[0],
                    'fullname':    row[1],
                    'monk_type':   row[2],
                    'vassa_years': row[3],
                    'residence':   row[4],
                    'position':    row[5]
                }
            })
        return jsonify({'exists': False})
    except Exception as e:
        return jsonify({'exists': False, 'error': str(e)})


@main_bp.route('/api/seat-order', methods=['GET'])
def get_seat_order():
    try:
        import json as _json
        conn = connect_db()
        cur  = conn.cursor()
        cur.execute("SELECT type, monk_ids, updated_at FROM seat_order")
        rows = cur.fetchall()
        cur.close(); conn.close()
        result = {'bhikkhu': None, 'samanera': None, 'grid_config': None, 'updated_at': None}
        for row in rows:
            result[row[0]] = _json.loads(row[1])
            ts = row[2].isoformat() if row[2] else None
            if ts and (result['updated_at'] is None or ts > result['updated_at']):
                result['updated_at'] = ts
        return jsonify({'success': True, **result})
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 500


@main_bp.route('/api/seat-order', methods=['POST'])
def save_seat_order():
    if not user_allowed(_session_user(), '/layout'):
        abort(403)
    import json as _json
    data  = request.get_json(silent=True) or {}
    type_ = str(data.get('type', '')).strip()
    ids   = data.get('ids', [])
    if type_ not in ('bhikkhu', 'samanera', 'grid_config'):
        return jsonify({'success': False, 'message': 'Invalid type'}), 400
    if type_ == 'grid_config':
        if not isinstance(ids, dict):
            return jsonify({'success': False, 'message': 'ids must be a dict for grid_config'}), 400
    elif not isinstance(ids, list):
        return jsonify({'success': False, 'message': 'ids must be a list'}), 400
    try:
        conn = connect_db()
        cur  = conn.cursor()
        cur.execute("""
            INSERT INTO seat_order (type, monk_ids, updated_at)
            VALUES (%s, %s, NOW())
            ON CONFLICT (type) DO UPDATE
                SET monk_ids = EXCLUDED.monk_ids, updated_at = NOW()
            RETURNING updated_at
        """, (type_, _json.dumps(ids)))
        row = cur.fetchone()
        conn.commit()
        cur.close(); conn.close()
        return jsonify({'success': True, 'updated_at': row[0].isoformat() if row else None})
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 500


@main_bp.route('/view')
def view_monks():
    return render_template('view.html')


@main_bp.route('/layout')
def layout():
    return render_template('layout.html', role=session.get('role', ''))


@main_bp.route('/classroom-layout')
def classroom_layout_page():
    return render_template('classroom_layout.html', username=session.get('username', ''))


@main_bp.route('/api/classroom-layout', methods=['GET'])
def get_classroom_layout():
    import json as _json
    try:
        conn = connect_db()
        cur = conn.cursor()
        cur.execute(
            "SELECT layout_data, updated_at FROM classroom_layout WHERE id = 1"
        )
        row = cur.fetchone()
        cur.close()
        conn.close()
        layout = row[0] if row else {'rows': []}
        if isinstance(layout, str):
            layout = _json.loads(layout)
        updated_at = row[1].isoformat() if row and row[1] else None
        return jsonify({
            'success': True,
            'layout': layout,
            'updated_at': updated_at,
        })
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 500


@main_bp.route('/api/classroom-layout', methods=['POST'])
def save_classroom_layout():
    if not user_allowed(_session_user(), '/classroom-layout'):
        abort(403)
    import json as _json
    data = request.get_json(silent=True) or {}
    layout = data.get('layout')
    if not isinstance(layout, dict) or 'rows' not in layout:
        return jsonify({'success': False, 'message': 'layout.rows is required'}), 400
    try:
        conn = connect_db()
        cur = conn.cursor()
        cur.execute("""
            INSERT INTO classroom_layout (id, layout_data, updated_at)
            VALUES (1, %s::jsonb, NOW())
            ON CONFLICT (id) DO UPDATE
                SET layout_data = EXCLUDED.layout_data,
                    updated_at = NOW()
            RETURNING updated_at
        """, (_json.dumps(layout),))
        row = cur.fetchone()
        conn.commit()
        cur.close()
        conn.close()
        return jsonify({
            'success': True,
            'updated_at': row[0].isoformat() if row else None,
        })
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 500


@main_bp.route('/api/check', methods=['GET'])
def check_system():
    """Check DB connection, record count, and existing triggers on monk_tbl"""
    try:
        conn = connect_db()
        cursor = conn.cursor()

        cursor.execute("SELECT COUNT(*) FROM monk_tbl;")
        row = cursor.fetchone()
        total = row[0] if row else 0

        cursor.execute("""
            SELECT created_at, updated_at FROM monk_tbl
            ORDER BY updated_at DESC NULLS LAST LIMIT 1;
        """)
        latest = cursor.fetchone()
        latest_created = latest[0].isoformat() if latest and latest[0] else None
        latest_updated = latest[1].isoformat() if latest and latest[1] else None

        cursor.execute("""
            SELECT trigger_name, event_manipulation, action_timing
            FROM information_schema.triggers
            WHERE event_object_table = 'monk_tbl'
            ORDER BY trigger_name, event_manipulation;
        """)
        triggers = [{'name': t[0], 'event': t[1], 'timing': t[2]} for t in cursor.fetchall()]

        cursor.close()
        conn.close()

        return jsonify({
            'success': True,
            'db_connected': True,
            'total_records': total,
            'latest_created_at': latest_created,
            'latest_updated_at': latest_updated,
            'triggers': triggers
        })
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 500


@main_bp.route('/api/init-db', methods=['POST'])
def init_database():
    """API endpoint to initialize the database (create tables)"""
    try:
        create_monks_table()
        return jsonify({'success': True, 'message': 'Database initialized successfully'})
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 500


@main_bp.route('/api/setup-trigger', methods=['POST'])
def setup_trigger():
    """Create or replace the updated_at trigger on monk_tbl"""
    try:
        conn = connect_db()
        cursor = conn.cursor()
        cursor.execute("""
            CREATE OR REPLACE FUNCTION set_updated_at()
            RETURNS TRIGGER AS $$
            BEGIN
                NEW.updated_at = NOW();
                RETURN NEW;
            END;
            $$ LANGUAGE plpgsql;
        """)
        cursor.execute("DROP TRIGGER IF EXISTS trg_set_updated_at ON monk_tbl;")
        cursor.execute("""
            CREATE TRIGGER trg_set_updated_at
                BEFORE UPDATE ON monk_tbl
                FOR EACH ROW
                EXECUTE FUNCTION set_updated_at();
        """)
        conn.commit()
        cursor.close()
        conn.close()
        return jsonify({'success': True})
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 500


@main_bp.route('/api/attendance', methods=['GET'])
def get_attendance():
    try:
        date_str = request.args.get('date', _date.today().isoformat())
        conn = connect_db()
        cursor = conn.cursor()
        cursor.execute("SELECT monk_id, status FROM attendance_tbl WHERE date = %s;", (date_str,))
        records = [{'monk_id': r[0], 'status': r[1]} for r in cursor.fetchall()]
        
        cursor.execute("""
            SELECT monk_id, start_date, end_date, reason, shift FROM monk_permission
            WHERE %s BETWEEN start_date AND end_date
        """, (date_str,))

        perms = {}
        target_date = _date.fromisoformat(date_str)
        for r in cursor.fetchall():
            monk_id, start_date, end_date, reason, shift = r
            days_left = (end_date - target_date).days
            if days_left >= 0:
                perms[monk_id] = {
                    'start_date': start_date.isoformat(),
                    'end_date': end_date.isoformat(),
                    'days_left': days_left,
                    'reason': reason or '',
                    'shift': shift or '',
                    'same_day': start_date == end_date,
                }
                
        cursor.close(); conn.close()
        return jsonify({'success': True, 'records': records, 'permissions_info': perms})
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 500


@main_bp.route('/api/attendance', methods=['POST'])
def set_attendance():
    try:
        data = request.get_json()
        monk_id = data.get('monk_id')
        status  = data.get('status')
        date_str = data.get('date', _date.today().isoformat())
        if status not in ('absent', 'permission'):
            return jsonify({'success': False, 'message': 'Invalid status'}), 400
        conn = connect_db()
        cursor = conn.cursor()
        cursor.execute("""
            INSERT INTO attendance_tbl (monk_id, status, date)
            VALUES (%s, %s, %s)
            ON CONFLICT (monk_id, date) DO UPDATE SET status = EXCLUDED.status;
        """, (monk_id, status, date_str))
        conn.commit()

        if status == 'absent':
            block_start, block_end = _get_block_dates(date_str)
            cursor.execute("""
                SELECT 
                    SUM(CASE WHEN status = 'absent' THEN 1 ELSE 0 END),
                    SUM(CASE WHEN status = 'permission' THEN 1 ELSE 0 END)
                FROM attendance_tbl
                WHERE monk_id = %s
                  AND date >= %s
                  AND date <= %s;
            """, (monk_id, block_start.isoformat(), block_end.isoformat()))
            row = cursor.fetchone()
            absent_count = int(row[0] or 0)
            perm_count = int(row[1] or 0)

            if absent_count in (3, 6):
                import threading
                def send_tg():
                    try:
                        _send_absent_alert_telegram(
                            monk_id, date_str, absent_count, perm_count, 'absent_alert',
                        )
                    except Exception:
                        pass
                threading.Thread(target=send_tg, daemon=True).start()
                    
        cursor.close(); conn.close()
        return jsonify({'success': True})
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 500


@main_bp.route('/api/permissions', methods=['POST'])
def add_permission():
    try:
        data = request.get_json()
        monk_id = data.get('monk_id')
        start_date = data.get('start_date')
        end_date = data.get('end_date')
        reason = data.get('reason', '')
        shift = str(data.get('shift') or '').strip()

        if not monk_id or not start_date or not end_date:
            return jsonify({'success': False, 'message': 'Missing required fields'}), 400

        s_date = _date.fromisoformat(start_date)
        e_date = _date.fromisoformat(end_date)

        if e_date < s_date:
            return jsonify({'success': False, 'message': 'ថ្ងៃបញ្ចប់ត្រូវតែក្រោយថ្ងៃចាប់ផ្តើម'}), 400

        same_day = s_date == e_date
        if same_day:
            if shift not in ('ព្រឹក', 'ល្ងាច'):
                shift = 'ល្ងាច'
        else:
            shift = None

        conn = connect_db()
        cursor = conn.cursor()

        # Sync SERIAL if it fell behind (fixes duplicate key on monk_permission_pkey)
        cursor.execute("""
            SELECT setval(
                pg_get_serial_sequence('monk_permission', 'id'),
                COALESCE((SELECT MAX(id) FROM monk_permission), 1),
                true
            );
        """)

        # Replace any existing leave range for this monk (one active leave record)
        cursor.execute("DELETE FROM monk_permission WHERE monk_id = %s", (monk_id,))
        cursor.execute("""
            INSERT INTO monk_permission (monk_id, reason, start_date, end_date, shift)
            VALUES (%s, %s, %s, %s, %s)
        """, (monk_id, reason, start_date, end_date, shift))

        from datetime import timedelta
        current_date = s_date
        while current_date <= e_date:
            cursor.execute("""
                INSERT INTO attendance_tbl (monk_id, status, date)
                VALUES (%s, 'permission', %s)
                ON CONFLICT (monk_id, date) DO UPDATE SET status = EXCLUDED.status;
            """, (monk_id, current_date.isoformat()))
            current_date += timedelta(days=1)

        conn.commit()
        cursor.close()
        conn.close()
        return jsonify({'success': True})
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 500


@main_bp.route('/api/attendance/monk/<int:monk_id>', methods=['GET'])
def get_monk_attendance(monk_id):
    try:
        start = request.args.get('start', '')
        end   = request.args.get('end',   '')
        conn   = connect_db()
        cursor = conn.cursor()
        where_parts = ["monk_id = %s"]
        params      = [monk_id]
        if start: where_parts.append("date >= %s"); params.append(start)
        if end:   where_parts.append("date <= %s"); params.append(end)
        cursor.execute(f"""
            SELECT date, status FROM attendance_tbl
            WHERE {' AND '.join(where_parts)}
            ORDER BY date DESC
        """, params)
        records = [{'date': str(r[0]), 'status': r[1]} for r in cursor.fetchall()]
        cursor.close(); conn.close()
        return jsonify({'success': True, 'records': records})
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 500


@main_bp.route('/api/attendance/<int:monk_id>', methods=['DELETE'])
def remove_attendance(monk_id):
    try:
        date_str = request.args.get('date', _date.today().isoformat())
        conn = connect_db()
        cursor = conn.cursor()
        cursor.execute("DELETE FROM attendance_tbl WHERE monk_id = %s AND date = %s;", (monk_id, date_str))
        conn.commit(); cursor.close(); conn.close()
        return jsonify({'success': True})
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 500


@main_bp.route('/report')
def report():
    return render_template('report.html')


def _get_block_dates(date_str):
    """Return the fixed 15-day block containing the given date.
    Days  1-15 → 1st–15th of the month.
    Days 16-31 → 16th–last day of the month."""
    import calendar
    d = _date.fromisoformat(date_str)
    if d.day <= 15:
        start = _date(d.year, d.month, 1)
        end   = _date(d.year, d.month, 15)
    else:
        start = _date(d.year, d.month, 16)
        end   = _date(d.year, d.month, calendar.monthrange(d.year, d.month)[1])
    return start, end


def _format_date_list(dates):
    """Format a list of dates in the same month as '03, 04, 05 - 07 - 2026',
    grouping by month/year when the list spans more than one month."""
    from itertools import groupby
    dates = [d for d in (dates or []) if d]
    if not dates:
        return ''
    groups = []
    for (year, month), grp in groupby(dates, key=lambda d: (d.year, d.month)):
        days = ', '.join(f'{d.day:02d}' for d in grp)
        groups.append(f'{days} - {month:02d} - {year}')
    return '; '.join(groups)


def _fetch_report_rows(args):
    date_str   = (args.get('date') or _date.today().isoformat())
    start_date, end_date = _get_block_dates(date_str)

    monk_type  = (args.get('monk_type')       or '').strip()
    kuti       = (args.get('kuti')            or '').strip()
    edu        = (args.get('education_level') or '').strip()
    acad       = (args.get('academic_year')   or '').strip()
    name_pfx   = (args.get('name')            or '').strip()

    where_parts, where_params = [], []
    if monk_type: where_parts.append("m.monk_type = %s");       where_params.append(monk_type)
    if kuti:      where_parts.append("m.residence = %s");        where_params.append(kuti)
    if edu:       where_parts.append("m.education_level = %s");  where_params.append(edu)
    if acad:      where_parts.append("m.academic_year = %s");    where_params.append(acad)
    if name_pfx:  where_parts.append("m.fullname ILIKE %s");     where_params.append(name_pfx + '%')

    where_sql = ('WHERE ' + ' AND '.join(where_parts)) if where_parts else ''

    sql = f"""
        SELECT m.id, m.fullname, m.monk_type, m.position, m.vassa_years,
               m.residence, m.education_level, m.academic_year,
               COUNT(CASE WHEN a.status = 'absent'     THEN 1 END) AS absent_count,
               COUNT(CASE WHEN a.status = 'permission' THEN 1 END) AS perm_count,
               ARRAY_AGG(a.date ORDER BY a.date) FILTER (WHERE a.status = 'absent')     AS absent_dates,
               ARRAY_AGG(a.date ORDER BY a.date) FILTER (WHERE a.status = 'permission') AS perm_dates
        FROM monk_tbl m
        LEFT JOIN attendance_tbl a
            ON a.monk_id = m.id AND a.date >= %s AND a.date <= %s
        {where_sql}
        GROUP BY m.id, m.fullname, m.monk_type, m.position, m.vassa_years,
                 m.residence, m.education_level, m.academic_year
        ORDER BY m.monk_type,
                 COUNT(CASE WHEN a.status = 'absent'     THEN 1 END) DESC,
                 COUNT(CASE WHEN a.status = 'permission' THEN 1 END) DESC,
                 m.fullname;
    """
    params = [start_date.isoformat(), end_date.isoformat()] + where_params

    conn = connect_db()
    cur  = conn.cursor()
    cur.execute(sql, params)
    rows = cur.fetchall()
    cur.close(); conn.close()

    monks = [{
        'id': r[0], 'fullname': r[1], 'monk_type': r[2], 'position': r[3],
        'vassa_years': r[4], 'residence': (r[5] or '').replace('_', ' '),
        'education_level': r[6] or '', 'academic_year': r[7] or '',
        'absent_count': int(r[8] or 0), 'permission_count': int(r[9] or 0),
        'absent_dates': _format_date_list(r[10]), 'perm_dates': _format_date_list(r[11]),
    } for r in rows]

    return monks, start_date, end_date


@main_bp.route('/api/attendance/report', methods=['GET'])
def attendance_report():
    try:
        monks, start_date, end_date = _fetch_report_rows(request.args)
        return jsonify({
            'success':    True,
            'start_date': start_date.isoformat(),
            'end_date':   end_date.isoformat(),
            'monks':      monks
        })
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 500


def _build_report_html(monks, start_date, end_date, filters_applied, ABSENT_LIMIT, PERM_LIMIT):
    import html as _html

    def row_html(m, idx):
        ab = m['absent_count']     >= ABSENT_LIMIT
        pr = m['permission_count'] >= PERM_LIMIT
        bg = '#fff5f5' if ab else ('#fffaf0' if pr else '#ffffff')
        badge = (
            '<span class="badge-danger">⚠ លើសអវត្តមាន</span>' if ab else
            '<span class="badge-warning">⚠ លើសច្បាប់</span>' if pr else
            '<span class="badge-ok">✓ ប្រក្រតី</span>'
        )
        edu = _html.escape(f"{m['education_level']} {m['academic_year']}".strip())
        ac  = 'color:#c53030;font-weight:bold' if ab else 'color:#718096'
        pc  = 'color:#c05621;font-weight:bold' if pr else 'color:#718096'
        d_parts = []
        if m['absent_dates']: d_parts.append(f'<span style="color:#c53030">❌ {_html.escape(m["absent_dates"])}</span>')
        if m['perm_dates']:   d_parts.append(f'<span style="color:#c05621">📋 {_html.escape(m["perm_dates"])}</span>')
        dates_cell = '<br>'.join(d_parts) if d_parts else '—'
        return (
            f'<tr style="background:{bg}">'
            f'<td class="num">{idx}</td>'
            f'<td><strong>{_html.escape(m["fullname"])}</strong></td>'
            f'<td>{_html.escape(m["position"])}</td>'
            f'<td class="num">{m["vassa_years"]} ឆ្នាំ</td>'
            f'<td>{_html.escape(m["residence"])}</td>'
            f'<td>{edu}</td>'
            f'<td class="num" style="{ac}">{m["absent_count"] or "—"}</td>'
            f'<td class="num" style="{pc}">{m["permission_count"] or "—"}</td>'
            f'<td class="dates">{dates_cell}</td>'
            f'<td>{badge}</td>'
            f'</tr>'
        )

    def section_html(section_monks, title, hc, bc):
        if not section_monks:
            return ''
        rows = ''.join(row_html(m, i + 1) for i, m in enumerate(section_monks))
        return (
            f'<h2 style="color:{hc};border-bottom:2px solid {bc};'
            f'padding:6px 0;margin:18px 0 6px;font-size:13px;">'
            f'{_html.escape(title)} ({len(section_monks)} នាក់)</h2>'
            f'<table><thead><tr>'
            f'<th>#</th><th>ឈ្មោះ</th><th>តួនាទី</th><th>វស្សា</th>'
            f'<th>ស្នាក់នៅ</th><th>ការសិក្សា</th><th>❌</th><th>📋</th><th>ថ្ងៃ</th><th>ស្ថានភាព</th>'
            f'</tr></thead><tbody>{rows}</tbody></table>'
        )

    bhikkhus    = [m for m in monks if m['monk_type'] == 'ភិក្ខុ']
    samaneras   = [m for m in monks if m['monk_type'] == 'សាមណេរ']
    absent_viol = sum(1 for m in monks if m['absent_count']     >= ABSENT_LIMIT)
    perm_viol   = sum(1 for m in monks if m['permission_count'] >= PERM_LIMIT)
    clean       = len(monks) - absent_viol - perm_viol
    filter_line = (
        f'<p class="sub">តម្រង: {" | ".join(filters_applied)}</p>'
        if filters_applied else ''
    )
    date_range  = f"{start_date.strftime('%d/%m/%Y')} ដល់ {end_date.strftime('%d/%m/%Y')}"

    return (
        '<!DOCTYPE html><html lang="km"><head><meta charset="UTF-8"><style>'
        "@import url('https://fonts.googleapis.com/css2?family=Battambang:wght@400;700&family=Moul&display=swap');"
        '*{box-sizing:border-box;margin:0;padding:0}'
        "body{font-family:'Battambang','Khmer MN','Khmer Sangam MN',sans-serif;"
        'color:#2d3748;font-size:10px;}'
        "h1{text-align:center;font-family:'Moul','Battambang',serif;font-size:16px;font-weight:400;color:#1a202c;margin-bottom:4px;}"
        ".sub{text-align:center;font-family:'Battambang',serif;color:#718096;font-size:14px;margin-bottom:3px;}"
        '.summary{display:flex;gap:12px;justify-content:center;margin:10px 0;flex-wrap:wrap;}'
        '.summary span{padding:3px 10px;border-radius:4px;font-weight:bold;font-size:9.5px;}'
        '.s-total{background:#edf2f7;color:#2d3748;}'
        '.s-absent{background:#fed7d7;color:#c53030;}'
        '.s-perm{background:#feebc8;color:#c05621;}'
        '.s-clean{background:#c6f6d5;color:#276749;}'
        'table{width:100%;border-collapse:collapse;margin-bottom:6px;}'
        'thead tr{background:#f7fafc;}'
        'th{padding:6px 7px;text-align:left;font-size:8.5px;font-weight:700;'
        'color:#718096;border-bottom:2px solid #e2e8f0;white-space:nowrap;}'
        'td{padding:6px 7px;border-bottom:1px solid #edf2f7;vertical-align:middle;font-size:9.5px;}'
        '.num{text-align:center;}'
        '.badge-danger{background:#fed7d7;color:#c53030;padding:2px 7px;'
        'border-radius:10px;font-size:8px;font-weight:bold;white-space:nowrap;}'
        '.badge-warning{background:#feebc8;color:#c05621;padding:2px 7px;'
        'border-radius:10px;font-size:8px;font-weight:bold;white-space:nowrap;}'
        '.badge-ok{background:#c6f6d5;color:#276749;padding:2px 7px;'
        'border-radius:10px;font-size:8px;font-weight:bold;white-space:nowrap;}'
        '.dates{font-size:8.5px;line-height:1.7;}'
        '@page{size:A4 portrait;margin:12mm 10mm;}'
        '@media print { @page { size: A4 portrait; margin: 12mm; } }'
        '</style></head><body>'
        '<h1>វត្តនិរោធរង្សី — របាយការណ៍វត្តមាន</h1>'
        f'<p class="sub">ចន្លោះ: {date_range} (១៥ ថ្ងៃ)</p>'
        f'{filter_line}'
        '<div class="summary">'
        f'<span class="s-total">ព្រះសង្ឃ: {len(monks)} នាក់</span>'
        f'<span class="s-absent">❌ លើសអវត្តមាន: {absent_viol} នាក់</span>'
        f'<span class="s-perm">📋 លើសច្បាប់: {perm_viol} នាក់</span>'
        f'<span class="s-clean">✓ ប្រក្រតី: {clean} នាក់</span>'
        '</div>'
        + section_html(bhikkhus,  'ផ្នែកទី ១ — ភិក្ខុ',   '#8a6100', '#f0c040')
        + section_html(samaneras, 'ផ្នែកទី ២ — សាមណេរ', '#1b5e20', '#66bb6a')
        + '</body></html>'
    )


@main_bp.route('/api/attendance/export-report-pdf', methods=['GET'])
def export_attendance_report_pdf():
    try:
        import io, requests as req
        
        action = request.args.get('action', 'download')
        monks, start_date, end_date = _fetch_report_rows(request.args)

        ABSENT_LIMIT = 2
        PERM_LIMIT   = 3

        filters_applied = []
        for k, label in [('monk_type','ប្រភេទ'),('kuti','កុដិ'),
                          ('education_level','ការសិក្សា'),('academic_year','ឆ្នាំ'),('name','ឈ្មោះ')]:
            v = request.args.get(k, '').strip()
            if v:
                filters_applied.append(f"{label}: {v}")

        html_str  = _build_report_html(monks, start_date, end_date, filters_applied, ABSENT_LIMIT, PERM_LIMIT)
        html_str += '<script>window.onload = function() { setTimeout(function(){ window.print(); }, 500); }</script>'
        return html_str

    except ImportError:
        return jsonify({'success': False,
                        'message': 'WeasyPrint មិនទាន់ install — សូម run: pip install weasyprint'}), 500
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 500


@main_bp.route('/api/attendance/export-report', methods=['GET'])
def export_attendance_report():
    try:
        import io, requests as req
        from docx import Document
        from docx.shared import Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_ALIGN_VERTICAL
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        action = request.args.get('action', 'download')
        monks, start_date, end_date = _fetch_report_rows(request.args)

        ABSENT_LIMIT = 2
        PERM_LIMIT   = 3

        def shade(cell, hex_color):
            tc  = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'),   'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'),  hex_color)
            tcPr.append(shd)

        def add_report_table(doc, section_monks):
            headers = ['#', 'ឈ្មោះ', 'តួនាទី', 'វស្សា', 'ស្នាក់នៅ', 'ការសិក្សា', '❌', '📋', 'ថ្ងៃ', 'ស្ថានភាព']
            widths  = [0.5, 3.0, 2.5, 1.0, 2.0, 1.6, 0.7, 0.7, 2.8, 1.8]

            tbl = doc.add_table(rows=1, cols=len(headers))
            tbl.style = 'Table Grid'
            for i, (h, w) in enumerate(zip(headers, widths)):
                cell = tbl.rows[0].cells[i]
                cell.width = Cm(w)
                shade(cell, 'F7FAFC')
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(h)
                run.bold = True
                run.font.size = Pt(8.5)
                run.font.color.rgb = RGBColor(0x71, 0x80, 0x96)

            for i, m in enumerate(section_monks, 1):
                ab_viol = m['absent_count']     >= ABSENT_LIMIT
                pr_viol = m['permission_count'] >= PERM_LIMIT
                row_color = 'FFF5F5' if ab_viol else ('FFFAF0' if pr_viol else 'FFFFFF')
                status = '⚠ លើសអវត្តមាន' if ab_viol else ('⚠ លើសច្បាប់' if pr_viol else '✓ ប្រក្រតី')

                row = tbl.add_row()
                dates_parts = []
                if m['absent_dates']:  dates_parts.append(f"❌ {m['absent_dates']}")
                if m['perm_dates']:    dates_parts.append(f"📋 {m['perm_dates']}")
                dates_text = '\n'.join(dates_parts) if dates_parts else '—'

                vals = [
                    str(i), m['fullname'], m['position'], f"{m['vassa_years']} ឆ្នាំ",
                    m['residence'], f"{m['education_level']} {m['academic_year']}".strip(),
                    str(m['absent_count']) if m['absent_count'] else '—',
                    str(m['permission_count']) if m['permission_count'] else '—',
                    dates_text,
                    status
                ]
                for j, (val, w) in enumerate(zip(vals, widths)):
                    cell = row.cells[j]
                    cell.width = Cm(w)
                    shade(cell, row_color)
                    p = cell.paragraphs[0]
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if j in (0,3,6,7) else WD_ALIGN_PARAGRAPH.LEFT
                    run = p.add_run(val)
                    run.font.size = Pt(8.5)
                    if j == 1: run.bold = True
                    if ab_viol and j == 6: run.font.color.rgb = RGBColor(0xC5, 0x30, 0x30)
                    if pr_viol and j == 7: run.font.color.rgb = RGBColor(0xC0, 0x56, 0x21)

        doc = Document()
        sec = doc.sections[0]
        sec.left_margin = sec.right_margin = Cm(1.5)
        sec.top_margin  = sec.bottom_margin = Cm(1.5)

        t = doc.add_heading('វត្តនិរោធរង្សី — របាយការណ៍វត្តមាន', 0)
        t.alignment = WD_ALIGN_PARAGRAPH.CENTER

        sub = doc.add_paragraph(
            f"ចន្លោះ: {start_date.strftime('%d/%m/%Y')} ដល់ {end_date.strftime('%d/%m/%Y')} (១៥ ថ្ងៃ)"
        )
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub.runs[0].font.size = Pt(10)
        sub.runs[0].font.color.rgb = RGBColor(0x71, 0x80, 0x96)

        # Applied filters note
        filters_applied = []
        for k, label in [('monk_type','ប្រភេទ'),('kuti','កុដិ'),('education_level','ការសិក្សា'),
                          ('academic_year','ឆ្នាំ'),('name','ឈ្មោះ')]:
            v = request.args.get(k, '').strip()
            if v: filters_applied.append(f"{label}: {v}")
        if filters_applied:
            fp = doc.add_paragraph('តម្រង: ' + ' | '.join(filters_applied))
            fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            fp.runs[0].font.size = Pt(9)
            fp.runs[0].font.color.rgb = RGBColor(0x4A, 0x55, 0x68)

        doc.add_paragraph()

        bhikkhus  = [m for m in monks if m['monk_type'] == 'ភិក្ខុ']
        samaneras = [m for m in monks if m['monk_type'] == 'សាមណេរ']

        if bhikkhus:
            h1 = doc.add_heading('ផ្នែកទី ១ — ភិក្ខុ', 1)
            h1.runs[0].font.color.rgb = RGBColor(0x8A, 0x61, 0x00)
            add_report_table(doc, bhikkhus)
            doc.add_paragraph()

        if samaneras:
            h2 = doc.add_heading('ផ្នែកទី ២ — សាមណេរ', 1)
            h2.runs[0].font.color.rgb = RGBColor(0x1B, 0x5E, 0x20)
            add_report_table(doc, samaneras)
            doc.add_paragraph()

        absent_viol = sum(1 for m in monks if m['absent_count']     >= ABSENT_LIMIT)
        perm_viol   = sum(1 for m in monks if m['permission_count'] >= PERM_LIMIT)
        sp = doc.add_paragraph(
            f"📊 សរុប {len(monks)} នាក់  |  "
            f"❌ លើសអវត្តមាន: {absent_viol} នាក់  |  "
            f"📋 លើសច្បាប់: {perm_viol} នាក់"
        )
        sp.runs[0].font.size = Pt(10)
        sp.runs[0].bold = True

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)

        if action in ('telegram', 'telegram-both'):
            TELEGRAM_TOKEN   = '8950898077:AAHNR0tTgtJWy17wMXooKwg4nfQLGdfe5aw'
            TELEGRAM_CHAT_ID = -1003960014484
            fname   = f"attendance_report_{end_date.isoformat()}.docx"
            caption = (
                f"📋 របាយការណ៍វត្តមាន — {end_date.strftime('%d/%m/%Y')}\n"
                f"📊 សរុប {len(monks)} នាក់  |  ❌ {absent_viol}  |  📋 {perm_viol}"
            )
            tg = req.post(
                f'https://api.telegram.org/bot{TELEGRAM_TOKEN}/sendDocument',
                data={'chat_id': TELEGRAM_CHAT_ID, 'caption': caption},
                files={'document': (fname, buf,
                       'application/vnd.openxmlformats-officedocument.wordprocessingml.document')},
                timeout=15
            ).json()
            if not tg.get('ok'):
                return jsonify({'success': False, 'message': f"Telegram (Word): {tg.get('description')}"}), 500

            

        fname = f"attendance_report_{end_date.isoformat()}.docx"
        return send_file(
            buf,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=fname
        )

    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 500


@main_bp.route('/api/attendance/history/<int:monk_id>', methods=['GET'])
def attendance_history(monk_id):
    try:
        date_str = request.args.get('date', _date.today().isoformat())
        block_start, block_end = _get_block_dates(date_str)
        conn = connect_db()
        cursor = conn.cursor()
        cursor.execute("""
            SELECT
                SUM(CASE WHEN status = 'absent'     THEN 1 ELSE 0 END),
                SUM(CASE WHEN status = 'permission' THEN 1 ELSE 0 END)
            FROM attendance_tbl
            WHERE monk_id = %s
              AND date >= %s
              AND date <= %s;
        """, (monk_id, block_start.isoformat(), block_end.isoformat()))
        row = cursor.fetchone()
        cursor.close(); conn.close()
        return jsonify({
            'success': True,
            'absent_count':     int(row[0] or 0),
            'permission_count': int(row[1] or 0)
        })
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 500


@main_bp.route('/api/attendance/full-history/<int:monk_id>', methods=['GET'])
def attendance_full_history(monk_id):
    try:
        date_str = request.args.get('date', _date.today().isoformat())
        block_start, block_end = _get_block_dates(date_str)
        conn = connect_db()
        cursor = conn.cursor()
        cursor.execute("""
            SELECT a.date, a.status, p.reason, p.end_date
            FROM attendance_tbl a
            LEFT JOIN monk_permission p
                   ON p.monk_id = a.monk_id
                  AND a.date >= p.start_date
                  AND a.date <= p.end_date
            WHERE a.monk_id = %s
              AND a.date >= %s
              AND a.date <= %s
            ORDER BY a.date DESC;
        """, (monk_id, block_start.isoformat(), block_end.isoformat()))
        rows = cursor.fetchall()
        cursor.close(); conn.close()

        absent_dates = []
        perm_dates   = []
        for date, status, reason, end_date in rows:
            entry = {'date': date.isoformat() if hasattr(date, 'isoformat') else str(date)}
            if status == 'absent':
                absent_dates.append(entry)
            elif status == 'permission':
                entry['reason']   = reason   or ''
                entry['end_date'] = end_date.isoformat() if end_date and hasattr(end_date, 'isoformat') else (str(end_date) if end_date else '')
                perm_dates.append(entry)

        return jsonify({
            'success':          True,
            'absent_count':     len(absent_dates),
            'permission_count': len(perm_dates),
            'absent_dates':     absent_dates,
            'perm_dates':       perm_dates,
        })
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 500


def _tg_send_message(token, chat_id, text, req_lib, timeout=10):
    """Send a text message to Telegram, splitting into chunks if > 4000 chars."""
    MAX = 4000
    if len(text) <= MAX:
        return req_lib.post(
            f'https://api.telegram.org/bot{token}/sendMessage',
            json={'chat_id': chat_id, 'text': text},
            timeout=timeout,
        ).json()

    # Split on double-newline paragraph boundaries to keep records intact
    paragraphs = text.split('\n\n')
    chunks, current = [], ''
    for para in paragraphs:
        candidate = (current + '\n\n' + para).lstrip('\n')
        if len(candidate) <= MAX:
            current = candidate
        else:
            if current:
                chunks.append(current)
            current = para[:MAX]          # hard-cut only if single para > MAX
    if current:
        chunks.append(current)

    last = None
    for chunk in chunks:
        last = req_lib.post(
            f'https://api.telegram.org/bot{token}/sendMessage',
            json={'chat_id': chat_id, 'text': chunk},
            timeout=timeout,
        ).json()
        if not last.get('ok'):
            return last
    return last


@main_bp.route('/api/attendance/submit', methods=['POST'])
def submit_attendance():
    TELEGRAM_TOKEN   = '8950898077:AAHNR0tTgtJWy17wMXooKwg4nfQLGdfe5aw'
    TELEGRAM_CHAT_ID = -1003960014484  # Channel: គ្រប់គ្រង អវត្តមាន-ច្បាប់ ថ្វាយបង្គំប្រចាំថ្ងៃ

    try:
        import requests as req
        data = request.get_json(silent=True) or {}
        date_str = (data.get('date') or '').strip() or _date.today().isoformat()
        try:
            report_day = _date.fromisoformat(date_str)
        except ValueError:
            return jsonify({'success': False, 'message': 'កាលបរិច្ឆេទមិនត្រឹមត្រូវ'}), 400

        conn = connect_db()
        cursor = conn.cursor()
        cursor.execute("""
            SELECT m.fullname, m.monk_type, m.position, m.vassa_years, m.residence,
                   a.status, p.reason
            FROM attendance_tbl a
            JOIN monk_tbl m ON m.id = a.monk_id
            LEFT JOIN monk_permission p
                   ON p.monk_id = a.monk_id
                  AND a.date BETWEEN p.start_date AND p.end_date
            WHERE a.date = %s
            ORDER BY m.monk_type, a.status, m.fullname;
        """, (date_str,))
        rows = cursor.fetchall()
        cursor.close(); conn.close()

        if not rows:
            return jsonify({'success': False, 'message': 'មិនមានការចុះឈ្មោះត្រូវបញ្ជូនទេ'}), 400

        absent_count     = sum(1 for r in rows if r[5] == 'absent')
        permission_count = sum(1 for r in rows if r[5] == 'permission')

        def fmt_group(monks, show_position=True):
            lines = []
            for i, (name, _, position, vassa, kuti, status, reason) in enumerate(monks, 1):
                icon  = '❌' if status == 'absent' else '📋'
                label = 'អវត្តមាន' if status == 'absent' else 'ច្បាប់'
                kuti_display = (kuti or '').replace('_', ' ')
                block = [
                    f'{i}. {icon} {name}',
                ]
                if show_position:
                    block.append(f'   ▸ តួនាទី: {position}')
                block += [
                    f'   ▸ វស្សា: {vassa} ឆ្នាំ',
                    f'   ▸ កុដិ: {kuti_display}',
                    f'   ▸ ស្ថានភាព: {label}',
                ]
                if status == 'permission' and (reason or '').strip():
                    block.append(f'   ▸ មូលហេតុ: {reason.strip()}')
                lines.append('\n'.join(block))
            return '\n\n'.join(lines)

        bhikkhus  = [r for r in rows if r[1] == 'ភិក្ខុ']
        samaneras = [r for r in rows if r[1] == 'សាមណេរ']

        d_fmt = report_day.strftime('%d/%m/%Y')
        parts = [
            f'🏛 វត្តនិរោធរង្សី',
            f'📋 ព័ត៌មានថ្វាយបង្គំប្រចាំថ្ងៃ — {d_fmt}',
            '═' * 15,
        ]
        if bhikkhus:
            parts += ['\n📿 ភិក្ខុ', '─' * 15, fmt_group(bhikkhus, show_position=False)]
        if samaneras:
            parts += ['\n🔰 សាមណេរ', '─' * 15, fmt_group(samaneras, show_position=True)]
        parts += [
            '\n' + '═' * 15,
            f'📊 សរុបចំនួន : {len(rows)} នាក់',
            f'   ❌ អវត្តមាន : {absent_count} នាក់',
            f'   📋 ច្បាប់    : {permission_count} នាក់',
        ]
        message = '\n'.join(parts)

        tg = _tg_send_message(TELEGRAM_TOKEN, TELEGRAM_CHAT_ID, message, req)

        if not tg.get('ok'):
            return jsonify({'success': False, 'message': f"Telegram: {tg.get('description', 'error')}"}), 500

        conn2 = connect_db()
        cur2 = conn2.cursor()
        cur2.execute("""
            SELECT a.monk_id, m.fullname, a.status
            FROM attendance_tbl a
            JOIN monk_tbl m ON m.id = a.monk_id
            WHERE a.date = %s AND a.status IN ('absent', 'permission')
        """, (date_str,))
        for mid, fname, status in cur2.fetchall():
            _log_telegram_notify(
                mid, fname, 'daily_submit', date_str,
                detail=f'status={status}',
            )
        cur2.close()
        conn2.close()

        _log_act('attendance_submit', 'layout', f'{len(rows)} monks — {date_str}')
        return jsonify({'success': True, 'total': len(rows)})

    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 500


@main_bp.route('/api/attendance/submit-image', methods=['POST'])
def submit_attendance_image():
    TELEGRAM_TOKEN   = '8950898077:AAHNR0tTgtJWy17wMXooKwg4nfQLGdfe5aw'
    TELEGRAM_CHAT_ID = -1003960014484

    try:
        import requests as req
        if 'image' not in request.files:
            return jsonify({'success': False, 'message': 'No image provided'}), 400

        image_bytes = request.files['image'].read()
        date_str = (request.form.get('date') or '').strip() or _date.today().isoformat()
        try:
            d_fmt = _date.fromisoformat(date_str).strftime('%d/%m/%Y')
        except ValueError:
            d_fmt = _date.today().strftime('%d/%m/%Y')
        caption = f'📋 បញ្ជីអវត្តមាន/ច្បាប់ — {d_fmt}'

        resp = req.post(
            f'https://api.telegram.org/bot{TELEGRAM_TOKEN}/sendPhoto',
            data={'chat_id': TELEGRAM_CHAT_ID, 'caption': caption},
            files={'photo': ('attendance.png', image_bytes, 'image/png')},
            timeout=30,
        ).json()

        if not resp.get('ok'):
            return jsonify({'success': False, 'message': resp.get('description', 'Telegram error')}), 500

        return jsonify({'success': True})

    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 500


@main_bp.route('/api/monks/export', methods=['GET'])
def export_monks():
    try:
        import io, html as _html
        from datetime import date
        fmt    = request.args.get('fmt',    'docx')

        name       = (request.args.get('name')            or '').strip()
        vassa      = (request.args.get('vassa_years')     or '').strip()
        monk_type  = (request.args.get('monk_type')       or '').strip()
        residence  = (request.args.get('residence')       or '').strip()
        position   = (request.args.get('position')        or '').strip()
        edu        = (request.args.get('education_level') or '').strip()
        acad       = (request.args.get('academic_year')   or '').strip()
        sort_vassa = (request.args.get('sort_vassa')      or '').strip()

        where_parts, where_params = [], []
        if name:      where_parts.append("fullname ILIKE %s");      where_params.append(f'%{name}%')
        if vassa:     where_parts.append("vassa_years = %s");        where_params.append(int(vassa))
        if monk_type: where_parts.append("monk_type = %s");          where_params.append(monk_type)
        if residence: where_parts.append("residence = %s");          where_params.append(residence)
        if position:  where_parts.append("position = %s");           where_params.append(position)
        if edu:       where_parts.append("education_level = %s");    where_params.append(edu)
        if acad:      where_parts.append("academic_year = %s");      where_params.append(acad)
        where_sql = ('WHERE ' + ' AND '.join(where_parts)) if where_parts else ''

        conn = connect_db()
        cur  = conn.cursor()
        cur.execute(f"""
            SELECT fullname, vassa_years, monk_type, residence, position,
                   education_level, academic_year, created_at
            FROM monk_tbl {where_sql}
            ORDER BY fullname;
        """, where_params)
        rows = cur.fetchall()
        cur.close(); conn.close()

        monks = [{
            'fullname': r[0], 'vassa_years': r[1], 'monk_type': r[2],
            'residence': (r[3] or '').replace('_', ' '), 'position': r[4],
            'education_level': r[5] or '', 'academic_year': r[6] or '',
            'created_at': r[7]
        } for r in rows]

        # Default: the same role > type > vassa hierarchy used everywhere else.
        # A monk clicking the vassa column header on /view wants a pure vassa
        # sort instead — sort_vassa carries that explicit override into export.
        if sort_vassa in ('asc', 'desc'):
            monks.sort(key=lambda m: m['vassa_years'] or 0, reverse=(sort_vassa == 'desc'))
        else:
            monks = sort_attendance_monks(monks)

        today      = date.today().strftime('%d/%m/%Y')
        fname_base = f"monks_{date.today().isoformat()}"

        # ── DOCX ─────────────────────────────────────────────────────────────
        if fmt == 'docx':
            from docx import Document
            from docx.shared import Pt, Cm, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement

            def _shade(cell, hex_color):
                tc = cell._tc; tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), hex_color); tcPr.append(shd)

            doc = Document()
            sec = doc.sections[0]
            sec.left_margin = sec.right_margin = Cm(1.5)
            sec.top_margin  = sec.bottom_margin = Cm(1.5)

            h = doc.add_heading('វត្តនិរោធរង្សី — បញ្ជីព្រះសង្ឃ', 0)
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER
            sub = doc.add_paragraph(f"ថ្ងៃទី {today}  |  ចំនួនសរុប: {len(monks)} នាក់")
            sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
            sub.runs[0].font.size = Pt(10)
            sub.runs[0].font.color.rgb = RGBColor(0x71, 0x80, 0x96)
            doc.add_paragraph()

            hdrs   = ['#', 'ឈ្មោះ', 'វស្សា', 'ប្រភេទ', 'ស្នាក់នៅ', 'តួនាទី', 'កម្រិតសិក្សា']
            widths = [0.8, 3.8, 1.4, 1.8, 2.8, 3.4, 3.0]
            tbl = doc.add_table(rows=1, cols=len(hdrs))
            tbl.style = 'Table Grid'
            for i, (hdr, w) in enumerate(zip(hdrs, widths)):
                cell = tbl.rows[0].cells[i]; cell.width = Cm(w)
                _shade(cell, 'EEF2FF')
                p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(hdr); run.bold = True
                run.font.size = Pt(8.5); run.font.color.rgb = RGBColor(0x4A, 0x55, 0x68)

            for idx, m in enumerate(monks, 1):
                bg  = 'FFF8E1' if m['monk_type'] == 'ភិក្ខុ' else 'F1F8E9'
                vals = [str(idx), m['fullname'], f"{m['vassa_years']} ឆ្នាំ", m['monk_type'],
                        m['residence'], m['position'], m['education_level']]
                row = tbl.add_row()
                for j, (val, w) in enumerate(zip(vals, widths)):
                    cell = row.cells[j]; cell.width = Cm(w); _shade(cell, bg)
                    p = cell.paragraphs[0]
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if j in (0, 2) else WD_ALIGN_PARAGRAPH.LEFT
                    run = p.add_run(val); run.font.size = Pt(8.5)
                    if j == 1: run.bold = True

            buf = io.BytesIO(); doc.save(buf); buf.seek(0)
            return send_file(buf,
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                as_attachment=True, download_name=f"{fname_base}.docx")

        # ── EXCEL ────────────────────────────────────────────────────────────
        elif fmt == 'excel':
            import openpyxl
            from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
            from openpyxl.utils import get_column_letter

            wb = openpyxl.Workbook()
            ws = wb.active; ws.title = 'ព្រះសង្ឃ'

            hdrs = ['#', 'ឈ្មោះ', 'វស្សា', 'ប្រភេទ', 'ស្នាក់នៅ',
                    'តួនាទី', 'កម្រិតសិក្សា', 'ថ្នាក់', 'ថ្ងៃបញ្ចូល']
            hfill = PatternFill(start_color='667EEA', end_color='667EEA', fill_type='solid')
            hfont = Font(bold=True, color='FFFFFF', size=11, name='Calibri')
            thin  = Side(border_style='thin', color='D1D5DB')
            bdr   = Border(left=thin, right=thin, top=thin, bottom=thin)

            for col, h in enumerate(hdrs, 1):
                c = ws.cell(row=1, column=col, value=h)
                c.font = hfont; c.fill = hfill; c.border = bdr
                c.alignment = Alignment(horizontal='center', vertical='center')
            ws.row_dimensions[1].height = 22

            bfill = PatternFill(start_color='FFF8E1', end_color='FFF8E1', fill_type='solid')
            sfill = PatternFill(start_color='F1F8E9', end_color='F1F8E9', fill_type='solid')
            dfont = Font(size=10, name='Calibri')

            for row_i, m in enumerate(monks, 2):
                fill = bfill if m['monk_type'] == 'ភិក្ខុ' else sfill
                cre  = m['created_at'].strftime('%d/%m/%Y') if m['created_at'] else '—'
                vals = [row_i - 1, m['fullname'], m['vassa_years'], m['monk_type'],
                        m['residence'], m['position'], m['education_level'], m['academic_year'], cre]
                for col, val in enumerate(vals, 1):
                    c = ws.cell(row=row_i, column=col, value=val)
                    c.fill = fill; c.border = bdr; c.font = dfont
                    c.alignment = Alignment(
                        horizontal='center' if col in (1, 3, 9) else 'left',
                        vertical='center')

            for col, w in enumerate([5, 24, 8, 12, 18, 20, 15, 10, 13], 1):
                ws.column_dimensions[get_column_letter(col)].width = w

            buf = io.BytesIO(); wb.save(buf); buf.seek(0)
            return send_file(buf,
                mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
                as_attachment=True, download_name=f"{fname_base}.xlsx")

        # ── PDF ──────────────────────────────────────────────────────────────
        elif fmt == 'html':

            bhikkhus  = [m for m in monks if m['monk_type'] == 'ភិក្ខុ']
            samaneras = [m for m in monks if m['monk_type'] == 'សាមណេរ']

            def _trow(m, idx, stripe):
                bg  = '#f8fafc' if stripe else '#ffffff'
                return (
                    f'<tr style="background:{bg}">'
                    f'<td class="c">{idx}</td>'
                    f'<td class="name">{_html.escape(m["fullname"])}</td>'
                    f'<td class="c">{m["vassa_years"]} ឆ្នាំ</td>'
                    f'<td>{_html.escape(m["monk_type"])}</td>'
                    f'<td>{_html.escape(m["residence"])}</td>'
                    f'<td>{_html.escape(m["position"])}</td>'
                    f'<td>{_html.escape(m["education_level"])}</td>'
                    f'</tr>'
                )

            def _section(section_monks, title, hdr_bg, hdr_color, border_color):
                if not section_monks:
                    return ''
                rows = ''.join(_trow(m, i + 1, i % 2 == 1) for i, m in enumerate(section_monks))
                return f'''
                <div class="sec-title" style="background:{hdr_bg};color:{hdr_color};border-left:4px solid {border_color}">
                    {_html.escape(title)} &nbsp;<span class="sec-count">({len(section_monks)} នាក់)</span>
                </div>
                <table>
                    <thead><tr>
                        <th class="c" style="width:28px">#</th>
                        <th>ឈ្មោះ</th>
                        <th class="c" style="width:48px">វស្សា</th>
                        <th>ប្រភេទ</th>
                        <th>ស្នាក់នៅ</th>
                        <th>តួនាទី</th>
                        <th>កម្រិតសិក្សា</th>
                    </tr></thead>
                    <tbody>{rows}</tbody>
                </table>'''

            # Active filters line
            fp = []
            if name:      fp.append(f'ឈ្មោះ: <strong>{_html.escape(name)}</strong>')
            if monk_type: fp.append(f'ប្រភេទ: <strong>{_html.escape(monk_type)}</strong>')
            if residence: fp.append(f'កុដិ: <strong>{_html.escape(residence.replace("_"," "))}</strong>')
            if position:  fp.append(f'តួនាទី: <strong>{_html.escape(position)}</strong>')
            if edu:       fp.append(f'កម្រិតសិក្សា: <strong>{_html.escape(edu)}</strong>')
            if acad:      fp.append(f'ថ្នាក់: <strong>{_html.escape(acad)}</strong>')
            filter_html = (
                f'<div class="filter-bar">🔎 &nbsp;' + ' &nbsp;|&nbsp; '.join(fp) + '</div>'
                if fp else ''
            )

            css = (
                "@import url('https://fonts.googleapis.com/css2?family=Battambang:wght@400;700&family=Moul&display=swap');"
                "*, *::before, *::after{box-sizing:border-box;margin:0;padding:0}"
                "body{font-family:'Battambang','Khmer MN','Khmer Sangam MN',sans-serif;"
                "color:#1a202c;font-size:9px;background:#fff}"

                # Header
                ".header{background:linear-gradient(135deg,#667eea 0%,#764ba2 100%);"
                "color:#fff;padding:14px 18px 12px;border-radius:8px;margin-bottom:14px}"
                ".hdr-top{display:flex;justify-content:space-between;align-items:flex-start}"
                ".hdr-name{font-family:'Moul','Battambang',serif;font-size:16px;font-weight:400;letter-spacing:.4px}"
                ".hdr-sub{font-family:'Battambang',serif;font-size:14px;opacity:.85;margin-top:3px}"
                ".hdr-date{font-size:8.5px;opacity:.8;text-align:right}"
                ".hdr-divider{border:none;border-top:1px solid rgba(255,255,255,.3);margin:10px 0 8px}"
                ".hdr-stats{display:flex;gap:20px}"
                ".hdr-stat{font-size:8.5px;opacity:.9}"
                ".hdr-stat strong{font-size:13px;display:block;font-weight:700}"

                # Summary chips
                ".chips{display:flex;gap:8px;margin-bottom:12px}"
                ".chip{flex:1;padding:7px 10px;border-radius:6px;text-align:center}"
                ".chip-lbl{font-size:7.5px;color:#718096;display:block;margin-bottom:2px}"
                ".chip-val{font-size:15px;font-weight:700;display:block}"
                ".chip-total{background:#edf2f7}.chip-total .chip-val{color:#2d3748}"
                ".chip-b{background:#fff8e1}.chip-b .chip-val{color:#8a6100}"
                ".chip-s{background:#f1f8e9}.chip-s .chip-val{color:#1b5e20}"

                # Filter bar
                ".filter-bar{background:#f7fafc;border:1px solid #e2e8f0;border-radius:5px;"
                "padding:5px 10px;font-size:8px;color:#4a5568;margin-bottom:12px}"

                # Section title
                ".sec-title{padding:5px 10px;font-size:10px;font-weight:700;"
                "border-radius:4px;margin-bottom:4px;margin-top:14px}"
                ".sec-count{font-weight:400;font-size:8.5px;opacity:.8}"

                # Table
                "table{width:100%;border-collapse:collapse;margin-bottom:4px}"
                "thead tr{background:#4a5568}"
                "th{padding:6px 6px;text-align:left;font-size:7.5px;font-weight:700;"
                "color:#fff;white-space:nowrap;border-right:1px solid rgba(255,255,255,.15)}"
                "th:last-child{border-right:none}"
                "td{padding:5px 6px;font-size:8.5px;border-bottom:1px solid #edf2f7;"
                "vertical-align:middle}"
                ".c{text-align:center}"
                ".name{font-weight:600;color:#1a202c}"

                # Footer
                ".footer{position:running(footer);display:flex;justify-content:space-between;"
                "font-size:7px;color:#a0aec0;border-top:1px solid #e2e8f0;padding-top:4px;"
                "margin-top:6px}"
                "@page{size:A4 portrait;margin:12mm 10mm 18mm;"
                "@bottom-center{content:element(footer)}}"
            )

            html_str = f'''<!DOCTYPE html>
<html lang="km"><head><meta charset="UTF-8">
<style>{css}
@media print {{ @page {{ size: A4 portrait; margin: 12mm; }} }}
</style></head><body>

<div class="header">
    <div class="hdr-top">
        <div>
            <div class="hdr-name">វត្តនិរោធរង្សី</div>
            <div class="hdr-sub">Pagoda Niroth Rangsay — បញ្ជីព្រះសង្ឃ</div>
        </div>
        <div class="hdr-date">
            ថ្ងៃទី {_html.escape(today)}<br>
            ស្ថានភាព​ ​បច្ចុប្បន្ន
        </div>
    </div>
    <hr class="hdr-divider">
    <div class="hdr-stats">
        <div class="hdr-stat"><strong>{len(monks)}</strong>ព្រះសង្ឃ​សរុប</div>
        <div class="hdr-stat"><strong>{len(bhikkhus)}</strong>ភិក្ខុ</div>
        <div class="hdr-stat"><strong>{len(samaneras)}</strong>សាមណេរ</div>
    </div>
</div>

{filter_html}

{_section(bhikkhus,  'ផ្នែកទី ១ — ភិក្ខុ',   '#fff8e1','#8a6100','#f0c040')}
{_section(samaneras, 'ផ្នែកទី ២ — សាមណេរ', '#f1f8e9','#1b5e20','#66bb6a')}

<div class="footer">
    <span>វត្តនិរោធរង្សី — បញ្ជីព្រះសង្ឃ</span>
    <span>ថ្ងៃទី {_html.escape(today)}</span>
</div>
</body></html>'''

            return html_str

        return jsonify({'success': False, 'message': 'Unknown format'}), 400

    except ImportError as e:
        return jsonify({'success': False, 'message': f'Missing package: {e}'}), 500
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 500


# ============ LAYOUT EXPORT HELPERS ============

_LAYOUT_BHIKKHU_RANK = ROLE_RANK
_LAYOUT_SAMANERA_ADMIN_RANK = {'មេកុដិ': 1, 'អនុកុដិ': 2}


def _load_seat_order_from_db():
    import json as _json
    conn = connect_db()
    cur = conn.cursor()
    cur.execute('SELECT type, monk_ids FROM seat_order')
    rows = cur.fetchall()
    cur.close()
    conn.close()
    result = {'bhikkhu': None, 'samanera': None, 'grid_config': None}
    for row in rows:
        result[row[0]] = _json.loads(row[1])
    return result


def _build_grid_layout(default_sorted, stored_ids, rows, cols, reserve_first_row=False):
    """Fit monks into a rows×cols grid using stored seat order."""
    total = rows * cols
    grid = [None] * total
    by_id = {m['id']: m for m in default_sorted}
    placed = set()
    first_data_index = cols if reserve_first_row else 0
    chief_seat = cols // 2 if reserve_first_row else 0
    chief = (
        next((m for m in default_sorted if m['position'] == 'ព្រះអធិការ'), None)
        if reserve_first_row else None
    )
    stored_is_valid = True
    if reserve_first_row and stored_ids:
        first_id = stored_ids[chief_seat] if chief_seat < len(stored_ids) else None
        stored_is_valid = first_id == (chief['id'] if chief else None)
        if stored_is_valid:
            stored_is_valid = all(
                i == chief_seat
                or (stored_ids[i] if i < len(stored_ids) else None) is None
                for i in range(cols)
            )
    stored_for_layout = stored_ids if stored_is_valid else None

    # Keep the first bhikkhu row exclusively for the chief monk.
    if reserve_first_row and total > 0:
        if chief:
            grid[chief_seat] = chief
            placed.add(chief['id'])

    if stored_for_layout:
        for i, mid in enumerate(stored_for_layout[:total]):
            if reserve_first_row and i < first_data_index:
                continue
            if mid is None:
                continue
            monk = by_id.get(mid)
            if monk and mid not in placed:
                grid[i] = monk
                placed.add(mid)

    remaining = [m for m in default_sorted if m['id'] not in placed]
    ri = 0
    for i in range(first_data_index, total):
        if grid[i] is None and ri < len(remaining):
            grid[i] = remaining[ri]
            placed.add(remaining[ri]['id'])
            ri += 1

    return grid


def _get_residing_monks_for_layout():
    raw = get_all_monks()
    return [
        {
            'id': m[0],
            'fullname': m[1],
            'vassa_years': m[2],
            'monk_type': m[3],
            'position': m[5],
        }
        for m in raw
        if (m[10] if len(m) > 10 and m[10] else _ACTIVE_LIVING_STATUS) != 'ឈប់ស្នាក់នៅ'
    ]


def _get_layout_seat_grids(br, bc, sr, sc):
    all_monks = _get_residing_monks_for_layout()
    bhikkhus = sorted(
        [m for m in all_monks if m['monk_type'] == 'ភិក្ខុ'],
        key=lambda m: (_LAYOUT_BHIKKHU_RANK.get(m['position'], 99), -m['vassa_years']),
    )
    samaneras = sorted(
        [m for m in all_monks if m['monk_type'] == 'សាមណេរ'],
        key=lambda m: (
            _LAYOUT_SAMANERA_ADMIN_RANK.get(m['position'], 99),
            -m['vassa_years'],
            m['fullname'],
        ),
    )
    seat_order = _load_seat_order_from_db()
    bhikkhu_grid = _build_grid_layout(
        bhikkhus, seat_order.get('bhikkhu'), br, bc, reserve_first_row=True
    )
    samanera_grid = _build_grid_layout(samaneras, seat_order.get('samanera'), sr, sc)
    return bhikkhu_grid, samanera_grid


def _build_layout_export_html(bhikkhu_grid, samanera_grid, br, bc, sr, sc):
    import html as _html
    from datetime import date

    def build_grid(monks, rows, cols, type_):
        cells = []
        for r in range(rows):
            for c in range(cols):
                idx = r * cols + c
                monk = monks[idx] if idx < len(monks) else None
                num = idx + 1
                if monk:
                    sub = monk['position'] if (
                        type_ == 'bhikkhu' or monk['position'] in _LAYOUT_SAMANERA_ADMIN_RANK
                    ) else f"វស្សា {monk['vassa_years']}"
                    cells.append(
                        f'<td class="cell filled">'
                        f'<span class="num">{num}</span>'
                        f'<span class="name">{_html.escape(monk["fullname"])}</span>'
                        f'<span class="sub">{_html.escape(sub)}</span>'
                        f'</td>'
                    )
                else:
                    cells.append(f'<td class="cell empty"><span class="num-e">{num}</span></td>')
        rows_html = ''
        for r in range(rows):
            rows_html += '<tr>' + ''.join(cells[r * cols:(r + 1) * cols]) + '</tr>'
        return f'<table>{rows_html}</table>'

    bhikkhu_count = sum(1 for m in bhikkhu_grid if m)
    samanera_count = sum(1 for m in samanera_grid if m)
    today = date.today().strftime('%d/%m/%Y')
    css = (
        "@import url('https://fonts.googleapis.com/css2?family=Battambang:wght@400;700&family=Moul&display=swap');"
        "*, *::before, *::after{box-sizing:border-box;margin:0;padding:0}"
        "body{font-family:'Battambang','Khmer MN',sans-serif;color:#1a202c;font-size:8.5px;background:#fff}"
        ".header{background:linear-gradient(135deg,#667eea 0%,#764ba2 100%);"
        "color:#fff;padding:10px 14px 9px;border-radius:7px;margin-bottom:10px}"
        ".hdr-row{display:flex;justify-content:space-between;align-items:center}"
        ".hdr-title{font-family:'Moul','Battambang',serif;font-size:16px;font-weight:400;letter-spacing:.3px}"
        ".hdr-sub{font-family:'Battambang',serif;font-size:14px;opacity:.85;margin-top:2px}"
        ".hdr-right{text-align:right;font-size:8px;opacity:.85}"
        ".hdr-divider{border:none;border-top:1px solid rgba(255,255,255,.3);margin:7px 0 6px}"
        ".hdr-stats{display:flex;gap:18px}"
        ".hdr-stat{font-size:8px;opacity:.9}"
        ".hdr-stat strong{font-size:12px;display:block;font-weight:700}"
        ".sec{padding:4px 9px;font-size:10px;font-weight:700;border-radius:4px;margin:10px 0 4px}"
        ".sec-b{background:#fff8e1;color:#8a6100;border-left:4px solid #f0c040}"
        ".sec-s{background:#f1f8e9;color:#1b5e20;border-left:4px solid #66bb6a}"
        ".sec-sub{font-size:7.5px;font-weight:400;opacity:.75;margin-left:6px}"
        "table{width:100%;border-collapse:collapse;margin-bottom:8px;table-layout:fixed}"
        "td{border:1px solid #e2e8f0;vertical-align:top;padding:3px 4px;height:42px}"
        ".filled{background:#f7f8fa}"
        ".empty{background:#fafbfc}"
        ".num{display:block;font-size:6.5px;color:#a0aec0;line-height:1}"
        ".num-e{display:block;font-size:6.5px;color:#e2e8f0}"
        ".name{display:block;font-size:8px;font-weight:700;color:#1a202c;margin:2px 0 1px;line-height:1.3}"
        ".sub{display:block;font-size:7px;color:#718096;line-height:1.2}"
        ".footer{display:flex;justify-content:space-between;font-size:7px;color:#a0aec0;"
        "border-top:1px solid #e2e8f0;padding-top:6px;margin-top:10px}"
        "@page{size:A4 portrait;margin:10mm 12mm 16mm}"
    )
    return f'''<!DOCTYPE html>
<html lang="km"><head><meta charset="UTF-8"><style>{css}
@media print {{ @page {{ size: A4 portrait; margin: 12mm; }} }}
</style></head><body>

<div class="header">
  <div class="hdr-row">
    <div>
      <div class="hdr-title">ប្លង់អាសនៈព្រះសង្ឃ — វត្តនិរោធរង្សី</div>
      <div class="hdr-sub">Pagoda Niroth Rangsay — Seating Layout</div>
    </div>
    <div class="hdr-right">ថ្ងៃទី {today}<br>ស្ថានភាពបច្ចុប្បន្ន</div>
  </div>
  <hr class="hdr-divider">
  <div class="hdr-stats">
    <div class="hdr-stat"><strong>{bhikkhu_count}</strong>ភិក្ខុ</div>
    <div class="hdr-stat"><strong>{samanera_count}</strong>សាមណេរ</div>
    <div class="hdr-stat"><strong>{bhikkhu_count + samanera_count}</strong>ព្រះសង្ឃសរុប</div>
  </div>
</div>

<div class="sec sec-b">ផ្នែកទី ១ — ភិក្ខុ
  <span class="sec-sub">{bhikkhu_count} នាក់ &nbsp;|&nbsp; ក្រឡា {br}×{bc}={br * bc}</span>
</div>
{build_grid(bhikkhu_grid, br, bc, 'bhikkhu')}

<div class="sec sec-s">ផ្នែកទី ២ — សាមណេរ
  <span class="sec-sub">{samanera_count} នាក់ &nbsp;|&nbsp; ក្រឡា {sr}×{sc}={sr * sc}</span>
</div>
{build_grid(samanera_grid, sr, sc, 'samanera')}

<div class="footer">
  <span>វត្តនិរោធរង្សី — ប្លង់អាសនៈព្រះសង្ឃ</span>
  <span>ថ្ងៃទី {today}</span>
</div>
</body></html>'''


@main_bp.route('/api/export-layout')
def export_layout():
    """Export the seating layout as a Word (.docx) document"""
    import io
    from datetime import date
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.enum.section import WD_ORIENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    BHIKKHU_RANK = _LAYOUT_BHIKKHU_RANK
    SAMANERA_ADMIN_RANK = _LAYOUT_SAMANERA_ADMIN_RANK

    def clamp(val, lo, hi):
        return max(lo, min(hi, int(val)))

    def shade_cell(cell, hex_color):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), hex_color)
        tcPr.append(shd)

    def fill_cell(cell, number, name, sub):
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        shade_cell(cell, 'F7F8FA')

        p0 = cell.paragraphs[0]
        p0.paragraph_format.space_before = Pt(2)
        p0.paragraph_format.space_after  = Pt(0)
        r0 = p0.add_run(str(number))
        r0.font.size = Pt(7)
        r0.font.color.rgb = RGBColor(0xA0, 0xAE, 0xC0)

        p1 = cell.add_paragraph()
        p1.paragraph_format.space_before = Pt(1)
        p1.paragraph_format.space_after  = Pt(0)
        r1 = p1.add_run(name)
        r1.font.size = Pt(9)
        r1.font.bold = True
        r1.font.color.rgb = RGBColor(0x1A, 0x20, 0x2C)

        p2 = cell.add_paragraph()
        p2.paragraph_format.space_before = Pt(1)
        p2.paragraph_format.space_after  = Pt(2)
        r2 = p2.add_run(sub)
        r2.font.size = Pt(7.5)
        r2.font.color.rgb = RGBColor(0x71, 0x80, 0x96)

    def fill_empty(cell, number):
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        p0 = cell.paragraphs[0]
        p0.paragraph_format.space_before = Pt(2)
        r0 = p0.add_run(str(number))
        r0.font.size = Pt(7)
        r0.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

    def add_grid_table(doc, monks, rows, cols, col_width_cm, type_):
        table = doc.add_table(rows=rows, cols=cols)
        table.style = 'Table Grid'
        for col in table.columns:
            col.width = Cm(col_width_cm)

        for r in range(rows):
            row_obj = table.rows[r]
            row_obj.height = Cm(1.6)
            for c in range(cols):
                idx = r * cols + c
                cell = row_obj.cells[c]
                monk = monks[idx] if idx < len(monks) else None
                if monk:
                    is_admin = monk['position'] in SAMANERA_ADMIN_RANK
                    sub = monk['position'] if (type_ == 'bhikkhu' or is_admin) else f"វស្សា {monk['vassa_years']}"
                    fill_cell(cell, idx + 1, monk['fullname'], sub)
                else:
                    fill_empty(cell, idx + 1)

    try:
        br = clamp(request.args.get('br', 3), 1, 30)
        bc = clamp(request.args.get('bc', 5), 1, 30)
        sr = clamp(request.args.get('sr', 12), 1, 50)
        sc = clamp(request.args.get('sc', 10), 1, 30)

        bhikkhus, samaneras = _get_layout_seat_grids(br, bc, sr, sc)

        doc = Document()

        # Portrait A4
        sec = doc.sections[0]
        sec.left_margin   = Cm(1.5)
        sec.right_margin  = Cm(1.5)
        sec.top_margin    = Cm(1.5)
        sec.bottom_margin = Cm(1.5)

        usable_cm = 18.0  # ~210mm - 3cm margins

        # Title
        t = doc.add_heading('ប្លង់អាសនៈព្រះសង្ឃ — វត្តនិរោធរង្សី', 0)
        t.alignment = WD_ALIGN_PARAGRAPH.CENTER

        sub_p = doc.add_paragraph(f'ថ្ងៃទី {date.today().strftime("%d/%m/%Y")}')
        sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub_p.runs[0].font.size = Pt(11)
        sub_p.runs[0].font.color.rgb = RGBColor(0x71, 0x80, 0x96)

        doc.add_paragraph()

        # ── Section 1: Bhikkhu ──
        h1 = doc.add_heading('ផ្នែកទី ១ — ភិក្ខុ', 1)
        h1.runs[0].font.color.rgb = RGBColor(0x8A, 0x61, 0x00)

        info1 = doc.add_paragraph(
            f'ចំនួនភិក្ខុ: {sum(1 for m in bhikkhus if m)} នាក់  |  ក្រឡា: {br} × {bc} = {br*bc}'
        )
        info1.runs[0].font.size = Pt(10)
        info1.runs[0].font.color.rgb = RGBColor(0x71, 0x80, 0x96)

        col_w1 = round(usable_cm / bc, 2)
        add_grid_table(doc, bhikkhus, br, bc, col_w1, 'bhikkhu')

        # Page break before Samanera
        doc.add_page_break()

        # ── Section 2: Samanera ──
        h2 = doc.add_heading('ផ្នែកទី ២ — សាមណេរ', 1)
        h2.runs[0].font.color.rgb = RGBColor(0x1B, 0x5E, 0x20)

        info2 = doc.add_paragraph(
            f'ចំនួនសាមណេរ: {sum(1 for m in samaneras if m)} នាក់  |  ក្រឡា: {sr} × {sc} = {sr*sc}'
        )
        info2.runs[0].font.size = Pt(10)
        info2.runs[0].font.color.rgb = RGBColor(0x71, 0x80, 0x96)

        col_w2 = round(usable_cm / sc, 2)
        add_grid_table(doc, samaneras, sr, sc, col_w2, 'samanera')

        output = io.BytesIO()
        doc.save(output)
        output.seek(0)

        filename = f'layout_{date.today().strftime("%Y%m%d")}.docx'
        return send_file(
            output,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=filename
        )

    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 500


@main_bp.route('/api/export-layout-pdf')
def export_layout_pdf():
    """Export the seating layout as PDF or HTML preview via WeasyPrint."""
    import io
    from datetime import date

    def clamp(val, lo, hi):
        return max(lo, min(hi, int(val or 0)))

    try:
        br = clamp(request.args.get('br', 3), 1, 30)
        bc = clamp(request.args.get('bc', 5), 1, 30)
        sr = clamp(request.args.get('sr', 12), 1, 50)
        sc = clamp(request.args.get('sc', 10), 1, 30)
        fmt = request.args.get('fmt', 'pdf').strip().lower()

        bhikkhu_grid, samanera_grid = _get_layout_seat_grids(br, bc, sr, sc)
        html_str = _build_layout_export_html(bhikkhu_grid, samanera_grid, br, bc, sr, sc)

        if fmt == 'html':
            return html_str, 200, {'Content-Type': 'text/html; charset=utf-8'}

        from weasyprint import HTML
        pdf_bytes = HTML(string=html_str).write_pdf()
        buf = io.BytesIO(pdf_bytes)
        fname = f'layout_{date.today().strftime("%Y%m%d")}.pdf'
        return send_file(buf, mimetype='application/pdf', as_attachment=True, download_name=fname)

    except ImportError:
        return jsonify({'success': False, 'message': 'WeasyPrint មិនទាន់ install — សូម run: pip install weasyprint'}), 500
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 500


# ================================================================
# MULTI-TIER REPORTING SYSTEM
# ================================================================


def _do_compile_period(conn, cur, period_start):
    """Aggregate the fixed 15-day block containing period_start into attendance_summaries
    and advance the period_tracker.  Returns (row_count, period_end)."""
    from datetime import timedelta
    _, period_end = _get_block_dates(period_start.isoformat())
    cur.execute("""
        SELECT monk_id,
               COALESCE(SUM(CASE WHEN status = 'absent'     THEN 1 ELSE 0 END), 0),
               COALESCE(SUM(CASE WHEN status = 'permission' THEN 1 ELSE 0 END), 0)
        FROM attendance_tbl
        WHERE date >= %s AND date <= %s
        GROUP BY monk_id
    """, (period_start.isoformat(), period_end.isoformat()))
    rows = cur.fetchall()
    for monk_id, absences, permissions in rows:
        cur.execute("""
            INSERT INTO attendance_summaries
                (monk_id, period_start, period_end, total_absences, total_permissions)
            VALUES (%s, %s, %s, %s, %s)
            ON CONFLICT (monk_id, period_start) DO UPDATE
                SET total_absences    = EXCLUDED.total_absences,
                    total_permissions = EXCLUDED.total_permissions
        """, (monk_id, period_start, period_end, int(absences), int(permissions)))
    new_start = period_end + timedelta(days=1)
    cur.execute("""
        UPDATE period_tracker
        SET current_period_start = %s, last_compiled_at = NOW()
        WHERE id = 1
    """, (new_start,))
    return len(rows), period_end


def _fetch_live_range(start_date, end_date):
    """Aggregate directly from attendance_tbl for any date range — no compiled summaries needed."""
    sql = """
        SELECT m.id, m.fullname, m.monk_type, m.position, m.vassa_years,
               m.residence, m.education_level, m.academic_year,
               COALESCE(SUM(CASE WHEN a.status = 'absent'     THEN 1 ELSE 0 END), 0),
               COALESCE(SUM(CASE WHEN a.status = 'permission' THEN 1 ELSE 0 END), 0)
        FROM monk_tbl m
        LEFT JOIN attendance_tbl a
            ON a.monk_id = m.id AND a.date >= %s AND a.date <= %s
        GROUP BY m.id, m.fullname, m.monk_type, m.position,
                 m.vassa_years, m.residence, m.education_level, m.academic_year
        HAVING COALESCE(SUM(CASE WHEN a.status = 'absent'     THEN 1 ELSE 0 END), 0) >= %s
            OR COALESCE(SUM(CASE WHEN a.status = 'permission' THEN 1 ELSE 0 END), 0) >= %s
        ORDER BY m.monk_type,
                 SUM(CASE WHEN a.status = 'absent'     THEN 1 ELSE 0 END) DESC,
                 SUM(CASE WHEN a.status = 'permission' THEN 1 ELSE 0 END) DESC
    """
    conn = connect_db()
    cur  = conn.cursor()
    cur.execute(sql, (start_date.isoformat(), end_date.isoformat(),
                      DISC_ABSENT_MIN, DISC_PERM_MIN))
    rows = cur.fetchall()
    cur.close(); conn.close()
    return [{
        'id':                r[0],
        'fullname':          r[1],
        'monk_type':         r[2],
        'position':          r[3],
        'vassa_years':       r[4],
        'residence':         (r[5] or '').replace('_', ' '),
        'education_level':   r[6] or '',
        'academic_year':     r[7] or '',
        'total_absences':    int(r[8] or 0),
        'total_permissions': int(r[9] or 0),
        'range_start':       start_date.isoformat(),
        'range_end':         end_date.isoformat(),
    } for r in rows]


def _summary_query(cur, start_str, end_str):
    """Aggregate summaries between two dates; apply disciplinary filter."""
    cur.execute("""
        SELECT m.id, m.fullname, m.monk_type, m.position, m.vassa_years,
               m.residence, m.education_level, m.academic_year,
               SUM(s.total_absences)    AS tot_abs,
               SUM(s.total_permissions) AS tot_perm,
               MIN(s.period_start)      AS range_start,
               MAX(s.period_end)        AS range_end
        FROM attendance_summaries s
        JOIN monk_tbl m ON m.id = s.monk_id
        WHERE s.period_start >= %s AND s.period_end <= %s
        GROUP BY m.id, m.fullname, m.monk_type, m.position,
                 m.vassa_years, m.residence, m.education_level, m.academic_year
        HAVING SUM(s.total_absences)    >= %s
            OR SUM(s.total_permissions) >= %s
        ORDER BY m.monk_type,
                 SUM(s.total_absences)    DESC,
                 SUM(s.total_permissions) DESC
    """, (start_str, end_str, DISC_ABSENT_MIN, DISC_PERM_MIN))
    return cur.fetchall()


def _rows_to_monks(rows):
    return [{
        'id':                 r[0],
        'fullname':           r[1],
        'monk_type':          r[2],
        'position':           r[3],
        'vassa_years':        r[4],
        'residence':          (r[5] or '').replace('_', ' '),
        'education_level':    r[6] or '',
        'academic_year':      r[7] or '',
        'total_absences':     int(r[8] or 0),
        'total_permissions':  int(r[9] or 0),
        'range_start':        r[10].isoformat() if r[10] else None,
        'range_end':          r[11].isoformat() if r[11] else None,
    } for r in rows]


# ---- Page routes ------------------------------------------------

@main_bp.route('/report/book')
def book_report_page():
    return render_template('book_report.html')


# ---- API: daily report -----------------------------------------

@main_bp.route('/api/attendance/daily-report', methods=['GET'])
def daily_report():
    """Return today's (or given date's) attendance records — no disciplinary filter."""
    date_str = request.args.get('date', _date.today().isoformat())
    try:
        conn = connect_db()
        cur  = conn.cursor()
        cur.execute("""
            SELECT m.id, m.fullname, m.monk_type, m.position, m.vassa_years,
                   m.residence, m.education_level, m.academic_year, a.status
            FROM attendance_tbl a
            JOIN monk_tbl m ON m.id = a.monk_id
            WHERE a.date = %s
            ORDER BY m.monk_type, a.status, m.fullname
        """, (date_str,))
        rows = cur.fetchall()
        cur.close(); conn.close()
        return jsonify({
            'success': True,
            'date': date_str,
            'records': [{
                'id':              r[0],
                'fullname':        r[1],
                'monk_type':       r[2],
                'position':        r[3],
                'vassa_years':     r[4],
                'residence':       (r[5] or '').replace('_', ' '),
                'education_level': r[6] or '',
                'academic_year':   r[7] or '',
                'status':          r[8],
            } for r in rows]
        })
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 500


# ---- API: compile period ----------------------------------------

@main_bp.route('/api/attendance/compile-period', methods=['POST'])
def compile_period_endpoint():
    """Compile the current 15-day block into attendance_summaries and advance the tracker."""
    conn = None
    try:
        conn = connect_db()
        cur  = conn.cursor()
        cur.execute("SELECT current_period_start FROM period_tracker WHERE id = 1")
        row = cur.fetchone()
        if not row:
            return jsonify({'success': False, 'message': 'Period tracker not initialised'}), 400
        period_start = row[0]
        count, period_end = _do_compile_period(conn, cur, period_start)
        conn.commit()
        cur.close()
        from datetime import timedelta as _td
        return jsonify({
            'success':           True,
            'compiled':          count,
            'period_start':      period_start.isoformat(),
            'period_end':        period_end.isoformat(),
            'next_period_start': (period_end + _td(days=1)).isoformat(),
        })
    except Exception as e:
        if conn: conn.rollback()
        return jsonify({'success': False, 'message': str(e)}), 500
    finally:
        if conn: conn.close()


# ---- API: list compiled periods ---------------------------------

@main_bp.route('/api/reports/periods', methods=['GET'])
def list_periods():
    """Return all compiled periods plus current active period info."""
    try:
        conn = connect_db()
        cur  = conn.cursor()
        cur.execute("""
            SELECT DISTINCT period_start, period_end
            FROM attendance_summaries
            ORDER BY period_start DESC
            LIMIT 100
        """)
        periods = [{'start': r[0].isoformat(), 'end': r[1].isoformat()}
                   for r in cur.fetchall()]
        cur.execute("SELECT current_period_start, last_compiled_at FROM period_tracker WHERE id = 1")
        row = cur.fetchone()
        cur.close(); conn.close()
        return jsonify({
            'success':              True,
            'compiled_periods':     periods,
            'current_period_start': row[0].isoformat() if row else None,
            'last_compiled_at':     row[1].isoformat() if row and row[1] else None,
        })
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 500


# ---- API: bi-weekly report --------------------------------------

@main_bp.route('/api/reports/biweekly', methods=['GET'])
def report_biweekly():
    """Bi-weekly report from summaries for a specific date (auto-block-snapped).
    Filter: absences >= 2 OR permissions >= 3."""
    date_str = request.args.get('period_start') or request.args.get('date', _date.today().isoformat())
    try:
        period_start, period_end = _get_block_dates(date_str)
        conn = connect_db()
        cur  = conn.cursor()
        rows = _summary_query(cur, period_start.isoformat(), period_end.isoformat())
        cur.close(); conn.close()
        return jsonify({
            'success':      True,
            'period_start': period_start.isoformat(),
            'period_end':   period_end.isoformat(),
            'monks':        _rows_to_monks(rows),
        })
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 500


# ---- API: monthly report ----------------------------------------

@main_bp.route('/api/reports/monthly', methods=['GET'])
def report_monthly():
    year_str  = request.args.get('year',  str(_date.today().year))
    month_str = request.args.get('month', str(_date.today().month))
    try:
        import calendar
        year, month = int(year_str), int(month_str)
        month_start = _date(year, month, 1)
        month_end   = _date(year, month, calendar.monthrange(year, month)[1])
        monks = _fetch_live_range(month_start, month_end)
        return jsonify({
            'success':      True,
            'year':         year,
            'month':        month,
            'period_start': month_start.isoformat(),
            'period_end':   month_end.isoformat(),
            'monks':        monks,
        })
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 500


# ---- API: annual report -----------------------------------------

@main_bp.route('/api/reports/annual', methods=['GET'])
def report_annual():
    year_str = request.args.get('year', str(_date.today().year))
    try:
        year       = int(year_str)
        year_start = _date(year, 1, 1)
        year_end   = _date(year, 12, 31)
        monks = _fetch_live_range(year_start, year_end)
        return jsonify({
            'success':      True,
            'year':         year,
            'period_start': year_start.isoformat(),
            'period_end':   year_end.isoformat(),
            'monks':        monks,
        })
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 500


# ---- API: 3-year report -----------------------------------------

@main_bp.route('/api/reports/triennial', methods=['GET'])
def report_triennial():
    start_year_str = request.args.get('start_year', str(_date.today().year - 2))
    try:
        start_year   = int(start_year_str)
        end_year     = start_year + 2
        period_start = _date(start_year, 1, 1)
        period_end   = _date(end_year, 12, 31)
        monks = _fetch_live_range(period_start, period_end)
        return jsonify({
            'success':      True,
            'start_year':   start_year,
            'end_year':     end_year,
            'period_start': period_start.isoformat(),
            'period_end':   period_end.isoformat(),
            'monks':        monks,
        })
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 500


# ---- Unified export: fetch + normalise data for any report type -

def _fetch_export_data(report_type, args):
    """Return (monks, type_label, subtitle, period_start_iso, period_end_iso).
    monks list is normalised: absent_count, permission_count, absent_dates, perm_dates, status."""
    import calendar as _cal

    if report_type == 'daily':
        date_str = args.get('date', _date.today().isoformat())
        conn = connect_db(); cur = conn.cursor()
        cur.execute("""
            SELECT m.id, m.fullname, m.monk_type, m.position, m.vassa_years,
                   m.residence, m.education_level, m.academic_year, a.status
            FROM attendance_tbl a
            JOIN monk_tbl m ON m.id = a.monk_id
            WHERE a.date = %s
            ORDER BY m.fullname
        """, (date_str,))
        rows = cur.fetchall(); cur.close(); conn.close()
        d_fmt = _date.fromisoformat(date_str).strftime('%d/%m/%Y')
        monks = [{
            'id': r[0], 'fullname': r[1], 'monk_type': r[2], 'position': r[3],
            'vassa_years': r[4], 'residence': (r[5] or '').replace('_', ' '),
            'education_level': r[6] or '', 'academic_year': r[7] or '',
            'absent_count':     1 if r[8] == 'absent'     else 0,
            'permission_count': 1 if r[8] == 'permission' else 0,
            'absent_dates':     d_fmt if r[8] == 'absent'     else '',
            'perm_dates':       d_fmt if r[8] == 'permission' else '',
            'status': r[8],
        } for r in rows]
        monks = sort_attendance_monks(monks)
        return monks, 'ប្រចាំថ្ងៃ', d_fmt, date_str, date_str

    elif report_type == 'biweekly':
        monks, start_date, end_date = _fetch_report_rows(args)
        subtitle = f"{start_date.strftime('%d/%m/%Y')} ដល់ {end_date.strftime('%d/%m/%Y')} (១៥ ថ្ងៃ)"
        return monks, 'ប្រចាំ ១៥ ថ្ងៃ', subtitle, start_date.isoformat(), end_date.isoformat()

    elif report_type == 'monthly':
        date_str = args.get('date', _date.today().isoformat())
        d = _date.fromisoformat(date_str)
        year  = int(args.get('year',  d.year))
        month = int(args.get('month', d.month))
        ps = _date(year, month, 1)
        pe = _date(year, month, _cal.monthrange(year, month)[1])
        monks = _norm_summary_monks(_fetch_live_range(ps, pe))
        subtitle = f"{ps.strftime('%d/%m/%Y')} ដល់ {pe.strftime('%d/%m/%Y')}"
        return monks, 'ប្រចាំខែ', subtitle, ps.isoformat(), pe.isoformat()

    elif report_type == 'annual':
        date_str = args.get('date', _date.today().isoformat())
        year = int(args.get('year', _date.fromisoformat(date_str).year))
        ps = _date(year, 1, 1); pe = _date(year, 12, 31)
        monks = _norm_summary_monks(_fetch_live_range(ps, pe))
        subtitle = f"{ps.strftime('%d/%m/%Y')} ដល់ {pe.strftime('%d/%m/%Y')}"
        return monks, 'ប្រចាំឆ្នាំ', subtitle, ps.isoformat(), pe.isoformat()

    else:  # triennial
        date_str   = args.get('date', _date.today().isoformat())
        d          = _date.fromisoformat(date_str)
        start_year = int(args.get('start_year', d.year - 2))
        ps = _date(start_year, 1, 1); pe = _date(start_year + 2, 12, 31)
        monks = _norm_summary_monks(_fetch_live_range(ps, pe))
        subtitle = f"{ps.strftime('%d/%m/%Y')} ដល់ {pe.strftime('%d/%m/%Y')}"
        return monks, 'ប្រចាំ ៣ ឆ្នាំ', subtitle, ps.isoformat(), pe.isoformat()


def _norm_summary_monks(monks):
    """Convert _rows_to_monks shape to unified export shape."""
    result = []
    for m in monks:
        rng = ''
        if m.get('range_start') and m.get('range_end'):
            rs  = _date.fromisoformat(m['range_start']).strftime('%d/%m/%Y')
            re  = _date.fromisoformat(m['range_end']).strftime('%d/%m/%Y')
            rng = f"{rs} → {re}"
        result.append({**m,
            'absent_count':     m['total_absences'],
            'permission_count': m['total_permissions'],
            'absent_dates':     '',
            'perm_dates':       rng,
            'status':           None,
        })
    return result


# ---- Unified export: Ministry document header (docx) --------------

def _add_ministry_header_docx(doc, lunar_str):
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    def no_border(table):
        tbl_pr = table._tbl.tblPr
        borders = OxmlElement('w:tblBorders')
        for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
            el = OxmlElement(f'w:{edge}')
            el.set(qn('w:val'), 'none')
            borders.append(el)
        tbl_pr.append(borders)

    def line(cell, text, *, bold=False, color=None, size=10, first=False):
        p = cell.paragraphs[0] if first else cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(size)
        if color: run.font.color.rgb = RGBColor(*color)

    hdr_tbl = doc.add_table(rows=1, cols=2)
    hdr_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    no_border(hdr_tbl)
    left, right = hdr_tbl.rows[0].cells
    left.width = right.width = Cm(8.5)

    line(left, '[ និមិត្តសញ្ញាក្រសួង ]', color=(0x2B, 0x6C, 0xB0), size=9, first=True)
    for txt in _HDR_LEFT_LINES:
        line(left, txt, color=(0x2B, 0x6C, 0xB0), size=10)

    line(right, _HDR_RIGHT_LINES[0], bold=True, size=12, first=True)
    line(right, _HDR_RIGHT_LINES[1], bold=True, size=11)

    doc.add_paragraph()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = t.add_run(_HDR_MAIN_TITLE)
    tr.bold = True
    tr.font.size = Pt(16)

    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = s.add_run(f"កាលបរិច្ឆេទ ជាចន្ទគតិ: {lunar_str}")
    sr.font.size = Pt(10)
    sr.font.color.rgb = RGBColor(0x71, 0x80, 0x96)
    doc.add_paragraph()


# ---- Unified export: build docx ----------------------------------

def _make_export_docx(monks, type_label, subtitle, report_type):
    import io
    from datetime import date as _d
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    def shade(cell, hex_color):
        tcPr = cell._tc.get_or_add_tcPr()
        shd  = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), hex_color); tcPr.append(shd)

    if report_type == 'daily':
        headers = ['#', 'ឈ្មោះ', 'ប្រភេទ', 'តួនាទី', 'វស្សា', 'ស្នាក់នៅ', 'ស្ថានភាព']
        widths  = [0.5, 3.0, 1.5, 2.5, 1.0, 2.0, 2.0]
    elif report_type == 'biweekly':
        headers = ['#', 'ឈ្មោះ', 'តួនាទី', 'វស្សា', 'ស្នាក់នៅ', 'ការសិក្សា', '❌', '📋', 'ថ្ងៃ', 'ស្ថានភាព']
        widths  = [0.5, 3.0, 2.5, 1.0, 2.0, 1.6, 0.7, 0.7, 2.8, 1.8]
    else:
        headers = ['#', 'ឈ្មោះ', 'ប្រភេទ', 'តួនាទី', 'វស្សា', '❌', '📋', 'ចន្លោះ', 'ស្ថានភាព']
        widths  = [0.5, 3.0, 1.5, 2.5, 1.0, 0.7, 0.7, 3.5, 1.8]

    def add_table(doc, section_monks):
        tbl = doc.add_table(rows=1, cols=len(headers))
        tbl.style = 'Table Grid'
        for i, (h, w) in enumerate(zip(headers, widths)):
            cell = tbl.rows[0].cells[i]; cell.width = Cm(w)
            shade(cell, 'F7FAFC')
            p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(h); run.bold = True
            run.font.size = Pt(8.5); run.font.color.rgb = RGBColor(0x71, 0x80, 0x96)

        for idx, m in enumerate(section_monks, 1):
            ab = m['absent_count']     >= DISC_ABSENT_MIN
            pr = m['permission_count'] >= DISC_PERM_MIN
            bg = 'FFF5F5' if ab else ('FFFAF0' if pr else 'FFFFFF')
            status_text = '⚠ លើសអវត្តមាន' if ab else ('⚠ លើសច្បាប់' if pr else '✓ ប្រក្រតី')
            row = tbl.add_row()

            if report_type == 'daily':
                s = '❌ អវត្តមាន' if m.get('status') == 'absent' else '📋 ច្បាប់'
                vals    = [str(idx), m['fullname'], m['monk_type'], m['position'],
                           f"{m['vassa_years']} ឆ្នាំ", m['residence'], s]
                centers = {0, 2, 4, 6}
                abs_col = perm_col = -1
            elif report_type == 'biweekly':
                edu    = f"{m['education_level']} {m['academic_year']}".strip()
                dp     = []
                if m['absent_dates']: dp.append(f"❌ {m['absent_dates']}")
                if m['perm_dates']:   dp.append(f"📋 {m['perm_dates']}")
                vals    = [str(idx), m['fullname'], m['position'], f"{m['vassa_years']} ឆ្នាំ",
                           m['residence'], edu,
                           str(m['absent_count'])     if m['absent_count']     else '—',
                           str(m['permission_count']) if m['permission_count'] else '—',
                           '\n'.join(dp) if dp else '—', status_text]
                centers = {0, 3, 6, 7}; abs_col, perm_col = 6, 7
            else:
                vals    = [str(idx), m['fullname'], m['monk_type'], m['position'],
                           f"{m['vassa_years']} ឆ្នាំ",
                           str(m['absent_count'])     if m['absent_count']     else '—',
                           str(m['permission_count']) if m['permission_count'] else '—',
                           m['perm_dates'] or '—', status_text]
                centers = {0, 2, 4, 5, 6}; abs_col, perm_col = 5, 6

            for j, (val, w) in enumerate(zip(vals, widths)):
                cell = row.cells[j]; cell.width = Cm(w)
                shade(cell, bg)
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if j in centers else WD_ALIGN_PARAGRAPH.LEFT
                run = p.add_run(val); run.font.size = Pt(8.5)
                if j == 1:            run.bold = True
                if ab and j == abs_col:  run.font.color.rgb = RGBColor(0xC5, 0x30, 0x30)
                if pr and j == perm_col: run.font.color.rgb = RGBColor(0xC0, 0x56, 0x21)

    from docx.enum.section import WD_ORIENT

    doc = Document()
    sec = doc.sections[0]
    sec.left_margin = sec.right_margin = Cm(1.5)
    sec.top_margin  = sec.bottom_margin = Cm(1.5)

    _add_ministry_header_docx(doc, khmer_lunar_date(_d.today()))

    t = doc.add_heading(f'របាយការណ៍វត្តមាន ({type_label})', 1)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph(f"ចន្លោះ: {subtitle}")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.size = Pt(10)
    sub.runs[0].font.color.rgb = RGBColor(0x71, 0x80, 0x96)
    doc.add_paragraph()

    bhikkhus  = [m for m in monks if m['monk_type'] == 'ភិក្ខុ']
    samaneras = [m for m in monks if m['monk_type'] == 'សាមណេរ']
    if bhikkhus:
        h1 = doc.add_heading('ផ្នែកទី ១ — ភិក្ខុ', 1)
        h1.runs[0].font.color.rgb = RGBColor(0x8A, 0x61, 0x00)
        add_table(doc, bhikkhus); doc.add_paragraph()
    if samaneras:
        h2 = doc.add_heading('ផ្នែកទី ២ — សាមណេរ', 1)
        h2.runs[0].font.color.rgb = RGBColor(0x1B, 0x5E, 0x20)
        add_table(doc, samaneras); doc.add_paragraph()

    av = sum(1 for m in monks if m['absent_count']     >= DISC_ABSENT_MIN)
    pv = sum(1 for m in monks if m['permission_count'] >= DISC_PERM_MIN)
    sp = doc.add_paragraph(
        f"📊 សរុប {len(monks)} នាក់  |  ❌ លើសអវត្តមាន: {av} នាក់  |  📋 លើសច្បាប់: {pv} នាក់"
    )
    sp.runs[0].font.size = Pt(10); sp.runs[0].bold = True

    buf = io.BytesIO(); doc.save(buf); buf.seek(0)
    return buf


# ---- Unified export: build HTML for PDF --------------------------

def _make_export_html(monks, type_label, subtitle, report_type):
    import html as _h

    if report_type == 'daily':
        thead = ('<th>#</th><th>ឈ្មោះ</th><th>ប្រភេទ</th><th>តួនាទី</th>'
                 '<th>វស្សា</th><th>ស្នាក់នៅ</th><th>ស្ថានភាព</th>')
    elif report_type == 'biweekly':
        thead = ('<th>#</th><th>ឈ្មោះ</th><th>តួនាទី</th><th>វស្សា</th><th>ស្នាក់នៅ</th>'
                 '<th>ការសិក្សា</th><th>❌</th><th>📋</th><th>ថ្ងៃ</th><th>ស្ថានភាព</th>')
    else:
        thead = ('<th>#</th><th>ឈ្មោះ</th><th>ប្រភេទ</th><th>តួនាទី</th>'
                 '<th>វស្សា</th><th>❌</th><th>📋</th><th>ចន្លោះ</th><th>ស្ថានភាព</th>')

    def row_html(m, idx):
        ab = m['absent_count']     >= DISC_ABSENT_MIN
        pr = m['permission_count'] >= DISC_PERM_MIN
        bg = '#fff5f5' if ab else ('#fffaf0' if pr else '#ffffff')
        badge = ('<span class="badge-danger">⚠ លើសអវត្តមាន</span>' if ab else
                 '<span class="badge-warning">⚠ លើសច្បាប់</span>' if pr else
                 '<span class="badge-ok">✓ ប្រក្រតី</span>')
        n  = f'<td class="num">{idx}</td>'
        nm = f'<td><strong>{_h.escape(m["fullname"])}</strong></td>'

        if report_type == 'daily':
            sb = ('<span class="badge-danger">❌ អវត្តមាន</span>'
                  if m.get('status') == 'absent' else
                  '<span class="badge-warning">📋 ច្បាប់</span>')
            return (f'<tr style="background:{bg}">{n}{nm}'
                    f'<td class="num">{_h.escape(m["monk_type"])}</td>'
                    f'<td>{_h.escape(m["position"])}</td>'
                    f'<td class="num">{m["vassa_years"]} ឆ្នាំ</td>'
                    f'<td>{_h.escape(m["residence"])}</td>'
                    f'<td>{sb}</td></tr>')

        elif report_type == 'biweekly':
            ac  = 'color:#c53030;font-weight:bold' if ab else 'color:#718096'
            pc  = 'color:#c05621;font-weight:bold' if pr else 'color:#718096'
            edu = _h.escape(f"{m['education_level']} {m['academic_year']}".strip())
            dp  = []
            if m['absent_dates']: dp.append(f'<span style="color:#c53030">❌ {_h.escape(m["absent_dates"])}</span>')
            if m['perm_dates']:   dp.append(f'<span style="color:#c05621">📋 {_h.escape(m["perm_dates"])}</span>')
            dc  = '<br>'.join(dp) if dp else '—'
            return (f'<tr style="background:{bg}">{n}{nm}'
                    f'<td>{_h.escape(m["position"])}</td>'
                    f'<td class="num">{m["vassa_years"]} ឆ្នាំ</td>'
                    f'<td>{_h.escape(m["residence"])}</td><td>{edu}</td>'
                    f'<td class="num" style="{ac}">{m["absent_count"] or "—"}</td>'
                    f'<td class="num" style="{pc}">{m["permission_count"] or "—"}</td>'
                    f'<td class="dates">{dc}</td><td>{badge}</td></tr>')

        else:
            ac  = 'color:#c53030;font-weight:bold' if ab else 'color:#718096'
            pc  = 'color:#c05621;font-weight:bold' if pr else 'color:#718096'
            rng = (f'<span style="font-size:15px">{_h.escape(m["perm_dates"])}</span>'
                   if m['perm_dates'] else '—')
            return (f'<tr style="background:{bg}">{n}{nm}'
                    f'<td class="num">{_h.escape(m["monk_type"])}</td>'
                    f'<td>{_h.escape(m["position"])}</td>'
                    f'<td class="num">{m["vassa_years"]} ឆ្នាំ</td>'
                    f'<td class="num" style="{ac}">{m["absent_count"] or "—"}</td>'
                    f'<td class="num" style="{pc}">{m["permission_count"] or "—"}</td>'
                    f'<td class="dates">{rng}</td><td>{badge}</td></tr>')

    def section_html(sm, title, sec_cls, border_color):
        if not sm: return ''
        rows = ''.join(row_html(m, i + 1) for i, m in enumerate(sm))
        return (
            f'<div class="sec {sec_cls}">'
            f'{_h.escape(title)}'
            f'<span class="sec-count">({len(sm)} នាក់)</span></div>'
            f'<table><thead><tr class="thead-row">{thead}</tr></thead>'
            f'<tbody>{rows}</tbody></table>'
        )

    bhikkhus  = [m for m in monks if m['monk_type'] == 'ភិក្ខុ']
    samaneras = [m for m in monks if m['monk_type'] == 'សាមណេរ']
    av = sum(1 for m in monks if m['absent_count']     >= DISC_ABSENT_MIN)
    pv = sum(1 for m in monks if m['permission_count'] >= DISC_PERM_MIN)
    cl = len(monks) - av - pv

    css = (
        "@import url('https://fonts.googleapis.com/css2?family=Battambang:wght@400;700&family=Moul&display=swap');"
        "*, *::before, *::after{box-sizing:border-box;margin:0;padding:0}"
        "body{font-family:'Battambang','Khmer MN','Khmer Sangam MN',sans-serif;"
        "color:#1a202c;font-size:17px;background:#fff}"

        # Header
        ".header{background:linear-gradient(135deg,#667eea 0%,#764ba2 100%);"
        "color:#fff;padding:16px 18px 14px;border-radius:7px;margin-bottom:14px}"
        ".hdr-top{display:flex;justify-content:space-between;align-items:flex-start}"
        ".hdr-title{font-family:'Moul','Battambang',serif;font-size:16px;font-weight:400;letter-spacing:.3px}"
        ".hdr-sub{font-family:'Battambang',serif;font-size:14px;opacity:.85;margin-top:4px}"
        ".hdr-right{text-align:right;font-size:15px;opacity:.85}"
        ".hdr-divider{border:none;border-top:1px solid rgba(255,255,255,.3);margin:10px 0 8px}"
        ".hdr-stats{display:flex;gap:20px}"
        ".hdr-stat{font-size:15px;opacity:.9}"
        ".hdr-stat strong{font-size:23px;display:block;font-weight:700}"

        # Chips
        ".chips{display:flex;gap:8px;margin-bottom:13px}"
        ".chip{flex:1;padding:7px 9px;border-radius:6px;text-align:center}"
        ".chip-lbl{font-size:13px;display:block;margin-bottom:2px}"
        ".chip-val{font-size:24px;font-weight:700;display:block}"
        ".c-total{background:#edf2f7}.c-total .chip-lbl{color:#718096}.c-total .chip-val{color:#2d3748}"
        ".c-abs{background:#fed7d7}.c-abs .chip-lbl{color:#c53030}.c-abs .chip-val{color:#c53030}"
        ".c-perm{background:#feebc8}.c-perm .chip-lbl{color:#c05621}.c-perm .chip-val{color:#c05621}"
        ".c-ok{background:#c6f6d5}.c-ok .chip-lbl{color:#276749}.c-ok .chip-val{color:#276749}"

        # Section
        ".sec{padding:6px 12px;font-size:18px;font-weight:700;border-radius:4px;margin:14px 0 5px}"
        ".sec-b{background:#fff8e1;color:#8a6100;border-left:4px solid #f0c040}"
        ".sec-s{background:#f1f8e9;color:#1b5e20;border-left:4px solid #66bb6a}"
        ".sec-count{font-weight:400;font-size:14px;opacity:.75;margin-left:7px}"

        # Table
        "table{width:100%;border-collapse:collapse;margin-bottom:5px}"
        ".thead-row{background:#4a5568}"
        "th{padding:8px 8px;text-align:left;font-size:14px;font-weight:700;"
        "color:#fff;white-space:nowrap;border-right:1px solid rgba(255,255,255,.15)}"
        "th:last-child{border-right:none}"
        "td{padding:7px 8px;font-size:16px;border-bottom:1px solid #edf2f7;vertical-align:middle}"
        "tr:nth-child(even) td{background:#f8fafc}"
        ".num{text-align:center}"
        ".dates{font-size:14px;line-height:1.6}"
        ".badge-danger{background:#fed7d7;color:#c53030;padding:2px 7px;"
        "border-radius:8px;font-size:14px;font-weight:700;white-space:nowrap}"
        ".badge-warning{background:#feebc8;color:#c05621;padding:2px 7px;"
        "border-radius:8px;font-size:14px;font-weight:700;white-space:nowrap}"
        ".badge-ok{background:#c6f6d5;color:#276749;padding:2px 7px;"
        "border-radius:8px;font-size:14px;font-weight:700;white-space:nowrap}"

        # Footer
        ".footer{position:running(footer);display:flex;justify-content:space-between;"
        "font-size:12px;color:#a0aec0;border-top:1px solid #e2e8f0;padding-top:5px}"
        "@page{size:A4 portrait;margin:12mm 10mm 17mm;"
        "@bottom-center{content:element(footer)}}"

        # Ministry document header
        ".mh{display:flex;justify-content:space-between;align-items:flex-start;margin-bottom:6px}"
        ".mh-left{display:flex;flex-direction:column;align-items:center;"
        "color:#2b6cb0;font-size:13px;line-height:1.5;flex:1}"
        ".mh-logo{width:44px;height:44px;border:1px dashed #2b6cb0;border-radius:4px;"
        "display:flex;align-items:center;justify-content:center;font-size:7px;"
        "margin-bottom:4px;text-align:center;color:#2b6cb0}"
        ".mh-right{text-align:center;font-weight:700;font-size:15px;line-height:1.6;flex:1}"
        ".mh-title{text-align:center;font-family:'Moul','Battambang',serif;font-weight:400;font-size:16px;margin-top:4px}"
        ".mh-sub{text-align:center;font-family:'Battambang',serif;font-size:14px;color:#718096;margin:2px 0 12px}"
    )

    lunar_str = khmer_lunar_date(_date.today())
    ministry_header = f'''
<div class="mh">
  <div class="mh-left">
    <div class="mh-logo">[ និមិត្តសញ្ញា<br>ក្រសួង ]</div>
    <div>{_h.escape(_HDR_LEFT_LINES[0])}</div>
    <div>{_h.escape(_HDR_LEFT_LINES[1])}</div>
    <div>{_h.escape(_HDR_LEFT_LINES[2])}</div>
  </div>
  <div class="mh-right">
    <div>{_h.escape(_HDR_RIGHT_LINES[0])}</div>
    <div>{_h.escape(_HDR_RIGHT_LINES[1])}</div>
  </div>
</div>
<div class="mh-title">{_h.escape(_HDR_MAIN_TITLE)}</div>
<div class="mh-sub">កាលបរិច្ឆេទ ជាចន្ទគតិ: {_h.escape(lunar_str)}</div>'''

    return f'''<!DOCTYPE html>
<html lang="km"><head><meta charset="UTF-8"><style>{css}
@media print {{ @page {{ size: A4 portrait; margin: 12mm; }} }}
</style></head><body>

{ministry_header}

<div class="header">
  <div class="hdr-top">
    <div>
      <div class="hdr-title">វត្តនិរោធរង្សី — របាយការណ៍វត្តមាន</div>
      <div class="hdr-sub">Pagoda Niroth Rangsay &nbsp;|&nbsp; {_h.escape(type_label)}</div>
    </div>
    <div class="hdr-right">ចន្លោះ: {_h.escape(subtitle)}</div>
  </div>
  <hr class="hdr-divider">
  <div class="hdr-stats">
    <div class="hdr-stat"><strong>{len(monks)}</strong>ព្រះសង្ឃ​សរុប</div>
    <div class="hdr-stat"><strong>{av}</strong>លើស​អវត្តមាន</div>
    <div class="hdr-stat"><strong>{pv}</strong>លើស​ច្បាប់</div>
    <div class="hdr-stat"><strong>{cl}</strong>ប្រក្រតី</div>
  </div>
</div>

<div class="chips">
  <div class="chip c-total"><span class="chip-lbl">ព្រះសង្ឃ</span><span class="chip-val">{len(monks)}</span></div>
  <div class="chip c-abs"><span class="chip-lbl">❌ លើស​អវត្តមាន</span><span class="chip-val">{av}</span></div>
  <div class="chip c-perm"><span class="chip-lbl">📋 លើស​ច្បាប់</span><span class="chip-val">{pv}</span></div>
  <div class="chip c-ok"><span class="chip-lbl">✓ ប្រក្រតី</span><span class="chip-val">{cl}</span></div>
</div>

{section_html(bhikkhus,  'ផ្នែកទី ១ — ភិក្ខុ',   'sec-b', '#f0c040')}
{section_html(samaneras, 'ផ្នែកទី ២ — សាមណេរ', 'sec-s', '#66bb6a')}

<div class="footer">
  <span>វត្តនិរោធរង្សី — របាយការណ៍វត្តមាន ({_h.escape(type_label)})</span>
  <span>{_h.escape(subtitle)}</span>
</div>
</body></html>'''


# ---- Unified export: build excel ----------------------------------

def _make_export_excel(monks, type_label, subtitle, report_type):
    import io
    from datetime import date as _d
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter

    if report_type == 'daily':
        headers = ['#', 'ឈ្មោះ', 'ប្រភេទ', 'តួនាទី', 'វស្សា', 'ស្នាក់នៅ', 'ស្ថានភាព']
        widths  = [5, 24, 12, 20, 8, 18, 16]
    elif report_type == 'biweekly':
        headers = ['#', 'ឈ្មោះ', 'តួនាទី', 'វស្សា', 'ស្នាក់នៅ', 'ការសិក្សា', '❌', '📋', 'ថ្ងៃ', 'ស្ថានភាព']
        widths  = [5, 24, 20, 8, 18, 14, 6, 6, 24, 16]
    else:
        headers = ['#', 'ឈ្មោះ', 'ប្រភេទ', 'តួនាទី', 'វស្សា', '❌', '📋', 'ចន្លោះ', 'ស្ថានភាព']
        widths  = [5, 24, 12, 20, 8, 6, 6, 24, 16]

    n = len(headers)
    half = (n + 1) // 2

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = 'របាយការណ៍'

    blue  = Font(color='2B6CB0', size=10)
    bold  = Font(bold=True, size=11)
    bold_big = Font(bold=True, size=15)
    gray  = Font(color='718096', size=10)
    center = Alignment(horizontal='center', vertical='center', wrap_text=True)

    def merged(row, c1, c2, text, font, height=16):
        ws.merge_cells(start_row=row, start_column=c1, end_row=row, end_column=c2)
        cell = ws.cell(row=row, column=c1, value=text)
        cell.font = font
        cell.alignment = center
        ws.row_dimensions[row].height = height

    merged(1, 1, half,     '[ និមិត្តសញ្ញាក្រសួង ]', blue)
    merged(1, half + 1, n, _HDR_RIGHT_LINES[0], bold)
    merged(2, 1, half,     _HDR_LEFT_LINES[0], blue)
    merged(2, half + 1, n, _HDR_RIGHT_LINES[1], bold)
    merged(3, 1, half,     _HDR_LEFT_LINES[1], blue)
    merged(4, 1, half,     _HDR_LEFT_LINES[2], blue)
    merged(5, 1, n, '', Font())
    merged(6, 1, n, _HDR_MAIN_TITLE, bold_big, height=24)
    merged(7, 1, n, f"កាលបរិច្ឆេទ ជាចន្ទគតិ: {khmer_lunar_date(_d.today())}", gray)
    merged(8, 1, n, f"របាយការណ៍វត្តមាន ({type_label})  |  ចន្លោះ: {subtitle}", Font(bold=True, size=10))
    merged(9, 1, n, '', Font())

    hdr_row = 10
    hfill = PatternFill(start_color='4A5568', end_color='4A5568', fill_type='solid')
    hfont = Font(bold=True, color='FFFFFF', size=10)
    thin  = Side(border_style='thin', color='D1D5DB')
    bdr   = Border(left=thin, right=thin, top=thin, bottom=thin)

    for col, h in enumerate(headers, 1):
        c = ws.cell(row=hdr_row, column=col, value=h)
        c.font = hfont; c.fill = hfill; c.border = bdr; c.alignment = center
    ws.row_dimensions[hdr_row].height = 20

    row_i = hdr_row + 1
    for idx, m in enumerate(monks, 1):
        ab = m['absent_count']     >= DISC_ABSENT_MIN
        pr = m['permission_count'] >= DISC_PERM_MIN
        fill = PatternFill(start_color='FFF5F5' if ab else 'FFFAF0' if pr else 'FFFFFF',
                            end_color='FFF5F5' if ab else 'FFFAF0' if pr else 'FFFFFF', fill_type='solid')
        status_text = 'លើសអវត្តមាន' if ab else ('លើសច្បាប់' if pr else 'ប្រក្រតី')

        if report_type == 'daily':
            s = 'អវត្តមាន' if m.get('status') == 'absent' else 'ច្បាប់'
            vals = [idx, m['fullname'], m['monk_type'], m['position'], m['vassa_years'], m['residence'], s]
        elif report_type == 'biweekly':
            edu = f"{m['education_level']} {m['academic_year']}".strip()
            dp = []
            if m['absent_dates']: dp.append(f"❌ {m['absent_dates']}")
            if m['perm_dates']:   dp.append(f"📋 {m['perm_dates']}")
            vals = [idx, m['fullname'], m['position'], m['vassa_years'], m['residence'], edu,
                     m['absent_count'] or '—', m['permission_count'] or '—',
                     '\n'.join(dp) if dp else '—', status_text]
        else:
            vals = [idx, m['fullname'], m['monk_type'], m['position'], m['vassa_years'],
                     m['absent_count'] or '—', m['permission_count'] or '—',
                     m['perm_dates'] or '—', status_text]

        for col, val in enumerate(vals, 1):
            c = ws.cell(row=row_i, column=col, value=val)
            c.fill = fill; c.border = bdr
            c.font = Font(bold=True, size=10) if col == 2 else Font(size=10)
            c.alignment = Alignment(horizontal='center' if col != 2 else 'left',
                                     vertical='center', wrap_text=True)
        row_i += 1

    for col, w in enumerate(widths, 1):
        ws.column_dimensions[get_column_letter(col)].width = w

    buf = io.BytesIO(); wb.save(buf); buf.seek(0)
    return buf


# ---- Unified export route ----------------------------------------


@main_bp.route('/api/reports/submit-image', methods=['POST'])
def submit_report_image():
    try:
        import requests as req
        if 'image' not in request.files:
            return jsonify({'success': False, 'message': 'រកមិនឃើញរូបភាព'}), 400
            
        file = request.files['image']
        TELEGRAM_TOKEN   = '8950898077:AAHNR0tTgtJWy17wMXooKwg4nfQLGdfe5aw'
        TELEGRAM_CHAT_ID = -1003960014484
        
        caption = request.form.get('caption', "ប្រគេនរបាយការណ៍វត្តមាន (បំប្លែងជារូបភាព)")
        
        tg = req.post(
            f'https://api.telegram.org/bot{TELEGRAM_TOKEN}/sendPhoto',
            data={'chat_id': TELEGRAM_CHAT_ID, 'caption': caption},
            files={'photo': ('report.png', file.read(), 'image/png')},
            timeout=25
        ).json()
        
        if not tg.get('ok'):
            return jsonify({'success': False, 'message': f"Telegram: {tg.get('description')}"}), 500
            
        return jsonify({'success': True})
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 500

def _apply_report_filters(monks, args):
    """Apply the same monk_type/kuti/education/academic_year/name/violation
    filters the live report page applies client-side, so exports match what
    the user sees on screen."""
    monk_type       = (args.get('monk_type')       or '').strip()
    kuti            = (args.get('kuti')            or '').strip()
    education_level = (args.get('education_level') or '').strip()
    academic_year   = (args.get('academic_year')   or '').strip()
    name            = (args.get('name')            or '').strip().lower()
    violation       = args.get('violation', 'all')

    def keep(m):
        if monk_type       and m.get('monk_type')       != monk_type:                  return False
        if kuti            and m.get('residence')       != kuti.replace('_', ' '):     return False
        if education_level and m.get('education_level') != education_level:            return False
        if academic_year   and m.get('academic_year')   != academic_year:              return False
        if name            and name not in (m.get('fullname') or '').lower():          return False
        if violation == 'violations' and not (m['absent_count'] >= DISC_ABSENT_MIN or m['permission_count'] >= DISC_PERM_MIN):
            return False
        if violation == 'absent'     and not (m['absent_count']     >= DISC_ABSENT_MIN):
            return False
        if violation == 'permission' and not (m['permission_count'] >= DISC_PERM_MIN):
            return False
        return True

    return [m for m in monks if keep(m)]


@main_bp.route('/api/reports/export', methods=['GET'])
def unified_export():
    """Export any report type (daily/biweekly/monthly/annual/triennial) as docx or pdf."""
    try:
        import io, requests as req

        report_type = request.args.get('type',   'daily')
        fmt         = request.args.get('fmt',    'docx')
        action      = request.args.get('action', 'download')

        TELEGRAM_TOKEN   = '8950898077:AAHNR0tTgtJWy17wMXooKwg4nfQLGdfe5aw'
        TELEGRAM_CHAT_ID = -1003960014484

        monks, type_label, subtitle, period_start, period_end = _fetch_export_data(report_type, request.args)
        monks = _apply_report_filters(monks, request.args)

        av = sum(1 for m in monks if m['absent_count']     >= DISC_ABSENT_MIN)
        pv = sum(1 for m in monks if m['permission_count'] >= DISC_PERM_MIN)

        if fmt == 'pdf' or fmt == 'html':
            html = _make_export_html(monks, type_label, subtitle, report_type)
            html += '<script>window.onload = function() { setTimeout(function(){ window.print(); }, 500); }</script>'
            return html
        elif fmt == 'excel':
            buf      = _make_export_excel(monks, type_label, subtitle, report_type)
            fname    = f"report_{report_type}_{period_start}.xlsx"
            mimetype = 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        else:
            buf      = _make_export_docx(monks, type_label, subtitle, report_type)
            fname    = f"report_{report_type}_{period_start}.docx"
            mimetype = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'

        caption = (
            f"📋 របាយការណ៍វត្តមាន ({type_label}) — {subtitle}\n"
            f"📊 សរុប {len(monks)} នាក់  |  ❌ {av}  |  📋 {pv}"
        )

        if action == 'telegram':
            tg = req.post(
                f'https://api.telegram.org/bot{TELEGRAM_TOKEN}/sendDocument',
                data={'chat_id': TELEGRAM_CHAT_ID, 'caption': caption},
                files={'document': (fname, buf, mimetype)},
                timeout=25
            ).json()
            if not tg.get('ok'):
                return jsonify({'success': False, 'message': f"Telegram: {tg.get('description')}"}), 500
            return jsonify({'success': True, 'total': len(monks)})

        

        return send_file(buf, mimetype=mimetype, as_attachment=True, download_name=fname)

    except ImportError as e:
        return jsonify({'success': False, 'message': f'Missing dependency: {e}'}), 500
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 500
